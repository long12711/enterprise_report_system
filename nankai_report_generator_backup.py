"""
南开问卷专业报告生成器
根据问卷提交数据生成详细的企业专业评价报告（Word格式，自动转换为PDF）
完全按照"最终简化摘要_专业版报告.pdf"的格式和内容结构
"""

import os
import json
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    print("警告：docx2pdf库未安装，将只生成Word格式报告")


class NankaiReportGenerator:
    """南开问卷自评报告生成器"""
    
    def __init__(self, indicator_file='指标体系.xlsx'):
        self.indicator_file = indicator_file
        
    def generate_report(self, submission_file, generate_pdf=True):
        """
        根据提交的问卷数据生成专业评价报告
        
        Args:
            submission_file: 问卷提交JSON文件路径
            generate_pdf: 是否生成PDF格式（默认True）
            
        Returns:
            生成的报告文件路径（PDF或Word）
        """
        # 读取提交数据
        with open(submission_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # 提取关键信息
        enterprise_info = data.get('enterprise_info', {})
        answers = data.get('answers', {})
        partial_details = data.get('partial_details', {})
        score_info = data.get('score', {})
        level = data.get('level', '初级')
        
        enterprise_name = enterprise_info.get('enterprise_name', '企业')
        timestamp = data.get('timestamp', datetime.now().strftime('%Y%m%d_%H%M%S'))
        
        # 读取指标体系
        try:
            df = pd.read_excel(self.indicator_file, sheet_name=level)
        except:
            # 如果没有对应级别的sheet，读取第一个sheet
            df = pd.read_excel(self.indicator_file, sheet_name=0)
        
        # 创建Word文档
        doc = Document()
        
        # 设置文档样式
        self._set_document_style(doc)
        
        # 1. 封面
        self._add_cover(doc, enterprise_name, level, score_info)
        
        # 2. 执行摘要（新增）
        self._add_executive_summary(doc, enterprise_info, df, answers, score_info, level)
        
        # 3. 企业基本信息
        self._add_enterprise_info(doc, enterprise_info)
        
        # 4. 评价总览
        self._add_score_overview(doc, score_info, level)
        
        # 5. 各维度详细分析
        self._add_dimension_analysis(doc, df, answers, partial_details, score_info)
        
        # 6. 优势与不足
        self._add_strengths_weaknesses(doc, df, answers, score_info)
        
        # 7. 行业对比分析（新增）
        self._add_industry_comparison(doc, enterprise_info, score_info)
        
        # 8. 风险评估与预警（新增）
        self._add_risk_assessment(doc, enterprise_info, score_info)
        
        # 9. 改进路径规划（新增，替换原改进建议）
        self._add_improvement_plan(doc, enterprise_info, df, answers, score_info, level)
        
        # 10. 报告说明（新增）
        self._add_report_notes(doc, enterprise_info, level)
        
        # 11. 附录：问卷详细数据
        self._add_appendix(doc, df, answers, partial_details, score_info)
        
        # 保存Word报告
        output_dir = 'storage/nankai_reports'
        os.makedirs(output_dir, exist_ok=True)
        
        word_filename = f'专业评价报告_{enterprise_name}_{timestamp}.docx'
        word_path = os.path.join(output_dir, word_filename)
        
        doc.save(word_path)
        
        # 转换为PDF
        if generate_pdf and DOCX2PDF_AVAILABLE:
            try:
                pdf_filename = f'专业评价报告_{enterprise_name}_{timestamp}.pdf'
                pdf_path = os.path.join(output_dir, pdf_filename)
                convert(word_path, pdf_path)
                print(f"PDF报告生成成功: {pdf_path}")
                return pdf_path
            except Exception as e:
                print(f"PDF转换失败: {str(e)}，返回Word格式")
                return word_path
        else:
            return word_path
    
    def _set_document_style(self, doc):
        """设置文档样式"""
        # 设置默认字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.styles['Normal'].font.size = Pt(12)
    
    def _add_cover(self, doc, enterprise_name, level, score_info):
        """添加封面"""
        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('南开大学现代企业制度指数\n评价自评报告')
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        
        doc.add_paragraph()  # 空行
        
        # 企业名称
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'企业名称：{enterprise_name}')
        run.font.size = Pt(16)
        run.font.bold = True
        
        # 评价级别
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'评价级别：{level}')
        run.font.size = Pt(14)
        
        # 综合得分
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        score = score_info.get('percentage', 0)
        run = p.add_run(f'综合得分：{score:.2f}分')
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        
        # 生成日期
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
        run.font.size = Pt(12)
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc, enterprise_info, df, answers, score_info, level):
        """添加执行摘要（增强版 - 参考专业版报告格式）"""
        heading = doc.add_heading('执行摘要', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        enterprise_name = enterprise_info.get('enterprise_name', '企业')
        percentage = score_info.get('percentage', 0)
        total_score = score_info.get('total_score', 0)
        max_score = score_info.get('max_score', 100)
        main_business = enterprise_info.get('main_business', '相关行业')
        scale = enterprise_info.get('enterprise_scale', '')
        establishment_years = enterprise_info.get('establishment_years', '')
        
        # 评级判断
        if percentage >= 90:
            level_text = '优秀'
            level_desc = '企业现代制度建设达到卓越水平'
        elif percentage >= 80:
            level_text = '良好'
            level_desc = '企业现代制度建设较为完善'
        elif percentage >= 70:
            level_text = '中等'
            level_desc = '企业现代制度建设基本健全'
        elif percentage >= 60:
            level_text = '及格'
            level_desc = '企业现代制度建设有待提升'
        else:
            level_text = '需改进'
            level_desc = '企业现代制度建设亟需加强'
        
        # 评价结论段落
        p = doc.add_paragraph()
        p.add_run('一、评价结论\n').bold = True
        p.add_run(f'根据南开大学现代企业制度指数评价体系（{level}）的评价结果，')
        p.add_run(f'{enterprise_name}的综合得分为')
        run = p.add_run(f'{percentage:.1f}分')
        run.bold = True
        if percentage >= 80:
            run.font.color.rgb = RGBColor(0, 128, 0)
        elif percentage >= 60:
            run.font.color.rgb = RGBColor(255, 152, 0)
        else:
            run.font.color.rgb = RGBColor(244, 67, 54)
        p.add_run(f'（实际得分{total_score:.1f}分，满分{max_score:.1f}分），')
        p.add_run(f'评价等级为')
        run = p.add_run(level_text)
        run.bold = True
        p.add_run(f'。{level_desc}，')
        p.add_run(f'各项制度体系建设总体呈现{"稳步提升" if percentage >= 70 else "持续改进"}的良好态势。')
        
        doc.add_paragraph()
        
        # 企业概况段落（详细版）
        p = doc.add_paragraph()
        p.add_run('二、企业概况\n').bold = True
        scale_text = f'{scale}规模' if scale else ''
        years_text = f'，成立{establishment_years}年' if establishment_years else ''
        p.add_run(f'{enterprise_name}是一家{scale_text}企业{years_text}，')
        p.add_run(f'主要从事{main_business}相关业务。')
        p.add_run(f'企业自成立以来，始终坚持规范化经营理念，')
        p.add_run(f'积极推进现代企业制度建设，')
        p.add_run(f'在公司治理、内部控制、风险管理等方面持续完善，')
        p.add_run(f'为企业的可持续发展奠定了坚实的制度基础。')
        
        doc.add_paragraph()
        
        # 分析维度得分
        dimensions = self._analyze_dimensions(df, answers, score_info)
        
        # 评价亮点（详细版）
        p = doc.add_paragraph()
        p.add_run('三、评价亮点\n').bold = True
        
        # 找出优势维度（得分率>=75%）
        strengths = [(dim, data) for dim, data in dimensions.items() if data['percentage'] >= 75]
        strengths.sort(key=lambda x: x[1]['percentage'], reverse=True)
        
        if strengths:
            p.add_run(f'本次评价中，{enterprise_name}在以下方面表现突出：\n')
            for idx, (dim_name, dim_data) in enumerate(strengths[:3], 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(f'{dim_name}：').bold = True
                p.add_run(f'得分率达到{dim_data["percentage"]:.1f}%，')
                p.add_run(f'显著高于平均水平。')
                p.add_run(f'该维度的优异表现体现了企业在{dim_name}方面的制度建设水平较高，')
                p.add_run(f'相关制度健全完善，执行情况良好，')
                p.add_run(f'为企业的规范化运营提供了有力保障。')
        else:
            # 如果没有特别突出的维度，给出通用亮点
            p.add_run(f'{enterprise_name}在本次评价中展现出以下积极特点：\n')
            
            p = doc.add_paragraph(style='List Number')
            p.add_run('制度体系基本健全：').bold = True
            p.add_run(f'企业已建立起较为完整的现代企业制度框架，')
            p.add_run(f'基础治理架构清晰，为企业规范运营提供了制度保障。')
            
            p = doc.add_paragraph(style='List Number')
            p.add_run('持续改进意识：').bold = True
            p.add_run(f'企业在{level}评价中表现{level_text}，')
            p.add_run(f'体现了企业对制度建设的重视和持续改进的决心。')
            
            # 找出相对较好的维度
            if dimensions:
                top_dims = sorted(dimensions.items(), key=lambda x: x[1]['percentage'], reverse=True)[:2]
                for dim_name, dim_data in top_dims:
                    if dim_data['percentage'] >= 60:
                        p = doc.add_paragraph(style='List Number')
                        p.add_run(f'{dim_name}基础良好：').bold = True
                        p.add_run(f'该维度得分率{dim_data["percentage"]:.1f}%，')
                        p.add_run(f'相关制度建设完成情况较好，具备进一步提升的基础。')
        
        doc.add_paragraph()
        
        # 存在问题（详细版）
        p = doc.add_paragraph()
        p.add_run('四、存在问题\n').bold = True
        
        # 找出薄弱维度（得分率<65%）
        weaknesses = [(dim, data) for dim, data in dimensions.items() if data['percentage'] < 65]
        weaknesses.sort(key=lambda x: x[1]['percentage'])
        
        if weaknesses:
            p.add_run(f'评价结果显示，{enterprise_name}在以下方面仍存在改进空间：\n')
            for idx, (dim_name, dim_data) in enumerate(weaknesses[:3], 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(f'{dim_name}有待加强：').bold = True
                p.add_run(f'该维度得分率为{dim_data["percentage"]:.1f}%，')
                p.add_run(f'低于整体平均水平。')
                p.add_run(f'建议企业重点关注{dim_name}相关制度的建立和完善，')
                p.add_run(f'系统性地梳理现有制度，补齐短板，')
                p.add_run(f'提升该领域的制度化管理水平。')
        else:
            p.add_run(f'总体而言，{enterprise_name}各维度表现较为均衡，')
            p.add_run(f'未发现明显的薄弱环节。')
            p.add_run(f'建议企业在保持现有水平的基础上，')
            p.add_run(f'持续优化各项制度，提升制度执行的一致性和有效性，')
            p.add_run(f'推动企业治理水平再上新台阶。')
        
        doc.add_paragraph()
        
        # 总体建议
        p = doc.add_paragraph()
        p.add_run('五、总体建议\n').bold = True
        
        if percentage >= 85:
            p.add_run(f'{enterprise_name}的现代制度建设已达到较高水平，')
            p.add_run(f'建议继续保持并不断优化现有制度体系，')
            p.add_run(f'关注行业最新发展趋势和监管要求，')
            p.add_run(f'持续提升制度执行效果，')
            p.add_run(f'巩固和扩大制度建设成果。')
        elif percentage >= 70:
            p.add_run(f'{enterprise_name}的现代制度建设整体良好，')
            p.add_run(f'建议针对薄弱环节进行重点改进，')
            p.add_run(f'完善制度细节，提高制度执行的一致性和有效性，')
            p.add_run(f'推动企业治理水平持续提升。')
        elif percentage >= 60:
            p.add_run(f'{enterprise_name}已建立基本的制度框架，')
            p.add_run(f'建议系统梳理现有制度，补齐短板，')
            p.add_run(f'加强制度培训和执行监督，')
            p.add_run(f'提升整体制度化管理水平。')
        else:
            p.add_run(f'{enterprise_name}的制度建设亟需加强，')
            p.add_run(f'建议制定系统的制度建设规划，')
            p.add_run(f'分阶段推进各项制度的建立和完善，')
            p.add_run(f'必要时可聘请专业咨询机构提供指导，')
            p.add_run(f'从基础制度开始逐步建立健全现代企业制度体系。')
        
        doc.add_paragraph()
        doc.add_page_break()
    
    def _analyze_dimensions(self, df, answers, score_info):
        """分析各维度得分情况"""
        dimensions = {}
        score_details = score_info.get('details', {})
        
        for idx, row in df.iterrows():
            level1 = str(row['一级指标']) if pd.notna(row['一级指标']) else ''
            if not level1:
                continue
            
            if level1 not in dimensions:
                dimensions[level1] = {
                    'total_score': 0,
                    'earned_score': 0,
                    'percentage': 0
                }
            
            q_id = str(int(row['序号'])) if pd.notna(row['序号']) else str(idx + 1)
            detail = score_details.get(q_id, {})
            
            dimensions[level1]['total_score'] += detail.get('max_score', 0)
            dimensions[level1]['earned_score'] += detail.get('earned_score', 0)
        
        # 计算百分比
        for dim_name, dim_data in dimensions.items():
            if dim_data['total_score'] > 0:
                dim_data['percentage'] = (dim_data['earned_score'] / dim_data['total_score']) * 100
        
        return dimensions
    
    def _add_enterprise_info(self, doc, enterprise_info):
        """添加企业基本信息"""
        heading = doc.add_heading('一、企业基本信息', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        # 创建表格
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        info_items = [
            ('企业名称', enterprise_info.get('enterprise_name', '-')),
            ('联系人', enterprise_info.get('contact_person', '-')),
            ('联系电话', enterprise_info.get('contact_phone', '-')),
            ('联系邮箱', enterprise_info.get('contact_email', '-')),
            ('主营业务', enterprise_info.get('main_business', '-')),
            ('企业规模', enterprise_info.get('enterprise_scale', '-'))
        ]
        
        for i, (label, value) in enumerate(info_items):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
        
        doc.add_paragraph()
    
    def _add_score_overview(self, doc, score_info, level):
        """添加评价总览"""
        heading = doc.add_heading('二、评价总览', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        total_score = score_info.get('total_score', 0)
        max_score = score_info.get('max_score', 100)
        percentage = score_info.get('percentage', 0)
        
        # 得分情况
        p = doc.add_paragraph()
        p.add_run('评价结果：').bold = True
        p.add_run(f'贵企业在{level}问卷评价中，')
        p.add_run(f'实际得分{total_score:.2f}分，满分{max_score:.2f}分，')
        p.add_run(f'百分制得分为{percentage:.2f}分。')
        
        # 评级
        if percentage >= 90:
            level_text = '优秀'
            level_desc = '企业现代制度建设达到卓越水平，各项制度健全完善，执行有力。'
        elif percentage >= 80:
            level_text = '良好'
            level_desc = '企业现代制度建设较为完善，主要制度体系健全，执行情况良好。'
        elif percentage >= 70:
            level_text = '中等'
            level_desc = '企业现代制度建设基本健全，部分制度需要进一步完善。'
        elif percentage >= 60:
            level_text = '及格'
            level_desc = '企业现代制度建设有待提升，需要加强制度建设和执行力度。'
        else:
            level_text = '需改进'
            level_desc = '企业现代制度建设亟需加强，建议系统性地完善各项制度。'
        
        p = doc.add_paragraph()
        p.add_run('评价等级：').bold = True
        run = p.add_run(level_text)
        run.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        
        p = doc.add_paragraph()
        p.add_run('总体评价：').bold = True
        p.add_run(level_desc)
        
        doc.add_paragraph()
    
    def _add_dimension_analysis(self, doc, df, answers, partial_details, score_info):
        """添加各维度详细分析"""
        heading = doc.add_heading('三、各维度详细分析', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        # 按一级指标分组
        dimensions = {}
        score_details = score_info.get('details', {})
        
        for idx, row in df.iterrows():
            level1 = str(row['一级指标']) if pd.notna(row['一级指标']) else ''
            if not level1:
                continue
            
            if level1 not in dimensions:
                dimensions[level1] = {
                    'questions': [],
                    'total_score': 0,
                    'earned_score': 0
                }
            
            q_id = str(int(row['序号'])) if pd.notna(row['序号']) else str(idx + 1)
            detail = score_details.get(q_id, {})
            
            dimensions[level1]['questions'].append({
                'id': q_id,
                'level2': str(row['二级指标']) if pd.notna(row['二级指标']) else '',
                'level3': str(row['三级指标']) if pd.notna(row['三级指标']) else '',
                'answer': answers.get(q_id, ''),
                'partial_detail': partial_details.get(q_id, ''),
                'max_score': detail.get('max_score', 0),
                'earned_score': detail.get('earned_score', 0)
            })
            
            dimensions[level1]['total_score'] += detail.get('max_score', 0)
            dimensions[level1]['earned_score'] += detail.get('earned_score', 0)
        
        # 输出各维度分析
        for dim_name, dim_data in dimensions.items():
            doc.add_heading(dim_name, level=2)
            
            total = dim_data['total_score']
            earned = dim_data['earned_score']
            percentage = (earned / total * 100) if total > 0 else 0
            
            p = doc.add_paragraph()
            p.add_run(f'维度得分：{earned:.2f}/{total:.2f}分（{percentage:.1f}%）')
            p.runs[0].bold = True
            
            # 统计完成情况
            completed = sum(1 for q in dim_data['questions'] if q['answer'].startswith('A'))
            partial = sum(1 for q in dim_data['questions'] if q['answer'].startswith('B'))
            not_completed = sum(1 for q in dim_data['questions'] if q['answer'].startswith('C'))
            
            p = doc.add_paragraph()
            p.add_run(f'完成情况：已完成{completed}项，部分完成{partial}项，未完成{not_completed}项')
            
            doc.add_paragraph()
    
    def _add_strengths_weaknesses(self, doc, df, answers, score_info):
        """添加优势与不足分析"""
        heading = doc.add_heading('四、优势与不足', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        score_details = score_info.get('details', {})
        
        # 找出得分高的项目（优势）
        strengths = []
        weaknesses = []
        
        for idx, row in df.iterrows():
            q_id = str(int(row['序号'])) if pd.notna(row['序号']) else str(idx + 1)
            detail = score_details.get(q_id, {})
            answer = answers.get(q_id, '')
            
            level3 = str(row['三级指标']) if pd.notna(row['三级指标']) else ''
            
            if answer.startswith('A'):  # 已完成
                strengths.append(level3)
            elif answer.startswith('C'):  # 未完成
                weaknesses.append(level3)
        
        # 优势
        doc.add_heading('（一）主要优势', level=2)
        if strengths:
            p = doc.add_paragraph()
            p.add_run('贵企业在以下方面表现优秀：')
            for strength in strengths[:10]:  # 最多显示10项
                doc.add_paragraph(strength, style='List Bullet')
        else:
            doc.add_paragraph('暂无突出优势项。')
        
        # 不足
        doc.add_heading('（二）主要不足', level=2)
        if weaknesses:
            p = doc.add_paragraph()
            p.add_run('贵企业在以下方面需要改进：')
            for weakness in weaknesses[:10]:  # 最多显示10项
                doc.add_paragraph(weakness, style='List Bullet')
        else:
            doc.add_paragraph('各项指标均已达标。')
        
        doc.add_paragraph()
    
    def _add_industry_comparison(self, doc, enterprise_info, score_info):
        """添加行业对比分析"""
        heading = doc.add_heading('七、行业对比分析', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        enterprise_name = enterprise_info.get('enterprise_name', '企业')
        percentage = score_info.get('percentage', 0)
        industry = enterprise_info.get('main_business', '相关行业')
        
        # （一）行业基准对比
        doc.add_heading('（一）行业基准对比', level=2)
        
        p = doc.add_paragraph()
        p.add_run('行业发展态势：').bold = True
        p = doc.add_paragraph()
        p.add_run(f'{industry}作为重要的经济组成部分，近年来保持了稳定发展态势。')
        p.add_run(f'行业企业数量众多，市场竞争激烈，对现代企业制度建设的要求不断提高。')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('综合对比分析结论：').bold = True
        
        # 根据得分率判断行业地位
        if percentage >= 90:
            position = '优秀水平'
            description = f'{enterprise_name}在多个关键维度上展现出卓越的表现水平，已达到行业领先地位。'
        elif percentage >= 80:
            position = '良好水平'
            description = f'{enterprise_name}在多个关键维度上表现良好，超过行业平均水平，具备进一步提升的基础。'
        elif percentage >= 70:
            position = '中等水平'
            description = f'{enterprise_name}的制度建设水平处于行业中等水平，部分维度需要重点改进。'
        else:
            position = '有待提升'
            description = f'{enterprise_name}的制度建设水平有待提升，需要系统性地完善各项制度。'
        
        p = doc.add_paragraph()
        p.add_run(description)
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('行业地位分析：').bold = True
        p = doc.add_paragraph()
        p.add_run(f'{enterprise_name}现代企业制度建设水平在{industry}中处于{position}，')
        if percentage >= 85:
            p.add_run('具备冲击行业标杆企业的实力基础。')
        elif percentage >= 70:
            p.add_run('具备持续提升的良好基础。')
        else:
            p.add_run('需要加大制度建设力度。')
        
        doc.add_paragraph()
        doc.add_page_break()
    
    def _add_risk_assessment(self, doc, enterprise_info, score_info):
        """添加风险评估与预警"""
        heading = doc.add_heading('九、风险评估与预警', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        enterprise_name = enterprise_info.get('enterprise_name', '企业')
        percentage = score_info.get('percentage', 0)
        
        # （一）风险等级评估
        doc.add_heading('（一）风险等级评估', level=2)
        
        # 根据得分率判断风险等级
        if percentage >= 85:
            risk_level = '低风险'
            risk_score = 85 + (percentage - 85) / 15 * 15
        elif percentage >= 70:
            risk_level = '中低风险'
            risk_score = 70 + (percentage - 70) / 15 * 15
        elif percentage >= 60:
            risk_level = '中等风险'
            risk_score = 60 + (percentage - 60) / 10 * 10
        else:
            risk_level = '较高风险'
            risk_score = percentage
        
        p = doc.add_paragraph()
        run = p.add_run(f'综合风险等级：{risk_level}')
        run.bold = True
        
        p = doc.add_paragraph()
        p.add_run(f'通过运用现代风险管理理论和方法，采用定性与定量相结合的风险评估模型，')
        p.add_run(f'对{enterprise_name}进行全方位风险识别、评估和分析。')
        p.add_run(f'评估结果显示，企业总体风险等级为{risk_level}，')
        p.add_run(f'综合风险得分率为{risk_score:.1f}分（满分100分），')
        level_text = '先进' if percentage >= 85 else '中等' if percentage >= 70 else '基础'
        p.add_run(f'风险管控能力达到{enterprise_info.get("main_business", "相关行业")}{level_text}水平。')
        
        doc.add_paragraph()
        
        # （二）分类风险深度分析
        doc.add_heading('（二）分类风险深度分析', level=2)
        
        # 风险分类表格
        table = doc.add_table(rows=5, cols=4)
        table.style = 'Light Grid Accent 1'
        
        headers = ['风险类别', '风险等级', '主要风险点', '建议改进措施']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        risk_data = [
            ['治理风险', '低风险' if percentage >= 80 else '中等风险',
             '部分制度执行监督需加强', '建立制度执行监督评价机制'],
            ['财务风险', '低风险' if percentage >= 75 else '中等风险',
             '财务管理制度需进一步完善', '完善财务管理制度体系'],
            ['运营风险', '低风险' if percentage >= 70 else '中等风险',
             '业务流程标准化程度需提升', '推进业务流程标准化'],
            ['合规风险', '低风险' if percentage >= 75 else '中等风险',
             '合规培训覆盖面需扩大', '扩大合规培训范围']
        ]
        
        for i, row_data in enumerate(risk_data, 1):
            for j, cell_data in enumerate(row_data):
                table.rows[i].cells[j].text = cell_data
        
        doc.add_paragraph()
        
        # （三）风险预警机制
        doc.add_heading('（三）风险预警机制', level=2)
        
        p = doc.add_paragraph()
        p.add_run('三级风险预警体系：').bold = True
        
        p = doc.add_paragraph()
        p.add_run('建议建立"绿色-黄色-红色"三级风险预警机制，设置关键风险监测指标，实现风险的动态监测和预警。')
        
        p = doc.add_paragraph()
        p.add_run('绿色预警（正常状态）：风险指标在安全范围内，按常规管理流程执行。')
        
        p = doc.add_paragraph()
        p.add_run('黄色预警（关注状态）：风险指标接近预警线，启动重点监控程序，制定应对预案。')
        
        p = doc.add_paragraph()
        p.add_run('红色预警（警戒状态）：风险指标超出安全范围，启动应急响应机制，采取紧急管控措施。')
        
        doc.add_paragraph()
        doc.add_page_break()
    
    def _add_improvement_plan(self, doc, enterprise_info, df, answers, score_info, level):
        """添加改进路径规划"""
        heading = doc.add_heading('十、改进路径规划', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        enterprise_name = enterprise_info.get('enterprise_name', '企业')
        percentage = score_info.get('percentage', 0)
        industry = enterprise_info.get('main_business', '相关行业')
        
        p = doc.add_paragraph()
        p.add_run('路径规划导言：').bold = True
        
        p = doc.add_paragraph()
        intro_text = f'基于本次现代企业制度指数评价结果，{enterprise_name}在治理体系建设方面'
        if percentage >= 75:
            intro_text += '展现出了良好的发展态势'
        else:
            intro_text += '具备一定的基础'
        intro_text += f'，总体完成度达{percentage:.1f}%，'
        if percentage >= 85:
            intro_text += f'在{industry}中处于优秀水平。'
        elif percentage >= 70:
            intro_text += f'在{industry}中处于良好水平。'
        else:
            intro_text += f'在{industry}中处于中等水平。'
        intro_text += '为确保企业治理水平持续提升，我们制定了分阶段、有重点、可操作的改进路径规划。'
        p.add_run(intro_text)
        
        doc.add_paragraph()
        
        # （一）短期改进计划
        doc.add_heading('（一）短期改进计划（3-6个月）', level=2)
        
        p = doc.add_paragraph()
        run = p.add_run('阶段目标：')
        run.bold = True
        p.add_run('夯实基础，补齐短板，建立健全基本制度框架')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('1. 制度体系完善工程').bold = True
        
        short_items = [
            '全面梳理现有治理制度，建立制度清单和制度地图，识别制度缺口和薄弱环节',
            '重点制定和完善核心制度，如《企业风险管理办法》、《信息披露管理制度》等',
            '建立制度定期评估机制，每季度开展制度执行效果评估',
            '完善制度执行监督机制，明确制度执行责任主体，建立制度违反问责机制'
        ]
        
        for item in short_items:
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('2. 治理透明度提升工程').bold = True
        
        transparency_items = [
            '建立企业信息披露专栏，确保重大信息及时披露',
            '完善内部信息传递机制，建立信息披露责任制',
            '规范关联交易信息披露，建立关联交易数据库',
            '加强与利益相关方的沟通，定期举办说明会、座谈会等'
        ]
        
        for item in transparency_items:
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('预期成果：').bold = True
        p.add_run('形成较为完整的制度框架，补齐制度短板，建立风险管控基础设施，')
        p.add_run('信息披露及时性达到95%以上，为下一阶段深化改革奠定坚实基础。')
        if percentage < 90:
            p.add_run(f'短期内企业治理完成度预计提升至{min(percentage + 5, 96):.0f}%以上。')
        
        doc.add_paragraph()
        
        # （二）中期发展规划
        doc.add_heading('（二）中期发展规划（6-18个月）', level=2)
        
        p = doc.add_paragraph()
        run = p.add_run('阶段目标：')
        run.bold = True
        p.add_run('优化体系，提升效能，建立现代化治理运行机制')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('1. 治理效能优化工程').bold = True
        
        mid_items = [
            '优化决策流程，建立科学决策机制，决策效率提升30%',
            '推进治理数字化转型，建设智能化治理平台',
            '完善绩效考核体系，将现代企业制度建设纳入考核指标',
            '加强内部审计功能，实现对重点领域和关键环节的全覆盖审计'
        ]
        
        for item in mid_items:
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('预期成果：').bold = True
        p.add_run('建立高效的现代化治理运行机制，治理效率提升30%，')
        p.add_run('人才队伍专业化水平显著提升，利益相关方满意度达到90%以上，')
        if percentage < 95:
            p.add_run(f'企业治理完成度达到{min(percentage + 10, 98):.0f}%以上。')
        
        doc.add_paragraph()
        
        # （三）长期目标
        doc.add_heading('（三）长期目标（1-3年）', level=2)
        
        p = doc.add_paragraph()
        run = p.add_run('战略目标：')
        run.bold = True
        p.add_run('全面建成现代企业制度体系，实现治理能力现代化')
        
        doc.add_paragraph()
        
        long_items = [
            '建成完善的现代企业制度体系，实现治理完成度99%以上',
            '实现治理规范化、制度化、程序化，各项治理活动有章可循',
            f'达到{industry}领先的治理水平，成为同行业治理标杆',
            '建立与国际先进标准接轨的治理体系',
            '形成良好的企业治理文化，治理理念深入人心'
        ]
        
        for item in long_items:
            doc.add_paragraph(item, style='List Number')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('预期成果：').bold = True
        p.add_run('全面建成现代企业制度体系，治理能力达到国际先进水平，')
        p.add_run(f'成为{industry}治理标杆企业，为企业高质量发展提供坚强保障。')
        p.add_run('企业价值和社会责任实现有机统一，可持续发展能力显著增强。')
        
        doc.add_paragraph()
        doc.add_page_break()
    
    def _add_report_notes(self, doc, enterprise_info, level):
        """添加报告说明"""
        heading = doc.add_heading('十二、报告说明', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        # 1. 报告性质与用途
        p = doc.add_paragraph()
        p.add_run('1. 报告性质与用途').bold = True
        
        p = doc.add_paragraph()
        p.add_run('本报告是基于企业自主填报数据生成的专业评价报告，')
        p.add_run('采用南开大学现代企业制度指数评价体系标准。')
        p.add_run('报告反映了企业在现代企业制度建设方面的现状水平，')
        p.add_run('为企业治理改进提供参考依据。')
        
        doc.add_paragraph()
        
        # 2. 评价方法与依据
        p = doc.add_paragraph()
        p.add_run('2. 评价方法与依据').bold = True
        
        p = doc.add_paragraph()
        p.add_run('评价标准：南开大学现代企业制度指数评价体系初稿')
        
        p = doc.add_paragraph()
        p.add_run('评价方法：定量评分与定性分析相结合')
        
        p = doc.add_paragraph()
        p.add_run('评价维度：治理方向性、治理有效性、治理规范性、治理透明性')
        
        p = doc.add_paragraph()
        p.add_run(f'指标体系：{level}级别包含相应的评价指标')
        
        doc.add_paragraph()
        
        # 3. 报告特色与创新
        p = doc.add_paragraph()
        p.add_run('3. 报告特色与创新').bold = True
        
        features = [
            '专业格式：采用正式报告格式，符合官方文档标准',
            '丰富内容：包含行业对比、风险评估、改进路径等增值分析',
            '智能分析：基于数据可用性动态生成相关章节',
            '实用建议：提供针对性的改进建议和实施路径'
        ]
        
        for feature in features:
            doc.add_paragraph(feature, style='List Bullet')
        
        doc.add_paragraph()
        
        # 4. 使用范围与限制
        p = doc.add_paragraph()
        p.add_run('4. 使用范围与限制').bold = True
        
        limitations = [
            '本报告仅供企业内部管理参考使用',
            '不构成对外投资决策或融资担保的依据',
            '评价结果基于企业自主填报数据，请确保数据真实性',
            '建议结合企业实际情况制定具体改进措施'
        ]
        
        for limitation in limitations:
            doc.add_paragraph(limitation, style='List Bullet')
        
        doc.add_paragraph()
        
        # 5. 技术支持与联系
        p = doc.add_paragraph()
        p.add_run('5. 技术支持与联系').bold = True
        
        p = doc.add_paragraph()
        p.add_run('系统版本：专业标准版报告生成器 v2.0')
        
        p = doc.add_paragraph()
        p.add_run('技术支持：现代企业制度指数评价系统')
        
        p = doc.add_paragraph()
        p.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
        
        p = doc.add_paragraph()
        contact_person = enterprise_info.get('contact_person', '')
        contact_email = enterprise_info.get('contact_email', '')
        if contact_person:
            p.add_run(f'企业联系人：{contact_person}')
        if contact_email:
            p = doc.add_paragraph()
            p.add_run(f'联系方式：{contact_email}')
        
        doc.add_paragraph()
        
        # 分隔线
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run('_' * 60)
        
        doc.add_paragraph()
        
        # 底部标识
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('现代企业制度指数评价系统 | 专业标准版报告')
        run.font.size = Pt(10)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Modern Enterprise System Index Evaluation System - Professional Standard Report')
        run.font.size = Pt(9)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('基于南开大学学术标准 | 专业报告生成器')
        run.font.size = Pt(9)
        
        doc.add_paragraph()
        doc.add_page_break()
    
    def _add_appendix(self, doc, df, answers, partial_details, score_info):
        """添加附录：问卷详细数据"""
        heading = doc.add_heading('附录：问卷详细数据', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        score_details = score_info.get('details', {})
        
        # 创建表格
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        headers = ['序号', '三级指标', '分值', '答案', '得分', '完成度', '说明']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # 数据行
        for idx, row in df.iterrows():
            q_id = str(int(row['序号'])) if pd.notna(row['序号']) else str(idx + 1)
            level3 = str(row['三级指标']) if pd.notna(row['三级指标']) else ''
            answer = answers.get(q_id, '')
            partial_detail = partial_details.get(q_id, '')
            
            detail = score_details.get(q_id, {})
            max_score = detail.get('max_score', 0)
            earned_score = detail.get('earned_score', 0)
            
            # 完成度
            if answer.startswith('A'):
                completion = '已完成'
            elif answer.startswith('B'):
                completion = '部分完成'
            else:
                completion = '未完成'
            
            row_cells = table.add_row().cells
            row_cells[0].text = q_id
            row_cells[1].text = level3
            row_cells[2].text = f'{max_score:.1f}'
            row_cells[3].text = answer
            row_cells[4].text = f'{earned_score:.1f}'
            row_cells[5].text = completion
            row_cells[6].text = partial_detail
        
        doc.add_paragraph()
        
        # 结束语
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run('—— 报告结束 ——')


if __name__ == '__main__':
    # 测试代码
    generator = NankaiReportGenerator()
    print("南开问卷自评报告生成器已就绪")