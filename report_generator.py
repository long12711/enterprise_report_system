"""
报告生成模块
根据企业填报信息生成自评报告
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os


class ReportGenerator:
    """自评报告生成器"""

    def __init__(self):
        self.report_template = None

    def generate_report(self, enterprise_data, output_folder='reports'):
        """
        生成企业自评报告

        Args:
            enterprise_data: 企业填报数据（字典）
            output_folder: 输出文件夹路径

        Returns:
            生成的报告文件路径
        """
        # 创建Word文档
        doc = Document()

        # 设置文档标题
        title = doc.add_heading('', level=0)
        title_run = title.add_run('中国特色现代企业制度评价')
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading('', level=1)
        subtitle_run = subtitle.add_run('企业自评报告')
        subtitle_run.font.size = Pt(18)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # 一、企业基本信息
        doc.add_heading('一、企业基本信息', level=1)

        # 创建基本信息表格
        basic_info_fields = [
            ('企业名称', '企业名称'),
            ('统一社会信用代码', '统一社会信用代码'),
            ('成立时间', '成立时间'),
            ('注册资本', '注册资本'),
            ('企业类型', '企业类型'),
            ('所属行业', '所属行业'),
            ('员工人数', '员工人数'),
            ('联系人', '联系人'),
            ('联系电话', '联系人手机'),
            ('联系邮箱', '联系人邮箱'),
        ]

        table = doc.add_table(rows=len(basic_info_fields), cols=2)
        table.style = 'Light Grid Accent 1'

        for idx, (label, key) in enumerate(basic_info_fields):
            row = table.rows[idx]
            row.cells[0].text = label
            row.cells[1].text = str(enterprise_data.get(key, ''))

        doc.add_paragraph()

        # 二、评价指标体系
        doc.add_heading('二、评价指标自评', level=1)

        # 获取所有评价指标
        indicators = self._get_indicators(enterprise_data)

        if indicators:
            for category, items in indicators.items():
                doc.add_heading(category, level=2)

                for item in items:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"{item['name']}: ").bold = True
                    p.add_run(f"{item['value']}")

                    if item.get('description'):
                        desc = doc.add_paragraph(style='List Bullet 2')
                        desc.add_run(f"说明: {item['description']}")

                doc.add_paragraph()

        # 三、评分汇总
        doc.add_heading('三、评分汇总', level=1)

        scores = self._calculate_scores(enterprise_data)

        if scores:
            score_table = doc.add_table(rows=len(scores) + 2, cols=3)
            score_table.style = 'Light Grid Accent 1'

            # 表头
            header_cells = score_table.rows[0].cells
            header_cells[0].text = '评价维度'
            header_cells[1].text = '得分'
            header_cells[2].text = '权重'

            # 明细
            for idx, (dimension, score, weight) in enumerate(scores, 1):
                row = score_table.rows[idx]
                row.cells[0].text = dimension
                row.cells[1].text = str(score)
                row.cells[2].text = f"{weight}%"

            # 总分
            total_row = score_table.rows[-1]
            total_row.cells[0].text = '综合得分'
            total_score = sum(s[1] * s[2] / 100 for s in scores)
            total_row.cells[1].text = f"{total_score:.2f}"
            total_row.cells[2].text = '100%'

            # 加粗总分行
            for cell in total_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

        doc.add_paragraph()

        # 四、评价结论
        doc.add_heading('四、评价结论', level=1)

        conclusion = self._generate_conclusion(enterprise_data, scores)
        doc.add_paragraph(conclusion)

        doc.add_paragraph()

        # 五、改进建议
        doc.add_heading('五、改进建议', level=1)

        suggestions = self._generate_suggestions(enterprise_data, scores)
        for suggestion in suggestions:
            doc.add_paragraph(suggestion, style='List Bullet')

        doc.add_paragraph()

        # 报告尾部
        footer = doc.add_paragraph()
        footer.add_run(f'\n报告生成时间: {datetime.now().strftime("%Y年%m月%d日")}')
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 保存文档
        os.makedirs(output_folder, exist_ok=True)
        enterprise_name = enterprise_data.get('企业名称', 'unknown')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{enterprise_name}_自评报告_{timestamp}.docx"

        # 清理文件名中的非法字符
        filename = self._clean_filename(filename)

        filepath = os.path.join(output_folder, filename)
        doc.save(filepath)

        return filepath

    def _get_indicators(self, data):
        """
        提取评价指标
        """
        indicators = {}

        # 从数据中提取所有以特定前缀开头的字段
        # 假设指标字段命名规则为: "维度-指标名称"

        for key, value in data.items():
            if '-' in str(key):
                parts = str(key).split('-', 1)
                if len(parts) == 2:
                    category, indicator_name = parts

                    if category not in indicators:
                        indicators[category] = []

                    indicators[category].append({
                        'name': indicator_name,
                        'value': value,
                        'description': ''
                    })

        # 如果没有找到指标,返回默认示例
        if not indicators:
            indicators = {
                '党建引领': [
                    {'name': '党组织建设情况', 'value': data.get('党组织建设情况', '未填写'), 'description': ''},
                ],
                '治理结构': [
                    {'name': '公司治理结构', 'value': data.get('公司治理结构', '未填写'), 'description': ''},
                ],
                '运营机制': [
                    {'name': '经营管理机制', 'value': data.get('经营管理机制', '未填写'), 'description': ''},
                ],
                '监督机制': [
                    {'name': '监督体系建设', 'value': data.get('监督体系建设', '未填写'), 'description': ''},
                ],
            }

        return indicators

    def _calculate_scores(self, data):
        """
        计算各维度得分
        返回: [(维度名, 得分, 权重)]
        """
        # 这里是示例逻辑,实际需要根据评价标准计算
        scores = []

        dimensions = {
            '党建引领': 25,
            '治理结构': 25,
            '运营机制': 25,
            '监督机制': 25,
        }

        for dimension, weight in dimensions.items():
            # 简化计算:从数据中查找相关字段
            score = self._calculate_dimension_score(data, dimension)
            scores.append((dimension, score, weight))

        return scores

    def _calculate_dimension_score(self, data, dimension):
        """
        计算单个维度得分
        """
        # 简化逻辑:根据填写情况给分
        score = 0
        relevant_fields = 0
        filled_fields = 0

        for key, value in data.items():
            if dimension in str(key):
                relevant_fields += 1
                if value and str(value).strip() not in ['', 'nan', 'None', '未填写']:
                    filled_fields += 1

        if relevant_fields > 0:
            score = (filled_fields / relevant_fields) * 100
        else:
            score = 60  # 默认基准分

        return round(score, 2)

    def _generate_conclusion(self, data, scores):
        """
        生成评价结论
        """
        if not scores:
            return "根据企业提供的信息,无法进行全面评价。"

        total_score = sum(s[1] * s[2] / 100 for s in scores)

        enterprise_name = data.get('企业名称', '贵企业')

        if total_score >= 90:
            level = "优秀"
            desc = "在现代企业制度建设方面表现突出"
        elif total_score >= 80:
            level = "良好"
            desc = "现代企业制度建设较为完善"
        elif total_score >= 70:
            level = "合格"
            desc = "基本建立现代企业制度框架"
        elif total_score >= 60:
            level = "待改进"
            desc = "现代企业制度建设有待加强"
        else:
            level = "不合格"
            desc = "现代企业制度建设存在明显不足"

        conclusion = (
            f"{enterprise_name}在中国特色现代企业制度建设评价中,综合得分为{total_score:.2f}分,"
            f"评价等级为\"{level}\"。{desc},符合现代企业治理的基本要求。"
        )

        return conclusion

    def _generate_suggestions(self, data, scores):
        """
        生成改进建议
        """
        suggestions = []

        if not scores:
            return ["建议完善企业信息填报,以便进行全面评价。"]

        # 找出得分较低的维度
        for dimension, score, weight in scores:
            if score < 80:
                if score < 60:
                    priority = "重点"
                else:
                    priority = "持续"

                suggestion = f"{priority}关注{dimension}方面的建设,建议加强相关制度完善和执行力度。"
                suggestions.append(suggestion)

        if not suggestions:
            suggestions.append("企业现代制度建设整体情况良好,建议继续保持并持续优化。")

        # 添加通用建议
        suggestions.append("建议定期开展企业制度评估,及时发现问题并改进。")
        suggestions.append("建议加强员工培训,提升全员对现代企业制度的认识和执行能力。")

        return suggestions

    def _clean_filename(self, filename):
        """
        清理文件名中的非法字符
        """
        invalid_chars = ['\\', '/', ':', '*', '?', '"', '<', '>', '|']
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        return filename
