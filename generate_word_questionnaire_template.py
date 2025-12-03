#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成Word格式的调查问卷模板
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

class WordQuestionnaireGenerator:
    """Word问卷生成器"""
    
    def __init__(self, excel_file='nankai_indicators.xlsx'):
        """初始化"""
        self.excel_file = excel_file
        self.output_dir = 'survey_generator/output/word_questionnaires'
        os.makedirs(self.output_dir, exist_ok=True)
    
    def set_chinese_font(self, run, font_name='宋体'):
        """设置中文字体"""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    def generate_questionnaire(self, level='初级'):
        """生成指定级别的问卷"""
        print(f"\n开始生成{level}问卷...")
        
        # 读取Excel数据
        df = pd.read_excel(self.excel_file, sheet_name=level)
        
        # 创建Word文档
        doc = Document()
        
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # 添加标题
        self._add_title(doc, level)
        
        # 添加说明
        self._add_instructions(doc)
        
        # 添加企业信息部分
        self._add_enterprise_info(doc)
        
        # 添加问卷题目
        self._add_questions(doc, df, level)
        
        # 添加页脚
        self._add_footer(doc)
        
        # 保存文档
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'调查问卷_{level}_{timestamp}.docx'
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        print(f"[OK] {level}问卷已生成: {filepath}")
        return filepath
    
    def _add_title(self, doc, level):
        """添加标题"""
        # 主标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('企业现代制度评价调查问卷')
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        self.set_chinese_font(run, '黑体')
        
        # 副标题
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f'（{level}版）')
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(46, 80, 144)
        self.set_chinese_font(run, '楷体')
        
        doc.add_paragraph()  # 空行
    
    def _add_instructions(self, doc):
        """添加填写说明"""
        # 说明标题
        p = doc.add_paragraph()
        run = p.add_run('【填写说明】')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(192, 0, 0)
        self.set_chinese_font(run, '黑体')
        
        # 说明内容
        instructions = [
            '1. 本问卷旨在评估企业现代制度建设情况，请如实填写。',
            '2. 请在相应选项前的方框内打"√"，多选题可选择多项。',
            '3. 对于"部分完成"的题目，请在详细说明栏中具体描述完成情况。',
            '4. 填写完成后，请将问卷发送至指定邮箱或在线提交。',
            '5. 如有疑问，请联系评价机构工作人员。'
        ]
        
        for instruction in instructions:
            p = doc.add_paragraph(instruction, style='List Number')
            p.paragraph_format.left_indent = Inches(0.25)
            for run in p.runs:
                run.font.size = Pt(10.5)
                self.set_chinese_font(run)
        
        doc.add_paragraph()  # 空行
    
    def _add_enterprise_info(self, doc):
        """添加企业信息部分"""
        # 标题
        p = doc.add_paragraph()
        run = p.add_run('一、企业基本信息')
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        self.set_chinese_font(run, '黑体')
        
        # 创建表格
        table = doc.add_table(rows=6, cols=4)
        table.style = 'Table Grid'
        
        # 设置表格内容
        info_fields = [
            ('企业名称', '', '统一社会信用代码', ''),
            ('企业类型', '□国有企业 □民营企业 □外资企业 □其他', '成立时间', ''),
            ('注册资本', '', '员工人数', ''),
            ('联系人', '', '职务', ''),
            ('联系电话', '', '电子邮箱', ''),
            ('企业地址', '', '', '')
        ]
        
        for i, (label1, value1, label2, value2) in enumerate(info_fields):
            row = table.rows[i]
            
            # 第一列（标签）
            cell = row.cells[0]
            cell.text = label1
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
            self.set_chinese_font(cell.paragraphs[0].runs[0])
            
            # 第二列（值）
            cell = row.cells[1]
            cell.text = value1
            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
            self.set_chinese_font(cell.paragraphs[0].runs[0])
            
            if label2:  # 如果有第二组
                # 第三列（标签）
                cell = row.cells[2]
                cell.text = label2
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                self.set_chinese_font(cell.paragraphs[0].runs[0])
                
                # 第四列（值）
                cell = row.cells[3]
                cell.text = value2
                cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                self.set_chinese_font(cell.paragraphs[0].runs[0])
            else:
                # 合并后三列
                row.cells[1].merge(row.cells[2]).merge(row.cells[3])
        
        doc.add_paragraph()  # 空行
    
    def _add_questions(self, doc, df, level):
        """添加问卷题目"""
        # 标题
        p = doc.add_paragraph()
        run = p.add_run('二、评价指标')
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        self.set_chinese_font(run, '黑体')
        
        current_level1 = None
        current_level2 = None
        question_num = 1
        
        for idx, row in df.iterrows():
            level1 = row.get('一级指标', '')
            level2 = row.get('二级指标', '')
            level3 = row.get('三级指标', '')
            question = row.get('题目', '')
            score = row.get('分值', '')
            evidence = row.get('佐证材料', '')
            
            # 跳过空行
            if pd.isna(question) or str(question).strip() == '':
                continue
            
            # 一级指标标题
            if level1 and level1 != current_level1:
                current_level1 = level1
                p = doc.add_paragraph()
                run = p.add_run(f'\n{level1}')
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(46, 80, 144)
                self.set_chinese_font(run, '黑体')
            
            # 二级指标标题
            if level2 and level2 != current_level2:
                current_level2 = level2
                p = doc.add_paragraph()
                run = p.add_run(f'（{level2}）')
                run.font.size = Pt(11)
                run.font.bold = True
                self.set_chinese_font(run, '黑体')
            
            # 题目
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            
            # 题号
            run = p.add_run(f'{question_num}. ')
            run.font.size = Pt(10.5)
            run.font.bold = True
            self.set_chinese_font(run)
            
            # 题目内容
            run = p.add_run(question)
            run.font.size = Pt(10.5)
            self.set_chinese_font(run)
            
            # 分值
            if pd.notna(score) and str(score).strip():
                run = p.add_run(f'  （{score}分）')
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)
                self.set_chinese_font(run)
            
            # 选项
            self._add_options(doc, row)
            
            # 佐证材料
            if pd.notna(evidence) and str(evidence).strip():
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run(f'佐证材料：{evidence}')
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
                self.set_chinese_font(run)
            
            question_num += 1
    
    def _add_options(self, doc, row):
        """添加选项"""
        options_str = row.get('选项', '')
        
        if pd.isna(options_str) or str(options_str).strip() == '':
            # 没有选项，添加填空行
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run('答：_' + '_' * 80)
            run.font.size = Pt(10.5)
            self.set_chinese_font(run)
            return
        
        # 解析选项
        options = str(options_str).split('|')
        
        for option in options:
            option = option.strip()
            if not option:
                continue
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            
            # 添加复选框
            run = p.add_run('□ ')
            run.font.size = Pt(10.5)
            self.set_chinese_font(run)
            
            # 添加选项文本
            run = p.add_run(option)
            run.font.size = Pt(10.5)
            self.set_chinese_font(run)
        
        # 如果有"部分完成"选项，添加详细说明栏
        if any('部分完成' in opt for opt in options):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.75)
            run = p.add_run('详细说明：_' + '_' * 70)
            run.font.size = Pt(9)
            run.font.italic = True
            self.set_chinese_font(run)
    
    def _add_footer(self, doc):
        """添加页脚"""
        doc.add_paragraph()
        doc.add_paragraph()
        
        # 签名栏
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run('填表人签字：_______________    日期：_______________')
        run.font.size = Pt(10.5)
        self.set_chinese_font(run)
        
        doc.add_paragraph()
        
        # 说明
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('— 感谢您的配合 —')
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        self.set_chinese_font(run, '楷体')
    
    def generate_all_levels(self):
        """生成所有级别的问卷"""
        levels = ['初级', '中级', '高级']
        generated_files = []
        
        print("\n" + "="*60)
        print("开始生成Word格式调查问卷")
        print("="*60)
        
        for level in levels:
            try:
                filepath = self.generate_questionnaire(level)
                generated_files.append(filepath)
            except Exception as e:
                print(f"[ERROR] 生成{level}问卷失败: {e}")
                import traceback
                traceback.print_exc()
        
        print("\n" + "="*60)
        print(f"问卷生成完成！共生成 {len(generated_files)} 个文件")
        print("="*60)
        print(f"\n输出目录: {self.output_dir}")
        
        for filepath in generated_files:
            print(f"  - {os.path.basename(filepath)}")
        
        return generated_files


def main():
    """主函数"""
    # 检查指标文件
    excel_file = 'nankai_indicators.xlsx'
    if not os.path.exists(excel_file):
        print(f"错误：找不到指标文件 {excel_file}")
        return
    
    # 创建生成器
    generator = WordQuestionnaireGenerator(excel_file)
    
    # 生成所有级别的问卷
    generator.generate_all_levels()


if __name__ == '__main__':
    main()