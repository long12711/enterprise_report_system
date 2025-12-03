#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""检查 Word 文件结构"""
import os
import sys
from docx import Document

docx_dir = r'D:\xwechat_files\wxid_nfuq3yq5zb4x22_dcf3\msg\file\2025-12'

try:
    files = os.listdir(docx_dir)
    docx_files = [f for f in files if '初级' in f and f.endswith('.docx')]
    
    if docx_files:
        file_path = os.path.join(docx_dir, docx_files[0])
        print(f"检查文件: {docx_files[0]}")
        print(f"文件路径: {file_path}")
        print(f"文件存在: {os.path.exists(file_path)}")
        
        doc = Document(file_path)
        print(f"表格数: {len(doc.tables)}")
        print(f"段落数: {len(doc.paragraphs)}")
        
        if len(doc.tables) > 0:
            table = doc.tables[0]
            print(f"第一个表格行数: {len(table.rows)}")
            print(f"第一个表格列数: {len(table.rows[0].cells)}")
            print("前3行内容:")
            for i, row in enumerate(table.rows[:3]):
                cells_text = [cell.text[:30] for cell in row.cells]
                print(f"  行{i}: {cells_text}")
        else:
            print("没有找到表格")
            print("前10个段落:")
            for i, p in enumerate(doc.paragraphs[:10]):
                text = p.text[:100] if p.text else "(空)"
                print(f"  {i}: {text}")
    else:
        print("未找到包含'初级'的 Word 文件")
        print(f"目录中的文件: {docx_files}")
except Exception as e:
    print(f"错误: {e}")
    import traceback
    traceback.print_exc()

