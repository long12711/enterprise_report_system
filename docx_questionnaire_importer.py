# -*- coding: utf-8 -*-
"""
从Word文档导入问卷题目
支持初级、中级、高级三个级别的问卷
"""
import os
import json
import uuid
from datetime import datetime
from docx import Document
from docx.table import Table
import logging
import os
import mysql.connector
from mysql.connector import Error

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocxQuestionnaireImporter:
    """从Word文档导入问卷题目"""

    def __init__(self, db_path='storage/questionnaires.json', db_config=None):
        """
        初始化导入器
        
        Args:
            db_path: 问卷JSON数据库文件路径（作为回退/备份）
            db_config: MySQL 连接配置（host/user/password/database）
        """
        self.db_path = db_path
        self.ensure_db_exists()
        self.db_config = db_config or {
            'host': os.environ.get('DB_HOST', 'localhost'),
            'user': os.environ.get('DB_USER', 'root'),
            'password': os.environ.get('DB_PASSWORD', '123456'),
            'database': os.environ.get('DB_NAME', 'localhost_3306')
        }

    def ensure_db_exists(self):
        """确保数据库文件存在"""
        os.makedirs(os.path.dirname(self.db_path), exist_ok=True)
        if not os.path.exists(self.db_path):
            with open(self.db_path, 'w', encoding='utf-8') as f:
                json.dump({
                    'surveys': [],
                    'questions': [],
                    'submissions': []
                }, f, ensure_ascii=False, indent=2)

    def load_db(self):
        """加载数据库"""
        try:
            with open(self.db_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            logger.error(f"加载数据库失败: {e}")
            return {'surveys': [], 'questions': [], 'submissions': []}

    def save_db(self, data):
        """保存数据库"""
        try:
            with open(self.db_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            logger.info(f"数据库已保存: {self.db_path}")
        except Exception as e:
            logger.error(f"保存数据库失败: {e}")

    def extract_questions_from_docx(self, docx_path):
        """
        从Word文档中提取问卷题目
        
        预期的Word文档结构：
        - 表格形式
        - 列：序号 | 一级指标 | 二级指标 | 题目 | 题目类型 | 分值 | 适用对象 | 补充数据/说明
        
        Args:
            docx_path: Word文档路径
            
        Returns:
            问题列表
        """
        if not os.path.exists(docx_path):
            logger.error(f"文件不存在: {docx_path}")
            return []

        try:
            doc = Document(docx_path)
            questions = []

            # 查找表格
            if not doc.tables:
                logger.warning(f"文档中没有找到表格: {docx_path}")
                return []

            # 遍历文档内所有表格（有的问卷一个文档包含多个表格）
            for t_index, table in enumerate(doc.tables, start=1):
                # 检查列数
                if len(table.rows) == 0:
                    continue
                col_count = len(table.rows[0].cells)
                logger.info(f"[表{t_index}] 列数: {col_count}")

                # 跳过表头行（第一行）
                for row_idx, row in enumerate(table.rows[1:], start=1):
                    try:
                        cells = row.cells

                        # 适配 7/8 列，两列含义固定：
                        # 0序号 1一级指标 2二级指标 3三级指标(题目) 4评分/类型 5分值/选项 6适用对象 7备注
                        if len(cells) < 7:
                            # 少于7列一般为说明或空行，跳过
                            continue

                        seq_no = self._extract_cell_text(cells[0]).strip()
                        level1 = self._extract_cell_text(cells[1]).strip()
                        level2 = self._extract_cell_text(cells[2]).strip()
                        question_text = self._extract_cell_text(cells[3]).strip()
                        question_type = self._extract_cell_text(cells[4]).strip()
                        score = self._extract_cell_text(cells[5]).strip()
                        applicable = self._extract_cell_text(cells[6]).strip()
                        remarks = self._extract_cell_text(cells[7]).strip() if len(cells) > 7 else ""

                        # 跳过空行
                        if not question_text:
                            continue

                        question = {
                            'id': str(uuid.uuid4()),
                            'seq_no': seq_no,
                            'level1': level1,
                            'level2': level2,
                            'question_text': question_text,
                            'question_type': question_type,
                            'score': score,
                            'applicable': applicable,
                            'remarks': remarks,
                            'requires_file': any(k in remarks for k in ['文件', '附件', '补充', '上传'])
                        }

                        questions.append(question)
                    except Exception as e:
                        logger.error(f"[表{t_index}] 处理第 {row_idx} 行时出错: {e}")
                        continue

            logger.info(f"共提取 {len(questions)} 个问题")
            return questions

        except Exception as e:
            logger.error(f"提取问卷题目失败: {e}")
            import traceback
            traceback.print_exc()
            return []

    def _extract_cell_text(self, cell):
        """提取单元格中的文本"""
        try:
            text = ''
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text.strip() + '\n'
            return text.strip()
        except Exception:
            return ''

    def import_questionnaire(self, docx_path, level, survey_name=None):
        """
        导入问卷
        
        Args:
            docx_path: Word文档路径
            level: 问卷级别 ('初级', '中级', '高级')
            survey_name: 问卷名称（可选）
            
        Returns:
            问卷ID
        """
        if level not in ['初级', '中级', '高级']:
            logger.error(f"无效的问卷级别: {level}")
            return None

        # 提取问题
        questions = self.extract_questions_from_docx(docx_path)
        if not questions:
            logger.error(f"无法从文档中提取问题: {docx_path}")
            return None

        # 创建问卷
        survey_id = str(uuid.uuid4())
        survey = {
            'id': survey_id,
            'name': survey_name or f'现代企业制度指数评价问卷_{level}',
            'level': level,
            'description': f'{level}级问卷，共 {len(questions)} 个问题',
            'total_questions': len(questions),
            'status': 'active',
            'created_at': datetime.now().isoformat(),
            'updated_at': datetime.now().isoformat(),
            'source_file': os.path.basename(docx_path)
        }

        # 为每个问题添加 survey_id
        for question in questions:
            question['survey_id'] = survey_id

        # 保存到 JSON 数据库（作为备份/回退）
        db = self.load_db()
        db['surveys'].append(survey)
        db['questions'].extend(questions)
        self.save_db(db)

        # 同步写入 MySQL（主存储）
        try:
            conn = mysql.connector.connect(**self.db_config)
            cur = conn.cursor(dictionary=True)
            # upsert template by level
            cur.execute("SELECT id FROM questionnaire_templates WHERE level=%s", (level,))
            row = cur.fetchone()
            if row:
                template_id = row['id']
                cur.execute(
                    "UPDATE questionnaire_templates SET name=%s, description=%s, total_questions=%s, source_file=%s, status='active', updated_at=NOW() WHERE id=%s",
                    (survey['name'], survey['description'], len(questions), survey['source_file'], template_id)
                )
            else:
                template_id = str(uuid.uuid4())
                cur.execute(
                    "INSERT INTO questionnaire_templates (id, level, name, description, total_questions, status, source_file) VALUES (%s,%s,%s,%s,%s,'active',%s)",
                    (template_id, level, survey['name'], survey['description'], len(questions), survey['source_file'])
                )
            # replace questions
            conn.commit()
            cur.execute("DELETE FROM questionnaire_template_questions WHERE template_id=%s", (template_id,))
            insert_sql = (
                "INSERT INTO questionnaire_template_questions "
                "(id, template_id, seq_no, level1, level2, question_text, question_type, score, applicable, remarks, requires_file, sort_order) "
                "VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
            )
            data = []
            for idx, q in enumerate(questions, start=1):
                # 解析分值为数字（如果可能）
                score_val = None
                try:
                    if isinstance(q.get('score'), (int, float)):
                        score_val = float(q.get('score'))
                    else:
                        score_val = float(str(q.get('score')).strip()) if str(q.get('score') or '').replace('.', '', 1).isdigit() else None
                except Exception:
                    score_val = None
                data.append((
                    str(uuid.uuid4()),
                    template_id,
                    q.get('seq_no'),
                    q.get('level1'),
                    q.get('level2'),
                    q.get('question_text'),
                    q.get('question_type'),
                    score_val,
                    q.get('applicable'),
                    q.get('remarks'),
                    1 if q.get('requires_file') else 0,
                    idx
                ))
            if data:
                cur.executemany(insert_sql, data)
            conn.commit()
            cur.close(); conn.close()
            logger.info(f"[MySQL] 模板已写入 level={level}, 模板题目={len(questions)}")
        except Exception as e:
            logger.error(f"[MySQL] 写入失败（已使用JSON备份）：{e}")

        logger.info(f"问卷导入成功: {survey_id} ({level}级, {len(questions)} 个问题)")
        return survey_id

    def import_all_questionnaires(self, docx_dir):
        """
        导入所有问卷（初级、中级、高级）
        
        Args:
            docx_dir: 包含Word文档的目录
            
        Returns:
            导入结果字典
        """
        result = {
            '初级': None,
            '中级': None,
            '高级': None
        }

        # 查找对应级别的文件
        for level in ['初级', '中级', '高级']:
            # 查找包含该级别的文件
            for filename in os.listdir(docx_dir):
                if level in filename and filename.endswith('.docx'):
                    docx_path = os.path.join(docx_dir, filename)
                    logger.info(f"导入 {level} 级问卷: {filename}")
                    
                    survey_id = self.import_questionnaire(docx_path, level)
                    result[level] = survey_id
                    break

        return result

    def get_survey(self, survey_id):
        """获取问卷信息"""
        db = self.load_db()
        for survey in db['surveys']:
            if survey['id'] == survey_id:
                return survey
        return None

    def get_survey_questions(self, survey_id):
        """获取问卷的所有问题"""
        db = self.load_db()
        return [q for q in db['questions'] if q['survey_id'] == survey_id]

    def get_survey_by_level(self, level):
        """按级别获取问卷"""
        db = self.load_db()
        for survey in db['surveys']:
            if survey['level'] == level:
                return survey
        return None

    def list_surveys(self):
        """列出所有问卷"""
        db = self.load_db()
        return db['surveys']

    def delete_survey(self, survey_id):
        """删除问卷及其所有问题"""
        db = self.load_db()
        
        # 删除问卷
        db['surveys'] = [s for s in db['surveys'] if s['id'] != survey_id]
        
        # 删除相关问题
        db['questions'] = [q for q in db['questions'] if q['survey_id'] != survey_id]
        
        self.save_db(db)
        logger.info(f"问卷已删除: {survey_id}")


# 使用示例
if __name__ == '__main__':
    # 初始化导入器
    importer = DocxQuestionnaireImporter()

    # 导入单个问卷
    docx_path = r'D:\xwechat_files\wxid_nfuq3yq5zb4x22_dcf3\msg\file\2025-12\南开大学现代企业制度指数评价问卷_初级_20251202_180815.docx'
    if os.path.exists(docx_path):
        survey_id = importer.import_questionnaire(docx_path, '初级')
        print(f"导入成功: {survey_id}")

        # 获取问卷信息
        survey = importer.get_survey(survey_id)
        print(f"问卷信息: {survey}")

        # 获取问题列表
        questions = importer.get_survey_questions(survey_id)
        print(f"问题数量: {len(questions)}")
        if questions:
            print(f"第一个问题: {questions[0]}")
    else:
        print(f"文件不存在: {docx_path}")

    # 列出所有问卷
    surveys = importer.list_surveys()
    print(f"\n所有问卷: {len(surveys)}")
    for survey in surveys:
        print(f"  - {survey['name']} ({survey['level']}级)")

