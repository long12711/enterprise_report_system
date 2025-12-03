#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""分析 Word 文件结构"""
import os
from docx import Document

docx_dir = r'D:\xwechat_files\wxid_nfuq3yq5zb4x22_dcf3\msg\file\2025-12'
files = os.listdir(docx_dir)
docx_files = [f for f in files if '初级' in f and f.endswith('.docx')]

if docx_files:
    file_path = os.path.join(docx_dir, docx_files[0])
    doc = Document(file_path)
    
    if len(doc.tables) > 0:
        table = doc.tables[0]
        print(f"表格信息:")
        print(f"  行数: {len(table.rows)}")
        print(f"  列数: {len(table.rows[0].cells)}")
        
        print(f"\n表头 (第一行):")
        header_row = table.rows[0]
        for i, cell in enumerate(header_row.cells):
            print(f"  列{i}: {cell.text}")
        
        print(f"\n前3个数据行:")
        for row_idx in range(1, min(4, len(table.rows))):
            row = table.rows[row_idx]
            print(f"\n  行{row_idx}:")
            for col_idx, cell in enumerate(row.cells):
                text = cell.text[:50] if cell.text else "(空)"
                print(f"    列{col_idx}: {text}")

