"""
企业自评报告生成器（Word格式）
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')  # 使用非交互式后端，适合多线程
import matplotlib.pyplot as plt
from datetime import datetime
import os
from score_calculator import ScoreCalculator

# 设置中文字体
matplotlib.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei']
matplotlib.rcParams['axes.unicode_minus'] = False


class EnterpriseReportGenerator:
    """企业自评报告生成器"""

    def __init__(self):
        """初始化报告生成器"""
        self.calculator = ScoreCalculator()

    def generate_report(self, questionnaire_file: str, output_path: str = None) -> str:
        """
        生成企业自评报告

        Args:
            questionnaire_file: 已填写的问卷文件路径
            output_path: 输出报告文件路径

        Returns:
            生成的报告文件路径
        """
        print(f"\n[INFO] 开始生成企业自评报告...")

        # 解析问卷
        questionnaire_data = self.calculator.parse_questionnaire(questionnaire_file)
        enterprise_info = questionnaire_data['enterprise_info']
        enterprise_name = enterprise_info.get('企业名称', '企业')

        # 计算得分
        score_summary = self.calculator.calculate_total_score(questionnaire_data)

        # 生成输出路径
        if output_path is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = f'{enterprise_name}_自评报告_{timestamp}.docx'

        # 创建Word文档
        doc = Document()

        # 设置文档样式
        self._setup_document_styles(doc)

        # 1. 封面
        self._create_cover_page(doc, enterprise_info)

        # 2. 目录（手动创建）
        doc.add_page_break()
        self._create_table_of_contents(doc)

        # 3. 报告摘要
        doc.add_page_break()
        self._create_executive_summary(doc, enterprise_info, score_summary)

        # 4. 企业基本信息
        doc.add_page_break()
        self._create_enterprise_info_section(doc, enterprise_info)

        # 5. 评价结果总览
        doc.add_page_break()
        self._create_score_overview(doc, score_summary, enterprise_name)

        # 6. 各维度详细分析
        doc.add_page_break()
        self._create_detailed_analysis(doc, score_summary)

        # 7. 问题清单
        doc.add_page_break()
        self._create_issue_list(doc, score_summary)

        # 8. 改进建议
        doc.add_page_break()
        self._create_recommendations(doc, score_summary)

        # 9. 附录：得分明细
        doc.add_page_break()
        self._create_score_details_appendix(doc, score_summary)

        # 保存文档
        doc.save(output_path)
        print(f"[OK] 报告生成成功: {output_path}")

        return output_path

    def _setup_document_styles(self, doc):
        """设置文档样式"""
        # 设置默认字体
        doc.styles['Normal'].font.name = '微软雅黑'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        doc.styles['Normal'].font.size = Pt(10.5)

    def _create_cover_page(self, doc, enterprise_info):
        """创建封面"""
        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = title.add_run('现代企业制度指数评价\n\n企业自评报告')
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        # 空行
        doc.add_paragraph('\n\n\n')

        # 企业信息
        info_para = doc.add_paragraph()
        info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        info_text = f"""
企业名称：{enterprise_info.get('企业名称', '')}

统一社会信用代码：{enterprise_info.get('统一社会信用代码', '')}

企业类型：{enterprise_info.get('企业类型', '')}

报告生成日期：{datetime.now().strftime('%Y年%m月%d日')}
        """
        run = info_para.add_run(info_text.strip())
        run.font.size = Pt(14)

        # 页脚声明
        doc.add_paragraph('\n\n\n\n\n')
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = footer_para.add_run('本报告仅供内部参考使用')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

    def _create_table_of_contents(self, doc):
        """创建目录"""
        heading = doc.add_heading('目  录', level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        toc_items = [
            '一、报告摘要',
            '二、企业基本信息',
            '三、评价结果总览',
            '四、各维度详细分析',
            '    4.1 党建引领',
            '    4.2 产权结构',
            '    4.3 公司治理结构和机制',
            '    4.4 战略管理',
            '    4.5 内控、风险与合规管理',
            '    4.6 科学民主管理',
            '    4.7 科技创新',
            '    4.8 社会责任与企业文化',
            '    4.9 家族企业治理',
            '五、问题清单与分析',
            '六、改进建议',
            '附录：得分明细表',
        ]

        for item in toc_items:
            para = doc.add_paragraph(item, style='List Number')
            para.paragraph_format.left_indent = Inches(0.5)

    def _create_executive_summary(self, doc, enterprise_info, score_summary):
        """创建报告摘要"""
        doc.add_heading('一、报告摘要', level=1)

        level, evaluation = self.calculator.get_evaluation_level(score_summary['score_percentage'])

        summary_text = f"""
本报告基于现代企业制度指数评价体系，对{enterprise_info.get('企业名称', '贵企业')}的制度建设情况进行了全面评估。

评价结果：
- 总得分：{score_summary['total_score']:.2f}分（满分{score_summary['max_possible_score']:.2f}分）
- 得分率：{score_summary['score_percentage']:.2f}%
- 评价等级：{level}级（{evaluation}）

完成情况：
- 适用问题总数：{score_summary['applicable_questions']}个
- 已完成问题：{score_summary['answered_count']}个
- 完成率：{score_summary['completion_rate']:.2f}%

主要发现：
"""
        doc.add_paragraph(summary_text.strip())

        # 找出得分最高和最低的维度
        level1_scores = score_summary['score_by_level1']
        if level1_scores:
            best_dim = max(level1_scores.items(), key=lambda x: x[1].get('percentage', 0))
            worst_dim = min(level1_scores.items(), key=lambda x: x[1].get('percentage', 0))

            findings = f"""
1. 制度建设最完善的维度是"{best_dim[0]}"，得分率为{best_dim[1].get('percentage', 0):.2f}%

2. 需要重点改进的维度是"{worst_dim[0]}"，得分率为{worst_dim[1].get('percentage', 0):.2f}%

3. 一票否决项扣分：{score_summary.get('negative_points', 0):.2f}分

4. 本次评价共识别出{len([d for d in score_summary['question_details'] if d['actual_score'] == 0 and d['base_score'] > 0])}个需要改进的问题
            """
            doc.add_paragraph(findings.strip())

    def _create_enterprise_info_section(self, doc, enterprise_info):
        """创建企业基本信息部分"""
        doc.add_heading('二、企业基本信息', level=1)

        # 创建表格
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'

        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '项目'
        header_cells[1].text = '内容'

        # 填充数据
        info_items = [
            ('企业名称', enterprise_info.get('企业名称', '')),
            ('统一社会信用代码', enterprise_info.get('统一社会信用代码', '')),
            ('企业类型', enterprise_info.get('企业类型', '')),
            ('成立时间', enterprise_info.get('成立时间', '')),
            ('注册资本', enterprise_info.get('注册资本（万元）', '') + '万元'),
            ('所属行业', enterprise_info.get('所属行业', '')),
            ('员工人数', enterprise_info.get('员工人数', '') + '人'),
            ('年营业收入', enterprise_info.get('年营业收入（万元）', '') + '万元'),
            ('企业地址', enterprise_info.get('企业地址', '')),
            ('联系人', enterprise_info.get('联系人姓名', '')),
            ('联系人职务', enterprise_info.get('联系人职务', '')),
            ('联系人电话', enterprise_info.get('联系人手机', '')),
            ('联系人邮箱', enterprise_info.get('联系人邮箱', '')),
        ]

        for item, value in info_items:
            row_cells = table.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = str(value)

    def _create_score_overview(self, doc, score_summary, enterprise_name):
        """创建评价结果总览"""
        doc.add_heading('三、评价结果总览', level=1)

        # 文字说明
        level, evaluation = self.calculator.get_evaluation_level(score_summary['score_percentage'])

        overview_text = f"""
根据现代企业制度指数评价体系，{enterprise_name}的总体评价结果如下：

总得分：{score_summary['total_score']:.2f}分 / {score_summary['max_possible_score']:.2f}分
得分率：{score_summary['score_percentage']:.2f}%
评价等级：{level}级（{evaluation}）
        """
        doc.add_paragraph(overview_text.strip())

        # 生成雷达图
        chart_path = self._generate_radar_chart(score_summary, enterprise_name)
        if chart_path and os.path.exists(chart_path):
            doc.add_paragraph('\n各维度得分雷达图：')
            doc.add_picture(chart_path, width=Inches(5.5))
            # 删除临时图片文件
            try:
                os.remove(chart_path)
            except:
                pass

        # 各维度得分表
        doc.add_paragraph('\n各维度得分详情：')
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'

        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '维度'
        header_cells[1].text = '得分'
        header_cells[2].text = '满分'
        header_cells[3].text = '得分率'
        header_cells[4].text = '评价'

        # 填充数据
        for level1, data in score_summary['score_by_level1'].items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(level1) if level1 else ''
            row_cells[1].text = f"{data['score']:.2f}"
            row_cells[2].text = f"{data['max_score']:.2f}"
            row_cells[3].text = f"{data['percentage']:.2f}%"

            # 评价
            if data['percentage'] >= 80:
                rating = '优秀'
            elif data['percentage'] >= 70:
                rating = '良好'
            elif data['percentage'] >= 60:
                rating = '中等'
            else:
                rating = '需改进'
            row_cells[4].text = rating

    def _generate_radar_chart(self, score_summary, enterprise_name):
        """生成雷达图"""
        try:
            level1_scores = score_summary['score_by_level1']
            if not level1_scores:
                return None

            categories = list(level1_scores.keys())
            values = [level1_scores[cat]['percentage'] for cat in categories]

            # 创建雷达图
            fig = plt.figure(figsize=(8, 8))
            ax = fig.add_subplot(111, projection='polar')

            # 计算角度
            angles = [n / float(len(categories)) * 2 * 3.14159 for n in range(len(categories))]
            values += values[:1]
            angles += angles[:1]

            # 绘制
            ax.plot(angles, values, 'o-', linewidth=2, label=enterprise_name)
            ax.fill(angles, values, alpha=0.25)
            ax.set_xticks(angles[:-1])
            ax.set_xticklabels(categories, size=10)
            ax.set_ylim(0, 100)
            ax.set_yticks([20, 40, 60, 80, 100])
            ax.set_yticklabels(['20%', '40%', '60%', '80%', '100%'])
            ax.grid(True)

            plt.title(f'{enterprise_name} - 各维度得分雷达图', size=14, pad=20)
            plt.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1))

            # 保存
            chart_path = f'temp_radar_{datetime.now().strftime("%Y%m%d%H%M%S")}.png'
            plt.savefig(chart_path, dpi=150, bbox_inches='tight')
            plt.close()

            return chart_path
        except Exception as e:
            print(f"[WARN] 生成雷达图失败: {e}")
            return None

    def _create_detailed_analysis(self, doc, score_summary):
        """创建各维度详细分析"""
        doc.add_heading('四、各维度详细分析', level=1)

        level1_list = list(score_summary['score_by_level1'].items())
        for idx, (level1, data) in enumerate(level1_list, 1):
            # 确保level1是字符串
            level1_str = str(level1) if level1 else '未分类'
            doc.add_heading(f'4.{idx} {level1_str}', level=2)

            analysis_text = f"""
得分：{data['score']:.2f}分 / {data['max_score']:.2f}分
得分率：{data['percentage']:.2f}%
问题数：{data['count']}个
            """
            doc.add_paragraph(analysis_text.strip())

            # 找出该维度下的二级指标
            level2_in_this_level1 = {
                k: v for k, v in score_summary['score_by_level2'].items()
                if str(v.get('level1', '')) == level1_str
            }

            if level2_in_this_level1:
                doc.add_paragraph('\n二级指标得分：')
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Light List Accent 1'

                header_cells = table.rows[0].cells
                header_cells[0].text = '二级指标'
                header_cells[1].text = '得分'
                header_cells[2].text = '满分'
                header_cells[3].text = '得分率'

                for level2, level2_data in level2_in_this_level1.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(level2) if level2 else ''
                    row_cells[1].text = f"{level2_data['score']:.2f}"
                    row_cells[2].text = f"{level2_data['max_score']:.2f}"
                    row_cells[3].text = f"{level2_data['percentage']:.2f}%"

            doc.add_paragraph()  # 空行

    def _create_issue_list(self, doc, score_summary):
        """创建问题清单"""
        doc.add_heading('五、问题清单与分析', level=1)

        # 筛选出得分为0的问题（需要改进的）
        issues = [
            q for q in score_summary['question_details']
            if q['actual_score'] == 0 and q['base_score'] > 0
        ]

        if not issues:
            doc.add_paragraph('恭喜！未发现明显问题。')
        else:
            doc.add_paragraph(f'共识别出{len(issues)}个需要改进的问题：\n')

            # 按一级指标分组
            issues_by_level1 = {}
            for issue in issues:
                level1 = issue['level1']
                if level1 not in issues_by_level1:
                    issues_by_level1[level1] = []
                issues_by_level1[level1].append(issue)

            for level1, level1_issues in sorted(issues_by_level1.items()):
                doc.add_heading(f'{level1}（{len(level1_issues)}个问题）', level=3)

                for idx, issue in enumerate(level1_issues, 1):
                    issue_text = f"{idx}. {issue['question']}\n   答案：{issue['answer']}\n   说明：{issue['comment']}"
                    if issue['remark']:
                        issue_text += f"\n   备注：{issue['remark']}"

                    doc.add_paragraph(issue_text, style='List Number')

        # 一票否决项
        veto_issues = [
            q for q in score_summary['question_details']
            if q['actual_score'] < 0
        ]

        if veto_issues:
            doc.add_paragraph('\n')
            doc.add_heading('一票否决项（需立即整改）', level=3)
            for issue in veto_issues:
                issue_text = f"- {issue['question']}\n  扣分：{issue['actual_score']:.2f}分"
                para = doc.add_paragraph(issue_text)
                para.runs[0].font.color.rgb = RGBColor(255, 0, 0)

    def _create_recommendations(self, doc, score_summary):
        """创建改进建议"""
        doc.add_heading('六、改进建议', level=1)

        # 根据得分情况给出建议
        level1_scores = score_summary['score_by_level1']

        # 找出需要改进的维度（得分率<70%）
        weak_dimensions = [
            (name, data) for name, data in level1_scores.items()
            if data['percentage'] < 70
        ]

        if weak_dimensions:
            doc.add_paragraph('根据评价结果，建议贵企业重点关注以下方面的制度建设：\n')

            for idx, (dim_name, dim_data) in enumerate(sorted(weak_dimensions, key=lambda x: x[1]['percentage']), 1):
                recommendation = f"""
{idx}. {dim_name}（当前得分率：{dim_data['percentage']:.2f}%）

   建议措施：
   - 组织专项培训，提升该领域的管理水平
   - 建立健全相关制度和流程
   - 加强日常监督和检查
   - 定期评估和持续改进
                """
                doc.add_paragraph(recommendation.strip())
        else:
            doc.add_paragraph('贵企业的现代制度建设总体表现良好，建议继续保持并持续优化。')

        # 通用建议
        doc.add_paragraph('\n\n通用建议：')
        general_recommendations = [
            '1. 建立制度定期review机制，确保制度与时俱进',
            '2. 加强员工培训，提升制度执行力',
            '3. 建立制度执行监督和问责机制',
            '4. 注重制度文化建设，营造良好的制度环境',
            '5. 借鉴行业最佳实践，持续优化制度体系'
        ]

        for rec in general_recommendations:
            doc.add_paragraph(rec, style='List Bullet')

    def _create_score_details_appendix(self, doc, score_summary):
        """创建得分明细附录"""
        doc.add_heading('附录：得分明细表', level=1)

        doc.add_paragraph('以下是所有问题的详细得分情况：\n')

        # 创建表格
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'

        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '序号'
        header_cells[1].text = '一级指标'
        header_cells[2].text = '问题'
        header_cells[3].text = '答案'
        header_cells[4].text = '基础分'
        header_cells[5].text = '实际得分'

        # 填充数据
        for question in score_summary['question_details']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(question['seq_no'])
            row_cells[1].text = str(question['level1']) if question['level1'] else ''
            question_text = str(question['question']) if question['question'] else ''
            row_cells[2].text = question_text[:50] + '...' if len(question_text) > 50 else question_text
            row_cells[3].text = str(question['answer'])
            row_cells[4].text = f"{question['base_score']:.2f}"
            row_cells[5].text = f"{question['actual_score']:.2f}"


if __name__ == '__main__':
    # 测试报告生成
    generator = EnterpriseReportGenerator()
    # 需要一个已填写的问卷文件
    # report_path = generator.generate_report('测试问卷_已填写.xlsx')
    # print(f"\n报告已生成: {report_path}")
    print("\n企业自评报告生成器已初始化")
