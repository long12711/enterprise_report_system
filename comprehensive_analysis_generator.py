"""
综合数据分析报告生成器
功能：
1. 汇总多个企业的评价数据
2. 生成行业/区��数据总结
3. 分析企业优劣势和特色
4. 输出专业的综合分析报告
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from datetime import datetime
import os
from collections import defaultdict
from score_calculator import ScoreCalculator

# 设置中文字体
matplotlib.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei']
matplotlib.rcParams['axes.unicode_minus'] = False


class ComprehensiveAnalysisGenerator:
    """综合数据分析报告生成器"""

    def __init__(self):
        """初始化"""
        self.calculator = ScoreCalculator()
        self.all_enterprises_data = []

    def generate_comprehensive_report(self, questionnaire_files: list, output_path: str = None) -> str:
        """
        生成综合分析报告

        Args:
            questionnaire_files: 多个企业的问卷文件列表
            output_path: 输出路径

        Returns:
            生成的报告路径
        """
        print(f"\n[INFO] 开始生成综合数据分析报告...")
        print(f"[INFO] 共{len(questionnaire_files)}家企业参与评价")

        # 解析所有企业数据
        for idx, file in enumerate(questionnaire_files, 1):
            try:
                print(f"  [{idx}/{len(questionnaire_files)}] 解析: {os.path.basename(file)}")
                questionnaire_data = self.calculator.parse_questionnaire(file)
                score_summary = self.calculator.calculate_total_score(questionnaire_data)

                self.all_enterprises_data.append({
                    'file': file,
                    'enterprise_info': questionnaire_data['enterprise_info'],
                    'score_summary': score_summary
                })
            except Exception as e:
                print(f"  [ERROR] 解析失败: {e}")

        if not self.all_enterprises_data:
            raise ValueError("没有成功解析的企业数据")

        # 生成输出路径
        if output_path is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = f'现代制度建设综合分析报告_{timestamp}.docx'

        # 创建Word文档
        doc = Document()

        # 设置样式
        self._setup_styles(doc)

        # 1. 封面
        self._create_cover(doc)

        # 2. 评价概况
        doc.add_page_break()
        self._create_overview(doc)

        # 3. 总体数据分析
        doc.add_page_break()
        self._create_overall_analysis(doc)

        # 4. 行业/规模对比分析
        doc.add_page_break()
        self._create_comparative_analysis(doc)

        # 5. 典型案例分析
        doc.add_page_break()
        self._create_case_studies(doc)

        # 6. 问题与建议
        doc.add_page_break()
        self._create_recommendations(doc)

        # 7. 附录：企业名单及得分
        doc.add_page_break()
        self._create_enterprise_list(doc)

        # 保存
        doc.save(output_path)
        print(f"[OK] 综合报告生成成功: {output_path}")

        return output_path

    def _setup_styles(self, doc):
        """设置文档样式"""
        doc.styles['Normal'].font.name = '仿宋_GB2312'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        doc.styles['Normal'].font.size = Pt(16)
        doc.styles['Normal'].paragraph_format.line_spacing = 1.5
        doc.styles['Normal'].paragraph_format.first_line_indent = Inches(0.5)

    def _create_cover(self, doc):
        """创建封面"""
        title = doc.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title.add_run('\n\n\n\n现代企业制度建设\n\n综合分析报告\n\n')
        title_run.font.name = '黑体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        title_run.font.size = Pt(26)
        title_run.font.bold = True

        date = doc.add_paragraph()
        date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        date_run = date.add_run(f'\n\n\n\n\n{datetime.now().strftime("%Y年%m月")}\n\n')
        date_run.font.name = '仿宋_GB2312'
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        date_run.font.size = Pt(18)

    def _create_overview(self, doc):
        """创建评价概况"""
        doc.add_heading('一、评价概况', level=1)

        total = len(self.all_enterprises_data)

        # 统计行业分布
        industry_dist = defaultdict(int)
        for data in self.all_enterprises_data:
            industry = data['enterprise_info'].get('所属行业', '其他')
            industry_dist[industry] += 1

        # 统计企业类型
        type_dist = defaultdict(int)
        for data in self.all_enterprises_data:
            etype = data['enterprise_info'].get('企业类型', '其他')
            type_dist[etype] += 1

        overview_text = f"""
本次评价共有{total}家企业参与，涵盖{len(industry_dist)}个行业��域。从企业类型看，{self._format_distribution(type_dist)}。从行业分布看，{self._format_industry_distribution(industry_dist)}。

评价工作严格按照现代企业制度指数评价体系开展，采取企业自评���专家评审相结合的方式，全面考察企业在党建引领、公司治理、战略管理、风险控制、科技创新等方面的制度建设情况。
        """

        para = doc.add_paragraph(overview_text.strip())
        para.paragraph_format.first_line_indent = Inches(0.5)

    def _format_distribution(self, dist_dict):
        """��式化分布描述"""
        items = sorted(dist_dict.items(), key=lambda x: x[1], reverse=True)
        if len(items) == 0:
            return "数据不足"
        elif len(items) == 1:
            return f"{items[0][0]}{items[0][1]}家"
        elif len(items) == 2:
            return f"{items[0][0]}{items[0][1]}家、{items[1][0]}{items[1][1]}家"
        else:
            main_types = f"{items[0][0]}{items[0][1]}家、{items[1][0]}{items[1][1]}家"
            others = sum(count for _, count in items[2:])
            return f"{main_types}、其他类型{others}家"

    def _format_industry_distribution(self, dist_dict):
        """格式化行业分布描述"""
        items = sorted(dist_dict.items(), key=lambda x: x[1], reverse=True)
        if len(items) <= 3:
            return "、".join([f"{name}领域{count}家" for name, count in items])
        else:
            main = "、".join([f"{name}领域{count}家" for name, count in items[:3]])
            return f"{main}等"

    def _create_overall_analysis(self, doc):
        """创建总体数据分析"""
        doc.add_heading('二、总体评价情况', level=1)

        # 计算统计数据
        scores = [data['score_summary']['score_percentage'] for data in self.all_enterprises_data]
        avg_score = np.mean(scores)
        max_score = np.max(scores)
        min_score = np.min(scores)

        # 等级分布
        level_dist = {'A': 0, 'B': 0, 'C': 0, 'D': 0}
        for data in self.all_enterprises_data:
            percentage = data['score_summary']['score_percentage']
            level, _ = self.calculator.get_evaluation_level(percentage)
            level_dist[level] += 1

        # 总体描述（叙述性）
        doc.add_heading('（一）总体水平', level=2)

        if avg_score >= 80:
            overall_assessment = "总体水平优良"
        elif avg_score >= 70:
            overall_assessment = "总体水平良好"
        elif avg_score >= 60:
            overall_assessment = "总体水平中等"
        else:
            overall_assessment = "仍有较大提升空间"

        overall_text = f"""
从总体情况看，参评企业现代制度建设{overall_assessment}。评价结果显示，企业普遍重视制度建设工作，基础制度框架基本建立，规��化管理水平逐步提升。

从评价等级分布看，达到优秀水平（A级）的企业{level_dist['A']}家，良好水平（B级）{level_dist['B']}家，中等水平（C级）{level_dist['C']}家，待提升水平（D级）{level_dist['D']}家。评价结果呈现出明显的层次性和差异性，反映了不同企业在制度建设方面的实际水平。
        """

        para = doc.add_paragraph(overall_text.strip())
        para.paragraph_format.first_line_indent = Inches(0.5)

        # 各维度平均水平
        doc.add_heading('（二）各维度建设情况', level=2)

        dimension_avg = self._calculate_dimension_average()

        # 找出优势和薄弱维度
        sorted_dims = sorted(dimension_avg.items(), key=lambda x: x[1], reverse=True)
        strong_dims = sorted_dims[:3]
        weak_dims = sorted_dims[-3:]

        dimension_text = f"""
从各维度平均水平看，参评企业在{strong_dims[0][0]}、{strong_dims[1][0]}、{strong_dims[2][0]}等方面制度建设相对完善，表现较为突出。这些领域的制度框架较为健全，执行情况良好，成为企业规范管理的重要支撑。

相对��言，{weak_dims[0][0]}、{weak_dims[1][0]}、{weak_dims[2][0]}等方面的制度建设存在一定短板，需要进一步加强和完善。这也是下一步推进企业现代制度建设的重点方向。
        """

        para = doc.add_paragraph(dimension_text.strip())
        para.paragraph_format.first_line_indent = Inches(0.5)

        # 生成雷达图
        self._create_dimension_radar_chart(dimension_avg, doc)

    def _calculate_dimension_average(self):
        """计算各维度平均得分率"""
        dimension_scores = defaultdict(list)

        for data in self.all_enterprises_data:
            for dim_name, dim_data in data['score_summary']['score_by_level1'].items():
                dimension_scores[dim_name].append(dim_data['percentage'])

        dimension_avg = {}
        for dim_name, scores in dimension_scores.items():
            dimension_avg[dim_name] = np.mean(scores)

        return dimension_avg

    def _create_dimension_radar_chart(self, dimension_avg, doc):
        """创建维度雷达图"""
        try:
            categories = list(dimension_avg.keys())
            values = list(dimension_avg.values())

            fig = plt.figure(figsize=(10, 10))
            ax = fig.add_subplot(111, projection='polar')

            angles = [n / float(len(categories)) * 2 * np.pi for n in range(len(categories))]
            values = values + [values[0]]
            angles = angles + [angles[0]]

            ax.plot(angles, values, 'o-', linewidth=2, label='平均水平', color='#2E5090')
            ax.fill(angles, values, alpha=0.25, color='#2E5090')
            ax.set_xticks(angles[:-1])
            ax.set_xticklabels(categories, size=12)
            ax.set_ylim(0, 100)
            ax.set_yticks([20, 40, 60, 80, 100])
            ax.set_yticklabels(['20%', '40%', '60%', '80%', '100%'])
            ax.grid(True)

            plt.title('各维度平均得分率', size=16, pad=20, fontweight='bold')

            chart_path = f'temp_comprehensive_radar_{datetime.now().strftime("%Y%m%d%H%M%S")}.png'
            plt.savefig(chart_path, dpi=150, bbox_inches='tight')
            plt.close()

            doc.add_paragraph('\n')
            doc.add_picture(chart_path, width=Inches(5.5))

            try:
                os.remove(chart_path)
            except:
                pass

        except Exception as e:
            print(f"[WARN] 生成雷达图失败: {e}")

    def _create_comparative_analysis(self, doc):
        """创建对比分析"""
        doc.add_heading('三、分类对比分析', level=1)

        # 按行业对比
        doc.add_heading('（一）行业对比', level=2)

        industry_analysis = self._analyze_by_industry()

        if len(industry_analysis) > 1:
            # 找出表现最好的行业
            best_industry = max(industry_analysis.items(), key=lambda x: x[1]['avg_score'])
            worst_industry = min(industry_analysis.items(), key=lambda x: x[1]['avg_score'])

            industry_text = f"""
从行业分布看，不同行业企业的制度建设水平呈现一定差异。{best_industry[0]}行业企业表现相对突出，平均水平较高，反映了该行业企业对现代制度建设的重视程度较高，管理规范化水平较好。

相比之下，{worst_industry[0]}行业企业在制度建设方面还有较大提升空间，建议加强行业交流学习，借鉴先进经验，提升整体管理水平。
            """

            para = doc.add_paragraph(industry_text.strip())
            para.paragraph_format.first_line_indent = Inches(0.5)

        # 按规模对比
        doc.add_heading('（二）规模对比', level=2)

        scale_analysis = self._analyze_by_scale()

        scale_text = """
从企业规模看，大型企业在制度建设的系统性、规范性方面通常表现更好，制度体系更���完善。中小企业虽然在制度建设方面可能不如大型企业系统全面，但普遍展现出较强的灵活性和适应性，能够根据自身实际情况建立适合的管理制度。
        """

        para = doc.add_paragraph(scale_text.strip())
        para.paragraph_format.first_line_indent = Inches(0.5)

    def _analyze_by_industry(self):
        """按行业分析"""
        industry_data = defaultdict(lambda: {'scores': [], 'count': 0})

        for data in self.all_enterprises_data:
            industry = data['enterprise_info'].get('所属行业', '其他')
            score = data['score_summary']['score_percentage']
            industry_data[industry]['scores'].append(score)
            industry_data[industry]['count'] += 1

        industry_analysis = {}
        for industry, data in industry_data.items():
            industry_analysis[industry] = {
                'avg_score': np.mean(data['scores']),
                'count': data['count']
            }

        return industry_analysis

    def _analyze_by_scale(self):
        """按规模分析"""
        scale_data = {'大型': [], '中型': [], '小型': []}

        for data in self.all_enterprises_data:
            revenue = float(data['enterprise_info'].get('年营业收入（万元）', 0))
            employees = int(data['enterprise_info'].get('员工人数', 0))
            score = data['score_summary']['score_percentage']

            if revenue > 50000 or employees > 500:
                scale_data['大型'].append(score)
            elif revenue > 10000 or employees > 200:
                scale_data['中型'].append(score)
            else:
                scale_data['小型'].append(score)

        return scale_data

    def _create_case_studies(self, doc):
        """创建典型案例分析"""
        doc.add_heading('四、典型案例', level=1)

        # 找出得分最高的前3名
        top_enterprises = sorted(
            self.all_enterprises_data,
            key=lambda x: x['score_summary']['score_percentage'],
            reverse=True
        )[:min(3, len(self.all_enterprises_data))]

        for idx, data in enumerate(top_enterprises, 1):
            enterprise_name = data['enterprise_info'].get('企业名称', '企业')
            score_summary = data['score_summary']

            doc.add_heading(f'（{self._num_to_chinese(idx)}）{enterprise_name}', level=2)

            # 找出该企业的优势维度
            level1_scores = score_summary['score_by_level1']
            strong_dims = sorted(level1_scores.items(), key=lambda x: x[1]['percentage'], reverse=True)[:3]

            case_text = f"""
该企业在现代制度建设方面表现突出，特别是在{strong_dims[0][0]}、{strong_dims[1][0]}等方面形成了较为完善的制度体系。企业高度重视规范化管理，将制度建设作为企业高质量发展的重要保障，值得其他企业学习借鉴。

主要特点：一是制度体系完整，覆盖企业运营的各个关键环节；二是制度执行有力，建立了有效的监督检查机制；三是持续优化改进，根据发展需要不断完善制度。
            """

            para = doc.add_paragraph(case_text.strip())
            para.paragraph_format.first_line_indent = Inches(0.5)

    def _create_recommendations(self, doc):
        """创建问题与建议"""
        doc.add_heading('五、工作建议', level=1)

        recommendations_text = """
基于本次评价情况，对进一步推进企业现代制度建设工作提出以下建议：

一是加强分类指导。针对不同行业、不同规模企业的特点，提供差异化的指导和支持，推动企业因地制宜建立完善现代企业制度。

二是强化典型引领。及时总结推广优秀企业的成功经验和创新做法，发挥示范带动作用，营造比学赶超的良好氛围。

三是完善服务体系。加强培训辅导，提升企业管理人员的制度建设能力和水平，为企业提供专业化、精准化的服务支持。

四是深化交流合作。搭建交流平台，组织企业间的学习交流活动，促进经验分享和互学互鉴。

五是强化动态管理。建立制度建设跟踪评估机制，及时了解企业制度建设进展，推动企业持续改进提升。
        """

        para = doc.add_paragraph(recommendations_text.strip())
        para.paragraph_format.first_line_indent = Inches(0.5)

    def _create_enterprise_list(self, doc):
        """创建企业名单附录"""
        doc.add_heading('附录：参评企业名单', level=1)

        # 按得分排序
        sorted_data = sorted(
            self.all_enterprises_data,
            key=lambda x: x['score_summary']['score_percentage'],
            reverse=True
        )

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        header_cells[0].text = '序号'
        header_cells[1].text = '企业名称'
        header_cells[2].text = '所属行业'
        header_cells[3].text = '评价等级'
        header_cells[4].text = '综合评价'

        for idx, data in enumerate(sorted_data, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = data['enterprise_info'].get('企业名称', '')
            row_cells[2].text = data['enterprise_info'].get('所属行业', '')

            percentage = data['score_summary']['score_percentage']
            level, evaluation = self.calculator.get_evaluation_level(percentage)

            row_cells[3].text = level + '级'
            row_cells[4].text = evaluation

    def _num_to_chinese(self, num):
        """数字转中文"""
        chinese_nums = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
        if num <= 10:
            return chinese_nums[num - 1]
        return str(num)


if __name__ == '__main__':
    # 测试
    generator = ComprehensiveAnalysisGenerator()
    print("\n综合数据分析报告生成器已初始化")
