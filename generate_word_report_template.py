#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成Word格式的企业评价报告模板
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

class WordReportTemplateGenerator:
    """Word报告模板生成器"""
    
    def __init__(self):
        """初始化"""
        self.output_dir = 'report_templates'
        os.makedirs(self.output_dir, exist_ok=True)
    
    def set_chinese_font(self, run, font_name='宋体'):
        """设置中文字体"""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    def add_page_break(self, doc):
        """添加分页符"""
        doc.add_page_break()
    
    def generate_report_template(self):
        """生成报告模板"""
        print("\n开始生成企业评价报告模板...")
        
        # 创建Word文档
        doc = Document()
        
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # 1. 封面
        self._add_cover(doc)
        self.add_page_break(doc)
        
        # 2. 目录（占位）
        self._add_toc(doc)
        self.add_page_break(doc)
        
        # 3. 执行摘要
        self._add_executive_summary(doc)
        self.add_page_break(doc)
        
        # 4. 第一章：企业概况
        self._add_chapter1(doc)
        self.add_page_break(doc)
        
        # 5. 第二章：评价结果
        self._add_chapter2(doc)
        self.add_page_break(doc)
        
        # 6. 第三章：维度分析
        self._add_chapter3(doc)
        self.add_page_break(doc)
        
        # 7. 第四章：亮点与优势
        self._add_chapter4(doc)
        self.add_page_break(doc)
        
        # 8. 第五章：问题与不足
        self._add_chapter5(doc)
        self.add_page_break(doc)
        
        # 9. 第六章：改进建议
        self._add_chapter6(doc)
        self.add_page_break(doc)
        
        # 10. 附录
        self._add_appendix(doc)
        
        # 保存文档
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'企业评价报告模板_{timestamp}.docx'
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        print(f"[OK] 报告模板已生成: {filepath}")
        return filepath
    
    def _add_cover(self, doc):
        """添加封面"""
        # 空行
        for _ in range(8):
            doc.add_paragraph()
        
        # 主标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('企业现代制度评价报告')
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        self.set_chinese_font(run, '黑体')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # 企业名称
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('【企业名称】')
        run.font.size = Pt(18)
        run.font.bold = True
        self.set_chinese_font(run, '楷体')
        
        # 空行
        for _ in range(10):
            doc.add_paragraph()
        
        # 底部信息
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('评价机构：南开大学企业制度研究中心')
        run.font.size = Pt(12)
        self.set_chinese_font(run)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'报告日期：{datetime.now().strftime("%Y年%m月%d日")}')
        run.font.size = Pt(12)
        self.set_chinese_font(run)
    
    def _add_toc(self, doc):
        """添加目录"""
        # 标题
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('目  录')
        run.font.size = Pt(18)
        run.font.bold = True
        self.set_chinese_font(run, '黑体')
        
        doc.add_paragraph()
        
        # 目录项
        toc_items = [
            ('执行摘要', '3'),
            ('第一章  企业概况', '4'),
            ('第二章  评价结果', '5'),
            ('第三章  维度分析', '6'),
            ('第四章  亮点与优势', '7'),
            ('第五章  问题与不足', '8'),
            ('第六章  改进建议', '9'),
            ('附录  详细评分表', '10')
        ]
        
        for item, page in toc_items:
            p = doc.add_paragraph()
            run = p.add_run(item)
            run.font.size = Pt(12)
            self.set_chinese_font(run)
            
            # 添加点线
            run = p.add_run(' ' + '.' * 50)
            run.font.size = Pt(12)
            
            # 添加页码
            run = p.add_run(f' {page}')
            run.font.size = Pt(12)
    
    def _add_executive_summary(self, doc):
        """添加执行摘要"""
        # 标题
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('执行摘要')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        self.set_chinese_font(run, '黑体')
        
        doc.add_paragraph()
        
        # 评价结论
        p = doc.add_paragraph()
        run = p.add_run('一、评价结论')
        run.font.size = Pt(14)
        run.font.bold = True
        self.set_chinese_font(run, '黑体')
        
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run('根据企业现代制度评价体系，【企业名称】的综合得分为【XX】分（满分100分），评价等级为【优秀/良好/合格/待改进】。')
        run.font.size = Pt(12)
        self.set_chinese_font(run)
        
        doc.add_paragraph()
        
        # 总体得分
        p = doc.add_paragraph()
        run = p.add_run('二、总体得分')
        run.font.size = Pt(14)
        run.font.bold = True
        self.set_chinese_font(run, '黑体')
        
        # 得分表格
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        
        # 表头
        headers = ['评价维度', '得分', '满分']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            self.set_chinese_font(cell.paragraphs[0].runs[0])
        
        # 数据行
        row = table.rows[1]
        row.cells[0].text = '综合得分'
        row.cells[1].text = '【XX】'
        row.cells[2].text = '100'
        
        doc.add_paragraph()
        
        # 主要亮点
        p = doc.add_paragraph()
        run = p.add_run('三、主要亮点')
        run.font.size = Pt(14)
        run.font.bold = True
        self.set_chinese_font(run, '黑体')
        
        highlights = [
            '【维度名称】表现优秀，得分率达【XX】%',
            '【具体指标】建设完善，符合现代企业制度要求',
            '【其他亮点】'
        ]
        
        for highlight in highlights:
            p = doc.add_paragraph(highlight, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
            for run in p.runs:
                run.font.size = Pt(12)
                self.set_chinese_font(run)
        
        doc.add_paragraph()
        
        # 主要问题
        p = doc.add_paragraph()
        run = p.add_run('四、主要问题')
        run.font.size = Pt(14)
        run.font.bold = True
        self.set_chinese_font(run, '黑体')
        
        issues = [
            '【维度名称】得分较低，需要加强建设',
            '【具体指标】存在不足，需要改进',
            '【其他问题】'
        ]
        
        for issue in issues:
            p = doc.add_paragraph(issue, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
            for run in p.runs:
                run.font.size = Pt(12)
                self.set_chinese_font(run)
    
    def _add_chapter1(self, doc):
        """第一章：企业概况"""
        # 章节标题
        p = doc.add_paragraph()
        run = p.add_run('第一章  企业概况')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        self.set_chinese_font(run, '黑体')
        
        doc.add_paragraph()
        
        # 基本信息
        p = doc.add_paragraph()
        run = p.add_run('一、基本信息')
        run.font.size = Pt(14)
        run.font.bold = True
        self.set_chinese_font(run, '黑体')
        
        # 信息表格
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        info_items = [
            ('企业名称', '【企业名称】'),
            ('统一社会信用代码', '【代码】'),
            ('企业类型', '【类型】'),
            ('成立时间', '【时间】'),
            ('注册资本', '【资本】'),
            ('员工人数', '【人数】')
        ]
        
        for i, (label, value) in enumerate(info_items):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
            self.set_chinese_font(row.cells[0].paragraphs[0].runs[0])
            
            row.cells[1].text = value
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
            self.set_chinese_font(row.cells[1].paragraphs[0].runs[0])
        
        doc.add_paragraph()
        
        # 企业简介
        p = doc.add_paragraph()
        run = p.add_run('二、企业简介')
        run.font.size = Pt(14)
        run.font.bold = True
        self.set_chinese_font(run, '黑体')
        
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run('【企业简介内容，包括主营业务、发展历程、市场地位等】')
        run.font.size = Pt(12)
        self.set_chinese_font(run)
    
    def _add_chapter2(self, doc):
        """第二章：评价结果"""
        # 章节标题
        p = doc.add_paragraph()
        run = p.add_run('第二章  评价结果')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        self.set_chinese_font(run, '黑体')
        
        doc.add_paragraph()
        
        # 综合得分
        p = doc.add_paragraph()
        run = p.add_run('一、综合得分')
        run.font.size = Pt(14)
        run.font.bold = True
        self.set_chinese_font(run, '黑体')
        
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run('【企业名称】在企业现代制度评价中的综合得分为【XX】分（满分100分），在同类企业中处于【领先/中等/落后】水平。')
        run.font.size = Pt(12)
        self.set_chinese_font(run)
        
        doc.add_paragraph()
        
        # 评价等级
        p = doc.add_paragraph()
        run = p.add_run('二、评价等级')
        run.font.size = Pt(14)
        run.font.bold = True
        self.set_chinese_font(run, '黑体')
        
        # 等级表格
        table = doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'
        
        # 表头
        headers = ['等级', '分数区间', '评价']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            self.set_chinese_font(cell.paragraphs[0].runs[0])
        
        # 数据行
        levels = [
            ('优秀', '90-100分', '制度建设完善，运行良好'),
            ('良好', '80-89分', '制度建设较好，运行正常'),
            ('合格', '60-79分', '制度基本健全，需要改进'),
            ('待改进', '60分以下', '制度建设不足，需要加强')
        ]
        
        for i, (level, score, desc) in enumerate(levels, 1):
            row = table.rows[i]
            row.cells[0].text = level
            row.cells[1].text = score
            row.cells[2].text = desc
            for j in range(3):
                row.cells[j].paragraphs[0].runs[0].font.size = Pt(10.5)
                self.set_chinese_font(row.cells[j].paragraphs[0].runs[0])
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run('根据评分结果，【企业名称】的评价等级为：【优秀/良好/合格/待改进】')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(192, 0, 0)
        self.set_chinese_font(run)
    
    def _add_chapter3(self, doc):
        """第三章：维度分析"""
        # 章节标题
        p = doc.add_paragraph()
        run = p.add_run('第三章  维度分析')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        self.set_chinese_font(run, '黑体')
        
        doc.add_paragraph()
        
        # 各维度得分
        p = doc.add_paragraph()
        run = p.add_run('一、各维度得分')
        run.font.size = Pt(14)
        run.font.bold = True
        self.set_chinese_font(run, '黑体')
        
        # 维度得分表格
        table = doc.add_table(rows=9, cols=4)
        table.style = 'Table Grid'
        
        # 表头
        headers = ['维度', '得分', '满分', '得分率']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            self.set_chinese_font(cell.paragraphs[0].runs[0])
        
        # 数据行
        dimensions = [
            '治理结构',
            '组织架构',
            '人力资源',
            '财务管理',
            '运营管理',
            '风险控制',
            '信息披露',
            '社会责任'
        ]
        
        for i, dim in enumerate(dimensions, 1):
            row = table.rows[i]
            row.cells[0].text = dim
            row.cells[1].text = '【XX】'
            row.cells[2].text = '【XX】'
            row.cells[3].text = '【XX】%'
            for j in range(4):
                row.cells[j].paragraphs[0].runs[0].font.size = Pt(10.5)
                self.set_chinese_font(row.cells[j].paragraphs[0].runs[0])
        
        doc.add_paragraph()
        
        # 维度分析
        p = doc.add_paragraph()
        run = p.add_run('二、维度分析')
        run.font.size = Pt(14)
        run.font.bold = True
        self.set_chinese_font(run, '黑体')
        
        for dim in dimensions[:3]:  # 示例前3个维度
            p = doc.add_paragraph()
            run = p.add_run(f'（一）{dim}')
            run.font.size = Pt(12)
            run.font.bold = True
            self.set_chinese_font(run)
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(f'【{dim}的详细分析内容，包括得分情况、主要表现、存在问题等】')
            run.font.size = Pt(12)
            self.set_chinese_font(run)
    
    def _add_chapter4(self, doc):
        """第四章：亮点与优势"""
        # 章节标题
        p = doc.add_paragraph()
        run = p.add_run('第四章  亮点与优势')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        self.set_chinese_font(run, '黑体')
        
        doc.add_paragraph()
        
        highlights = [
            ('治理结构完善', '【具体描述企业在治理结构方面的优势】'),
            ('制度建设规范', '【具体描述企业在制度建设方面的亮点】'),
            ('运营管理高效', '【具体描述企业在运营管理方面的特色】')
        ]
        
        for i, (title, content) in enumerate(highlights, 1):
            p = doc.add_paragraph()
            run = p.add_run(f'{i}. {title}')
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 128, 0)
            self.set_chinese_font(run)
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(content)
            run.font.size = Pt(12)
            self.set_chinese_font(run)
            
            doc.add_paragraph()
    
    def _add_chapter5(self, doc):
        """第五章：问题与不足"""
        # 章节标题
        p = doc.add_paragraph()
        run = p.add_run('第五章  问题与不足')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        self.set_chinese_font(run, '黑体')
        
        doc.add_paragraph()
        
        issues = [
            ('某维度建设不足', '【具体描述存在的问题】'),
            ('某制度执行不到位', '【具体描述执行中的问题】'),
            ('某方面需要改进', '【具体描述需要改进的地方】')
        ]
        
        for i, (title, content) in enumerate(issues, 1):
            p = doc.add_paragraph()
            run = p.add_run(f'{i}. {title}')
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(192, 0, 0)
            self.set_chinese_font(run)
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(content)
            run.font.size = Pt(12)
            self.set_chinese_font(run)
            
            doc.add_paragraph()
    
    def _add_chapter6(self, doc):
        """第六章：改进建议"""
        # 章节标题
        p = doc.add_paragraph()
        run = p.add_run('第六章  改进建议')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        self.set_chinese_font(run, '黑体')
        
        doc.add_paragraph()
        
        suggestions = [
            ('加强制度建设', '【具体的改进建议】'),
            ('完善管理流程', '【具体的改进措施】'),
            ('提升执行效果', '【具体的实施方案】')
        ]
        
        for i, (title, content) in enumerate(suggestions, 1):
            p = doc.add_paragraph()
            run = p.add_run(f'{i}. {title}')
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 255)
            self.set_chinese_font(run)
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(content)
            run.font.size = Pt(12)
            self.set_chinese_font(run)
            
            doc.add_paragraph()
    
    def _add_appendix(self, doc):
        """添加附录"""
        # 标题
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('附录  详细评分表')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        self.set_chinese_font(run, '黑体')
        
        doc.add_paragraph()
        
        # 评分表
        table = doc.add_table(rows=11, cols=5)
        table.style = 'Table Grid'
        
        # 表头
        headers = ['序号', '指标', '分值', '得分', '备注']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
            self.set_chinese_font(cell.paragraphs[0].runs[0])
        
        # 数据行（示例）
        for i in range(1, 11):
            row = table.rows[i]
            row.cells[0].text = str(i)
            row.cells[1].text = f'【指标{i}】'
            row.cells[2].text = '【分值】'
            row.cells[3].text = '【得分】'
            row.cells[4].text = '【备注】'
            for j in range(5):
                row.cells[j].paragraphs[0].runs[0].font.size = Pt(10)
                self.set_chinese_font(row.cells[j].paragraphs[0].runs[0])


def main():
    """主函数"""
    print("\n" + "="*60)
    print("开始生成企业评价报告模板")
    print("="*60)
    
    # 创建生成器
    generator = WordReportTemplateGenerator()
    
    # 生成报告模板
    filepath = generator.generate_report_template()
    
    print("\n" + "="*60)
    print("报告模板生成完成！")
    print("="*60)
    print(f"\n输出文件: {filepath}")
    print("\n使用说明：")
    print("1. 打开生成的Word文档")
    print("2. 将【】中的占位符替换为实际内容")
    print("3. 根据需要调整格式和内容")
    print("4. 保存为最终报告")


if __name__ == '__main__':
    main()