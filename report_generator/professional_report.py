"""
专业版企业报告生成器 - 独立模块
生成叙述性、专业性的企业制度建设评价报告
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from datetime import datetime
import os
import sys

# 设置中文字体
matplotlib.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei']
matplotlib.rcParams['axes.unicode_minus'] = False


class ProfessionalReportGenerator:
    """专业版企业报告生成器 - 叙述性风格"""

    def __init__(self, score_calculator=None):
        """
        初始化报告生成器

        Args:
            score_calculator: 评分计算器实例
                             如果为None，将尝试导入默认的计算器
        """
        self.calculator = score_calculator
        if self.calculator is None:
            self._init_default_calculator()

    def _init_default_calculator(self):
        """初始化默认的评分计算器"""
        try:
            # 尝试从项目根目录导入
            from score_calculator import ScoreCalculator
            self.calculator = ScoreCalculator()
        except ImportError:
            print("[WARN] 无法导入默认的ScoreCalculator，某些功能可能不可用")
            self.calculator = None

    def generate_report(self, questionnaire_file: str, output_path: str = None) -> str:
        """
        生成专业版企业报告

        Args:
            questionnaire_file: 已填写的问卷文件路径
            output_path: 输出报告文件路径
                        如果为None，将在reports目录生成

        Returns:
            生成的报告文件路径

        Raises:
            FileNotFoundError: 问卷文件不存在
            Exception: 生成失败
        """
        if not os.path.exists(questionnaire_file):
            raise FileNotFoundError(f"问卷文件不存在: {questionnaire_file}")

        print(f"\n[INFO] 开始生成专业版企业报告...")

        # 解析问卷
        questionnaire_data = self.calculator.parse_questionnaire(questionnaire_file)
        enterprise_info = questionnaire_data['enterprise_info']
        enterprise_name = enterprise_info.get('企业名称', '企业')

        # 计算得分
        score_summary = self.calculator.calculate_total_score(questionnaire_data)

        # 从问卷文件名生成报告文件名，确保关联性
        if output_path is None:
            os.makedirs('reports', exist_ok=True)
            base_name = os.path.basename(questionnaire_file)
            report_base_name = base_name.replace('问卷_', '专业报告_').replace('.xlsx', '')
            output_path = os.path.join('reports', f"{report_base_name}.docx")

        # 创建Word文档
        doc = Document()

        # 设置文档样式
        self._setup_professional_styles(doc)

        # 1. 封面
        self._create_professional_cover(doc, enterprise_info)

        # 2. 企业概况
        doc.add_page_break()
        self._create_enterprise_overview(doc, enterprise_info, score_summary)

        # 3. 制度建设总体评价
        doc.add_page_break()
        self._create_overall_assessment(doc, enterprise_info, score_summary)

        # 4. 制度建设主要成效（分维度叙述）
        doc.add_page_break()
        self._create_achievements_narrative(doc, score_summary)

        # 5. 存在的问题与挑战
        doc.add_page_break()
        self._create_challenges_analysis(doc, score_summary)

        # 6. 改进建议与发展方向
        doc.add_page_break()
        self._create_strategic_recommendations(doc, score_summary, enterprise_info)

        # 7. 附录：评价指标体系
        doc.add_page_break()
        self._create_appendix(doc, score_summary)

        # 保存文档
        doc.save(output_path)
        print(f"[OK] 专业报告生成成功: {output_path}")

        return output_path

    def _setup_professional_styles(self, doc):
        """设置专业文档样式"""
        # 正文样式
        doc.styles['Normal'].font.name = '仿宋_GB2312'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        doc.styles['Normal'].font.size = Pt(16)  # 四号字
        doc.styles['Normal'].paragraph_format.line_spacing = 1.5  # 1.5倍行距
        doc.styles['Normal'].paragraph_format.first_line_indent = Inches(0.5)  # 首行缩进

        # 标题样式
        heading1 = doc.styles['Heading 1']
        heading1.font.name = '黑体'
        heading1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        heading1.font.size = Pt(22)  # 二号字
        heading1.font.bold = True
        heading1.font.color.rgb = RGBColor(0, 0, 0)
        heading1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        heading2 = doc.styles['Heading 2']
        heading2.font.name = '黑体'
        heading2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        heading2.font.size = Pt(18)  # 三号字
        heading2.font.bold = True

        heading3 = doc.styles['Heading 3']
        heading3.font.name = '黑体'
        heading3._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        heading3.font.size = Pt(16)  # 四号字
        heading3.font.bold = True

    def _create_professional_cover(self, doc, enterprise_info):
        """创建专业封面"""
        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title.add_run('\n\n\n\n现代企业制度建设评价报告\n\n')
        title_run.font.name = '黑体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        title_run.font.size = Pt(26)  # 一号字
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)

        # 企业名称
        company = doc.add_paragraph()
        company.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        company_run = company.add_run(f'\n\n{enterprise_info.get("企业名称", "")}\n\n\n')
        company_run.font.name = '黑体'
        company_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        company_run.font.size = Pt(22)  # 二号字
        company_run.font.bold = True

        # 日期
        date = doc.add_paragraph()
        date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        date_run = date.add_run(f'\n\n\n\n\n{datetime.now().strftime("%Y年%m月")}\n\n')
        date_run.font.name = '仿宋_GB2312'
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        date_run.font.size = Pt(18)

    def _create_enterprise_overview(self, doc, enterprise_info, score_summary):
        """创建企业概况（叙述性）"""
        doc.add_heading('一、企业概况', level=1)

        # （一）基本情况
        doc.add_heading('（一）基本情况', level=2)

        overview_text = f"""
{enterprise_info.get('企业名称', '')}成立于{enterprise_info.get('成立时间', '')}，是一家专注于{enterprise_info.get('所属行业', '')}领域的{enterprise_info.get('企业类型', '')}。企业注册资本{enterprise_info.get('注册资本（万元）', '')}万元，年营业收入达{enterprise_info.get('年营业收入（万元）', '')}万元，现有员工{enterprise_info.get('员工人数', '')}人。

企业统一社会信用代码为{enterprise_info.get('统一社会信用代码', '')}，坐落于{enterprise_info.get('企业地址', '')}。
        """

        para = doc.add_paragraph(overview_text.strip())
        para.paragraph_format.first_line_indent = Inches(0.5)

        # （二）发展定位
        doc.add_heading('（二）发展定位', level=2)

        # 根据行业自动生成定位描述
        industry = enterprise_info.get('所属行业', '')
        positioning = self._generate_positioning_description(enterprise_info, score_summary)

        para = doc.add_paragraph(positioning)
        para.paragraph_format.first_line_indent = Inches(0.5)

    def _generate_positioning_description(self, enterprise_info, score_summary):
        """生成企业定位描述"""
        industry = enterprise_info.get('所属行业', '')
        revenue = float(enterprise_info.get('年营业收入（万元）', 0))
        employees = int(enterprise_info.get('员工人数', 0))

        # 企业规模判断
        if revenue > 50000 or employees > 500:
            scale = "行业领军企业"
        elif revenue > 10000 or employees > 200:
            scale = "区域骨干企业"
        else:
            scale = "成长型企业"

        positioning = f"""
企业立足{industry}行业，定位为{scale}。近年来，企业始终坚持创新驱动发展战略，不断优化产业结构，持续提升核心竞争力。通过完善现代企业制度建设，企业在规范化管理、科学决策、风险防控等方面取得显著成效，为企业高质量发展奠定了坚实基础。
        """
        return positioning.strip()

    def _create_overall_assessment(self, doc, enterprise_info, score_summary):
        """创建总体评价（叙述性，弱化数据）"""
        doc.add_heading('二、现代制度建设总体评价', level=1)

        # 获取评价等级
        level, evaluation = self.calculator.get_evaluation_level(score_summary['score_percentage'])

        # 等级描述映射
        level_descriptions = {
            'A': '优秀',
            'B': '良好',
            'C': '中等',
            'D': '有待提升'
        }

        level_analysis = {
            'A': '表明企业现代制度建设已达到较高水平，制度体系健全完善，运行规范有效',
            'B': '说明企业现代制度建设总体良好，基础制度较为健全，部分领域仍有提升空间',
            'C': '反映企业现代制度建设处于中等水平，基本框架已建立，但系统性、规范性有待加强',
            'D': '显示企业现代制度建设相对薄弱，亟需系统规划和全面提升'
        }

        # 生成叙述性评价
        assessment_text = f"""
经对{enterprise_info.get('企业名称', '贵企业')}现代制度建设情况进行全面评估，企业制度建设水平总体达到{level_descriptions.get(level, '中等')}标准，{level_analysis.get(level, '')}。

本次评价立足于现代企业制度指数评价体系，从党建引领、产权结构、公司治理、战略管理、风险管控、科技创新、社会责任等多个维度进行综合考量。评价结果显示，企业在适用指标中完成度较高，制度建设取得积极成效。
        """

        para = doc.add_paragraph(assessment_text.strip())
        para.paragraph_format.first_line_indent = Inches(0.5)

        # 添加总体评价表格
        doc.add_heading('（一）总体评价结果', level=2)

        overall_table = doc.add_table(rows=5, cols=2)
        overall_table.style = 'Light Grid Accent 1'

        # 设置表格内容
        overall_table.rows[0].cells[0].text = '评价项目'
        overall_table.rows[0].cells[1].text = '评价结果'
        overall_table.rows[1].cells[0].text = '评价等级'
        overall_table.rows[1].cells[1].text = level_descriptions.get(level, '中等')
        overall_table.rows[2].cells[0].text = '总体水平'
        overall_table.rows[2].cells[1].text = f"{score_summary['score_percentage']:.1f}%"
        overall_table.rows[3].cells[0].text = '适用维度数'
        overall_table.rows[3].cells[1].text = f"{len(score_summary['score_by_level1'])}个"
        overall_table.rows[4].cells[0].text = '评价结论'
        overall_table.rows[4].cells[1].text = level_analysis.get(level, '')

        # 美化表格
        self._format_table(overall_table, header_row=True)

        # 表格后的解释
        doc.add_paragraph()  # 空行
        table_explanation = f"""
上表展示了企业现代制度建设的总体评价结果。从评价等级来看，企业达到了{level_descriptions.get(level, '中等')}水平，总体建设水平为{score_summary['score_percentage']:.1f}%，这一成绩在同类企业中{self._get_ranking_phrase(score_summary['score_percentage'])}。本次评价涵盖{len(score_summary['score_by_level1'])}个适用维度，全面考察了企业在各个领域的制度建设情况。{level_analysis.get(level, '')}，这为企业未来的发展奠定了{self._get_foundation_phrase(score_summary['score_percentage'])}。
        """
        para = doc.add_paragraph(table_explanation.strip())
        para.paragraph_format.first_line_indent = Inches(0.5)

        # 添加维度得分详细表格
        doc.add_heading('（二）各维度建设水平详表', level=2)

        # 找出表现突出和需改进的维度
        level1_scores = score_summary['score_by_level1']
        sorted_dimensions = sorted(level1_scores.items(), key=lambda x: x[1]['percentage'], reverse=True)

        # 创建维度得分表
        dimension_table = doc.add_table(rows=len(sorted_dimensions) + 1, cols=4)
        dimension_table.style = 'Light Grid Accent 1'

        # 表头
        header_cells = dimension_table.rows[0].cells
        header_cells[0].text = '序号'
        header_cells[1].text = '评价维度'
        header_cells[2].text = '建设水平'
        header_cells[3].text = '水平评价'

        # 填充数据
        for idx, (dimension_name, dimension_data) in enumerate(sorted_dimensions, 1):
            row_cells = dimension_table.rows[idx].cells
            row_cells[0].text = str(idx)
            row_cells[1].text = dimension_name

            percentage = dimension_data['percentage']
            row_cells[2].text = f"{percentage:.1f}%"

            # 水平评价
            if percentage >= 90:
                row_cells[3].text = "优秀"
            elif percentage >= 80:
                row_cells[3].text = "良好"
            elif percentage >= 70:
                row_cells[3].text = "中等"
            elif percentage >= 60:
                row_cells[3].text = "合格"
            else:
                row_cells[3].text = "待提升"

        # 美化表格
        self._format_table(dimension_table, header_row=True)

        # 表格后的详细解释
        doc.add_paragraph()  # 空行

        top_dimensions = sorted_dimensions[:3]  # 前3名
        bottom_dimensions = sorted_dimensions[-2:] if len(sorted_dimensions) > 2 else []  # 后2名

        dimension_explanation = f"""
上表详细列出了企业在各个维度的建设水平，按得分率从高到低排序。从表中数据可以看出，企业在不同维度的表现存在明显差异，呈现出"有优势、有短板"的特点。
        """
        para = doc.add_paragraph(dimension_explanation.strip())
        para.paragraph_format.first_line_indent = Inches(0.5)

        # 优势维度分析
        if top_dimensions:
            strengths = "、".join([f"{dim[0]}（{dim[1]['percentage']:.1f}%）" for dim in top_dimensions])
            strength_text = f"""
从优势维度来看，企业在{strengths}等方面表现突出，相关制度体系健全，执行情况良好，成为企业规范管理的重要保障。特别是{top_dimensions[0][0]}维度，以{top_dimensions[0][1]['percentage']:.1f}%的建设水平位居首位，充分说明企业在该领域建立了较为系统的制度框架，有效支撑了企业的规范化运作。这些优势领域的成功经验，值得在其他维度中借鉴推广。
            """
            para = doc.add_paragraph(strength_text.strip())
            para.paragraph_format.first_line_indent = Inches(0.5)

        # 提升空间维度分析
        if bottom_dimensions:
            improvement = "、".join([f"{dim[0]}（{dim[1]['percentage']:.1f}%）" for dim in bottom_dimensions])
            improvement_text = f"""
同时，从表格末端可以看出，企业在{improvement}等领域的建设水平相对较低，存在明显的提升空间。这些维度的制度完善程度有待加强，建议企业高度重视，结合实际情况，有针对性地补齐制度短板，进一步提升管理规范化水平。需要特别关注的是{bottom_dimensions[-1][0]}维度，其建设水平仅为{bottom_dimensions[-1][1]['percentage']:.1f}%，应当作为未来制度建设的重点突破方向。
            """
            para = doc.add_paragraph(improvement_text.strip())
            para.paragraph_format.first_line_indent = Inches(0.5)

    def _create_achievements_narrative(self, doc, score_summary):
        """创建主要成效（叙述性，按维度）"""
        doc.add_heading('三、制度建设主要成效', level=1)

        level1_scores = score_summary['score_by_level1']

        # 只展示得分率>60%的维度（优势维度）
        strong_dimensions = [(name, data) for name, data in level1_scores.items() if data['percentage'] >= 60]
        strong_dimensions.sort(key=lambda x: x[1]['percentage'], reverse=True)

        for idx, (dimension_name, dimension_data) in enumerate(strong_dimensions, 1):
            # 维度标题
            doc.add_heading(f'（{self._num_to_chinese(idx)}）{dimension_name}', level=2)

            # 创建该维度的详细表格
            dimension_detail_table = doc.add_table(rows=4, cols=2)
            dimension_detail_table.style = 'Light Grid Accent 1'

            dimension_detail_table.rows[0].cells[0].text = '评价指标'
            dimension_detail_table.rows[0].cells[1].text = '评价结果'
            dimension_detail_table.rows[1].cells[0].text = '维度名称'
            dimension_detail_table.rows[1].cells[1].text = dimension_name
            dimension_detail_table.rows[2].cells[0].text = '建设水平'
            dimension_detail_table.rows[2].cells[1].text = f"{dimension_data['percentage']:.1f}%"

            # 水平评价
            percentage = dimension_data['percentage']
            if percentage >= 90:
                level_text = "优秀（制度健全完善）"
            elif percentage >= 80:
                level_text = "良好（制度较为健全）"
            elif percentage >= 70:
                level_text = "中等（制度基本健全）"
            else:
                level_text = "合格（制度初步建立）"

            dimension_detail_table.rows[3].cells[0].text = '水平评价'
            dimension_detail_table.rows[3].cells[1].text = level_text

            # 美化表格
            self._format_table(dimension_detail_table, header_row=True)

            doc.add_paragraph()  # 空行

            # 表格后的数据解读
            table_analysis = f"""
上表展示了{dimension_name}维度的具体评价结果。该维度建设水平达到{dimension_data['percentage']:.1f}%，在所有评价维度中排名第{idx}位，属于企业的{self._get_advantage_type(percentage)}。这一成绩的取得，充分体现了企业在{dimension_name}领域的重视程度和投入力度。
            """
            para = doc.add_paragraph(table_analysis.strip())
            para.paragraph_format.first_line_indent = Inches(0.5)

            # 生成该维度的叙述性评价
            narrative = self._generate_dimension_narrative(
                dimension_name,
                dimension_data,
                score_summary
            )

            para = doc.add_paragraph(narrative)
            para.paragraph_format.first_line_indent = Inches(0.5)

            # 添加该维度的亮点总结
            highlights = self._generate_dimension_highlights(dimension_name, percentage)
            if highlights:
                doc.add_paragraph()  # 空行
                highlight_text = f"【主要亮点】{highlights}"
                para = doc.add_paragraph(highlight_text)
                para.paragraph_format.first_line_indent = Inches(0.5)
                # 设置亮点文字样式
                for run in para.runs:
                    if '【主要亮点】' in run.text:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0, 102, 204)  # 蓝色

    def _get_advantage_type(self, percentage):
        """获取优势类型描述"""
        if percentage >= 90:
            return "核心优势领域"
        elif percentage >= 80:
            return "优势领域"
        elif percentage >= 70:
            return "相对优势领域"
        else:
            return "基础领域"

    def _generate_dimension_highlights(self, dimension_name, percentage):
        """生成维度亮点总结"""
        if percentage < 70:
            return None

        highlights = {
            '党建引领': "党组织作用充分发挥，党建工作与企业发展深度融合，为企业提供了坚强的政治保证和组织保证。",
            '产权结构': "产权关系清晰明确，股权结构合理规范，为企业长远发展奠定了坚实的产权基础。",
            '公司治理结构和机制': "治理结构健全规范，'三会一层'运作有效，决策程序科学民主，实现了有效的权力制衡。",
            '战略管理': "战略规划科学系统，战略执行机制完善，战略引领作用明显，为企业发展指明了方向。",
            '内控、风险与合规管理': "内控体系完善，风险防控到位，合规意识强，有效保障了企业稳健运营。",
            '科学民主管理': "民主管理制度健全，职工参与渠道畅通，充分调动了员工积极性和创造性。",
            '科技创新': "创新体系完善，研发投入充足，创新成果丰硕，持续增强企业核心竞争力。",
            '社会责任与企业文化': "社会责任意识强，企业文化特色鲜明，树立了良好的企业形象和社会口碑。",
            '家族企业治理': "家族治理机制规范，传承制度完善，实现了家族控制与现代管理的有机结合。"
        }

        return highlights.get(dimension_name, f"制度建设成效显著，为企业规范化管理提供了有力支撑。")

    def _generate_dimension_narrative(self, dimension_name, dimension_data, score_summary):
        """生成维度叙述性评价"""
        percentage = dimension_data['percentage']

        # 维度成效描述模板
        narratives = {
            '党建引领': f"""
企业深入贯彻党建引领发展理念，{self._get_level_phrase(percentage)}。企业党组织建设规范有序，组织架构健全，充分发挥了政治核心和领导作用。党建工作与企业经营发展深度融合，为企业改革发展提供了坚强政治保证和组织保证。
            """,
            '产权结构': f"""
企业产权结构{self._get_clarity_phrase(percentage)}，股权关系明晰，出资情况规范。通过建立科学合理的产权制度，企业有效保障了各方股东的合法权益，为企业长期稳定发展奠定了产权基础。
            """,
            '公司治理结构和机制': f"""
企业{self._get_establishment_phrase(percentage)}了较为完善的公司治理结构，股东大会、董事会、监事会等机构设置规范，职责分工明确。"三会一层"运作{self._get_operation_phrase(percentage)}，决策程序科学规范，有效形成了权责明确、运转协调、有效制衡的治理机制。
            """,
            '战略管理': f"""
企业高度重视战略管理工作，{self._get_formulation_phrase(percentage)}了符合企业实际的发展战略。战略规划体系较为健全，战略执行机制有效，战略评估与调整机制逐步完善。通过系统的战略管理，企业明确了发展方向和路径。
            """,
            '内控、风险与合规管理': f"""
企业{{self._get_attention_phrase(percentage)}}内部控制、风险管理和合规经营。建立了相对完善的内控制度体系，风险识别、评估和应对机制逐步健全，合规管理意识不断增强。通过加强内控风险管理，企业有效防范了各类经营风险。
            """,
            '科学民主管理': f"""
企业积极推行科学民主管理，{{self._get_implementation_phrase(percentage)}}了多种形式的民主管理制度。职工参与企业管理的渠道较为畅通，民主决策、民主管理、民主监督机制逐步完善，有效调动了职工的积极性和创造性。
            """,
            '科技创新': f"""
企业坚持创新驱动发展战略，{{self._get_innovation_phrase(percentage)}}。研发投入持续增加，创新体系不断完善，创新成果转化机制逐步健全。通过持续创新，企业核心竞争力不断提升。
            """,
            '社会责任与企业文化': f"""
企业积极履行社会责任，{{self._get_culture_phrase(percentage)}}。企业文化建设有声有色，价值理念深入人心。在环境保护、职工权益保障、社会公益等方面表现良好，树立了良好的企业形象。
            """,
            '家族企业治理': f"""
企业作为家族企业，{{self._get_governance_phrase(percentage)}}了较为规范的治理机制。家族成员与职业经理人的权责关系逐步理顺，家族传承制度初步建立，为企业可持续发展提供了制度保障。
            """
        }

        return narratives.get(dimension_name, f"企业在{dimension_name}方面制度建设取得积极成效。").strip()

    def _get_level_phrase(self, percentage):
        """根据得分率返回水平描述"""
        if percentage >= 90: return "制度建设成效显著"
        elif percentage >= 80: return "制度体系较为完善"
        elif percentage >= 70: return "制度建设总体良好"
        elif percentage >= 60: return "基本建立了相关制度"
        else: return "制度建设有待加强"

    def _get_clarity_phrase(self, percentage):
        """产权清晰度描述"""
        if percentage >= 90: return "清晰明确"
        elif percentage >= 80: return "较为清晰"
        elif percentage >= 70: return "基本清晰"
        else: return "有待进一步明晰"

    def _get_establishment_phrase(self, percentage):
        """建立程度描述"""
        if percentage >= 90: return "建立健全"
        elif percentage >= 80: return "建立"
        elif percentage >= 70: return "初步建立"
        else: return "正在建立"

    def _get_operation_phrase(self, percentage):
        """运作效果描述"""
        if percentage >= 90: return "规范高效"
        elif percentage >= 80: return "规范有效"
        elif percentage >= 70: return "总体规范"
        else: return "有待规范"

    def _get_formulation_phrase(self, percentage):
        """制定水平描述"""
        if percentage >= 90: return "科学制定并有效实施"
        elif percentage >= 80: return "制定"
        elif percentage >= 70: return "初步制定"
        else: return "正在制定"

    def _get_attention_phrase(self, percentage):
        """重视程度描述"""
        if percentage >= 90: return "高度重视并有效推进"
        elif percentage >= 80: return "较为重视"
        elif percentage >= 70: return "重视"
        else: return "逐步重视"

    def _get_implementation_phrase(self, percentage):
        """实施程度描述"""
        if percentage >= 90: return "全面实施"
        elif percentage >= 80: return "较好实施"
        elif percentage >= 70: return "实施"
        else: return "正在推进"

    def _get_innovation_phrase(self, percentage):
        """创新水平描述"""
        if percentage >= 90: return "创新能力突出，成效显著"
        elif percentage >= 80: return "创新体系较为完善"
        elif percentage >= 70: return "创新意识较强"
        else: return "创新能力有待提升"

    def _get_culture_phrase(self, percentage):
        """文化建设描述"""
        if percentage >= 90: return "企业文化特色鲜明"
        elif percentage >= 80: return "企业文化建设成效明显"
        elif percentage >= 70: return "注重企业文化建设"
        else: return "企业文化建设逐步推进"

    def _get_governance_phrase(self, percentage):
        """治理水平描述"""
        if percentage >= 90: return "建立健全"
        elif percentage >= 80: return "建立"
        elif percentage >= 70: return "初步建立"
        else: return "正在探索建立"

    def _create_challenges_analysis(self, doc, score_summary):
        """创建问题与挑战分析（委婉表述）"""
        doc.add_heading('四、存在的问题与挑战', level=1)

        # 找出得分率<60%的维度
        level1_scores = score_summary['score_by_level1']
        weak_dimensions = [(name, data) for name, data in level1_scores.items() if data['percentage'] < 60]
        weak_dimensions.sort(key=lambda x: x[1]['percentage'])

        if not weak_dimensions:
            para = doc.add_paragraph("""
评价结果显示，企业现代制度建设总体水平良好，各项制度基本健全。建议企业继续巩固现有成果，持续优化完善，不断提升制度执行力和有效性。
            """.strip())
            para.paragraph_format.first_line_indent = Inches(0.5)
        else:
            intro = doc.add_paragraph("""
通过评价发现，企业现代制度建设虽已取得积极成效，但对标高质量发展要求，部分领域仍有提升空间，主要体现在以下几个方面：
            """.strip())
            intro.paragraph_format.first_line_indent = Inches(0.5)

            doc.add_paragraph()  # 空行

            # 添加问题分析表格
            problem_table = doc.add_table(rows=len(weak_dimensions) + 1, cols=4)
            problem_table.style = 'Light Grid Accent 1'

            # 表头
            header_cells = problem_table.rows[0].cells
            header_cells[0].text = '序号'
            header_cells[1].text = '薄弱维度'
            header_cells[2].text = '当前水平'
            header_cells[3].text = '改进优先级'

            # 填充数据
            for idx, (dimension_name, dimension_data) in enumerate(weak_dimensions, 1):
                row_cells = problem_table.rows[idx].cells
                row_cells[0].text = str(idx)
                row_cells[1].text = dimension_name
                row_cells[2].text = f"{dimension_data['percentage']:.1f}%"

                # 改进优先级
                percentage = dimension_data['percentage']
                if percentage < 40:
                    row_cells[3].text = "高优先级"
                elif percentage < 50:
                    row_cells[3].text = "较高优先级"
                else:
                    row_cells[3].text = "中等优先级"

            # 美化表格
            self._format_table(problem_table, header_row=True)

            doc.add_paragraph()  # 空行

            # 表格后的详细分析
            problem_analysis = f"""
上表列出了企业制度建设中相对薄弱的{len(weak_dimensions)}个维度及其改进优先级。这些领域的制度建设水平尚未达到60%的合格线，说明在制度的系统性、完整性和执行力方面还存在明显短板。从改进优先级来看，建议企业按照"先急后缓、突出重点"的原则，优先加强高优先级领域的制度建设，逐步补齐短板，提升整体制度建设水平。
            """
            para = doc.add_paragraph(problem_analysis.strip())
            para.paragraph_format.first_line_indent = Inches(0.5)

    def _create_strategic_recommendations(self, doc, score_summary, enterprise_info):
        """创建改进建议与发展方向（战略性、指导性）"""
        doc.add_heading('五、改进建议与发展方向', level=1)

        # 总体建议
        doc.add_heading('（一）总体思路', level=2)

        overall_text = f"""
建议{enterprise_info.get('企业名称', '企业')}以本次评价为契机，系统谋划现代制度建设工作，坚持问题导向和目标导向相结合，补短板、强弱项、促提升，不断完善中国特色现代企业制度，为企业高质量发展提供坚实制度保障。
        """
        para = doc.add_paragraph(overall_text.strip())
        para.paragraph_format.first_line_indent = Inches(0.5)

        # 具体建议（根据薄弱环节）
        doc.add_heading('（二）重点举措', level=2)

        level1_scores = score_summary['score_by_level1']
        weak_dimensions = [(name, data) for name, data in level1_scores.items() if data['percentage'] < 70]

        if weak_dimensions:
            for idx, (dimension_name, dimension_data) in enumerate(weak_dimensions[:5], 1):  # 最多5条建议
                recommendation = self._generate_dimension_recommendation(dimension_name, dimension_data)
                para = doc.add_paragraph(f"{idx}. {recommendation}")
                para.paragraph_format.first_line_indent = Inches(0.5)
        else:
            general_recommendations = [
                "持续深化改革创新，不断完善公司治理结构，提升治理能力和水平。",
                "加强制度执行监督，确保各项制度落地见效，切实发挥制度效能。",
                "强化数字化转型，推进制度管理信息化建设，提升管理效率。",
                "注重人才队伍建设，提升管理人员专业能力和制度执行力。",
                "加强交流学习，借鉴先进经验，持续优化完善制度体系。"
            ]
            for idx, rec in enumerate(general_recommendations, 1):
                para = doc.add_paragraph(f"{idx}. {rec}")
                para.paragraph_format.first_line_indent = Inches(0.5)

    def _generate_dimension_recommendation(self, dimension_name, dimension_data):
        """生成维度改进建议"""
        recommendations = {
            '党建引领': "健全党建工作制度，加强党组织建设，充分发挥党组织的领导作用和政治核心作用，推动党建工作与生产经营深度融合。",
            '产权结构': "进一步明晰产权关系，规范股权管理，完善产权保护制度，为企业长远发展夯实产权基础。",
            '公司治理结构和机制': "完善公司治理结构，规范'三会一层'运作，建立科学决策机制，提升治理规范化水平。",
            '战略管理': "加强战略规划工作，建立健全战略管理体系，完善战略实施和评估机制，增强战略引领作用。",
            '内控、风险与合规管理': "强化内部控制建设，完善风险管理体系，加强合规管理，有效防范化解各类风险。",
            '科学民主管理': "深化民主管理，畅通职工参与渠道，完善民主决策机制，充分调动职工积极性。",
            '科技创新': "加大研发投入，完善创新体系，健全激励机制，增强企业创新能力和核心竞争力。",
            '社会责任与企业文化': "加强企业文化建设，积极履行社会责任，树立良好企业形象，促进企业可持续发展。",
            '家族企业治理': "完善家族企业治理机制，理顺家族成员与职业经理人关系，建立科学传承制度。"
        }
        return recommendations.get(dimension_name, f"进一步完善{dimension_name}相关制度，提升管理规范化水平。")

    def _create_appendix(self, doc, score_summary):
        """创建附录"""
        doc.add_heading('附录：评价指标概览', level=1)

        para = doc.add_paragraph("""
本次评价依据现代企业制度指数评价体系，涵盖党建引领、产权结构、公司治理、战略管理、内控风险、科技创新、社会责任等多个维度，全面考察企业制度建设情况。各维度得分情况如下：
        """.strip())
        para.paragraph_format.first_line_indent = Inches(0.5)

        # 创建简要的得分表
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        header_cells[0].text = '评价维度'
        header_cells[1].text = '建设水平'
        header_cells[2].text = '评价'

        for level1, data in score_summary['score_by_level1'].items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(level1) if level1 else ''

            # 用文字描述代替百分比
            percentage = data['percentage']
            if percentage >= 90:
                level_desc = "优秀"
                assessment = "制度健全完善"
            elif percentage >= 80:
                level_desc = "良好"
                assessment = "制度较为健全"
            elif percentage >= 70:
                level_desc = "中等"
                assessment = "制度基本健全"
            elif percentage >= 60:
                level_desc = "合格"
                assessment = "制度初步建立"
            else:
                level_desc = "待提升"
                assessment = "需进一步完善"

            row_cells[1].text = level_desc
            row_cells[2].text = assessment

    def _get_ranking_phrase(self, percentage):
        """根据得分率返回排名描述"""
        if percentage >= 90:
            return "处于领先地位"
        elif percentage >= 80:
            return "位居前列"
        elif percentage >= 70:
            return "处于中上游水平"
        elif percentage >= 60:
            return "处于中等水平"
        else:
            return "仍有较大提升空间"

    def _get_foundation_phrase(self, percentage):
        """根据得分率返回基础描述"""
        if percentage >= 90:
            return "坚实的制度基础"
        elif percentage >= 80:
            return "良好的制度基础"
        elif percentage >= 70:
            return "一定的制度基础"
        else:
            return "初步的制度基础"

    def _format_table(self, table, header_row=False):
        """美化表格格式"""
        for i, row in enumerate(table.rows):
            # 设置行高
            row.height = Inches(0.35)

            for j, cell in enumerate(row.cells):
                # 设置字体
                for paragraph in cell.paragraphs:
                    # 清除默认格式
                    paragraph.paragraph_format.space_before = Pt(3)
                    paragraph.paragraph_format.space_after = Pt(3)

                    for run in paragraph.runs:
                        run.font.name = '仿宋_GB2312'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(12)

                    # 设置段落对齐
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # 表头加粗并加大字号
                    if header_row and i == 0:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(13)
                            run.font.color.rgb = RGBColor(0, 0, 0)

    def _num_to_chinese(self, num):
        """数字转中文"""
        chinese_nums = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
        if num <= 10:
            return chinese_nums[num - 1]
        return str(num)


if __name__ == '__main__':
    # 测试
    generator = ProfessionalReportGenerator()
    print("\n专业版企业报告生成器已初始化")

