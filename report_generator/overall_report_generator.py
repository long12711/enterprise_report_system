"""
整体分析报告生成器（汇总多个企业的评价结果）
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
import matplotlib
import pandas as pd
import numpy as np
from datetime import datetime
import os
from typing import List, Dict
from score_calculator import ScoreCalculator

# 设置中文字体
matplotlib.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei']
matplotlib.rcParams['axes.unicode_minus'] = False


class OverallAnalysisReportGenerator:
    """整体分析报告生成器"""

    def __init__(self):
        """初始化报告生成器"""
        self.calculator = ScoreCalculator()

    def generate_report(self, questionnaire_files: List[str], output_path: str = None) -> str:
        """
        生成整体分析报告

        Args:
            questionnaire_files: 已填写的问卷文件列表
            output_path: 输出报告文件路径

        Returns:
            生成的报告文件路径
        """
        print(f"\n[INFO] 开始生成整体分析报告，共{len(questionnaire_files)}家企业...")

        # 收集所有企业的数据
        all_enterprise_data = []
        for qf in questionnaire_files:
            try:
                questionnaire_data = self.calculator.parse_questionnaire(qf)
                score_summary = self.calculator.calculate_total_score(questionnaire_data)
                all_enterprise_data.append({
                    'file': qf,
                    'info': questionnaire_data['enterprise_info'],
                    'score': score_summary
                })
                print(f"[OK] 已处理: {questionnaire_data['enterprise_info'].get('企业名称', qf)}")
            except Exception as e:
                print(f"[WARN] 处理失败 {qf}: {e}")

        if not all_enterprise_data:
            raise ValueError("没有有效的问卷数据")

        # 生成输出路径
        if output_path is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = f'整体分析报告_{timestamp}.docx'

        # 创建Word文档
        doc = Document()
        self._setup_document_styles(doc)

        # 1. 封面
        self._create_cover_page(doc, len(all_enterprise_data))

        # 2. 报告摘要
        doc.add_page_break()
        self._create_executive_summary(doc, all_enterprise_data)

        # 3. 参评企业概况
        doc.add_page_break()
        self._create_enterprise_overview(doc, all_enterprise_data)

        # 4. 总体得分分析
        doc.add_page_break()
        self._create_overall_score_analysis(doc, all_enterprise_data)

        # 5. 各维度对比分析
        doc.add_page_break()
        self._create_dimension_comparison(doc, all_enterprise_data)

        # 6. 企业排名
        doc.add_page_break()
        self._create_enterprise_ranking(doc, all_enterprise_data)

        # 7. 行业对比分析
        doc.add_page_break()
        self._create_industry_comparison(doc, all_enterprise_data)

        # 8. 共性问题分析
        doc.add_page_break()
        self._create_common_issues_analysis(doc, all_enterprise_data)

        # 9. 最佳实践
        doc.add_page_break()
        self._create_best_practices(doc, all_enterprise_data)

        # 10. 整体建议
        doc.add_page_break()
        self._create_overall_recommendations(doc, all_enterprise_data)

        # 11. 附录：企业得分明细
        doc.add_page_break()
        self._create_score_details_appendix(doc, all_enterprise_data)

        # 保存文档
        doc.save(output_path)
        print(f"[OK] 整体分析报告生成成功: {output_path}")

        return output_path

    def _setup_document_styles(self, doc):
        """设置文档样式"""
        doc.styles['Normal'].font.name = '微软雅黑'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        doc.styles['Normal'].font.size = Pt(10.5)

    def _create_cover_page(self, doc, enterprise_count):
        """创建封面"""
        title = doc.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = title.add_run('现代企业制度指数评价\n\n整体分析报告')
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph('\n\n\n')

        info_para = doc.add_paragraph()
        info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        info_text = f"""
参评企业数量：{enterprise_count}家

报告生成日期：{datetime.now().strftime('%Y年%m月%d日')}
        """
        run = info_para.add_run(info_text.strip())
        run.font.size = Pt(14)

        doc.add_paragraph('\n\n\n\n\n')
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = footer_para.add_run('本报告仅供内部参考使用')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

    def _create_executive_summary(self, doc, all_data):
        """创建报告摘要"""
        doc.add_heading('一、报告摘要', level=1)

        # 计算统计数据
        scores = [d['score']['total_score'] for d in all_data]
        percentages = [d['score']['score_percentage'] for d in all_data]

        summary_text = f"""
本报告基于现代企业制度指数评价体系，对{len(all_data)}家企业的制度建设情况进行了整体分析评估。

总体情况：
- 参评企业数：{len(all_data)}家
- 平均得分：{np.mean(scores):.2f}分
- 平均得分率：{np.mean(percentages):.2f}%
- 最高得分率：{np.max(percentages):.2f}%
- 最低得分率：{np.min(percentages):.2f}%

得分分布：
- A级（90%以上）：{len([p for p in percentages if p >= 90])}家
- B级（80-90%）：{len([p for p in percentages if 80 <= p < 90])}家
- C级（70-80%）：{len([p for p in percentages if 70 <= p < 80])}家
- D级（60-70%）：{len([p for p in percentages if 60 <= p < 70])}家
- E级（60%以下）：{len([p for p in percentages if p < 60])}家
        """
        doc.add_paragraph(summary_text.strip())

    def _create_enterprise_overview(self, doc, all_data):
        """创建参评企业概况"""
        doc.add_heading('二、参评企业概况', level=1)

        doc.add_paragraph(f'本次评价共有{len(all_data)}家企业参与，基本情况如下：\n')

        # 创建表格
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'

        # 表头
        header_cells = table.rows[0].cells
        headers = ['序号', '企业名称', '企业类型', '所属行业', '员工人数', '成立时间']
        for i, header in enumerate(headers):
            header_cells[i].text = header

        # 填充数据
        for idx, data in enumerate(all_data, 1):
            info = data['info']
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = info.get('企业名称', '')
            row_cells[2].text = info.get('企业类型', '')
            row_cells[3].text = info.get('所属行业', '')
            row_cells[4].text = info.get('员工人数', '')
            row_cells[5].text = info.get('成立时间', '')

        # 统计企业类型分布
        doc.add_paragraph('\n')
        doc.add_heading('企业类型分布', level=3)
        enterprise_types = {}
        for data in all_data:
            etype = data['info'].get('企业类型', '未知')
            enterprise_types[etype] = enterprise_types.get(etype, 0) + 1

        for etype, count in enterprise_types.items():
            doc.add_paragraph(f'- {etype}：{count}家', style='List Bullet')

    def _create_overall_score_analysis(self, doc, all_data):
        """创建总体得分分析"""
        doc.add_heading('三、总体得分分析', level=1)

        percentages = [d['score']['score_percentage'] for d in all_data]

        # 描述性统计
        stats_text = f"""
总体得分统计：
- 平均得分率：{np.mean(percentages):.2f}%
- 中位数得分率：{np.median(percentages):.2f}%
- 标准差：{np.std(percentages):.2f}%
- 最高得分率：{np.max(percentages):.2f}%
- 最低得分率：{np.min(percentages):.2f}%
        """
        doc.add_paragraph(stats_text.strip())

        # 生成得分分布柱状图
        chart_path = self._generate_score_distribution_chart(all_data)
        if chart_path and os.path.exists(chart_path):
            doc.add_paragraph('\n得分分布图：')
            doc.add_picture(chart_path, width=Inches(6))
            try:
                os.remove(chart_path)
            except:
                pass

    def _generate_score_distribution_chart(self, all_data):
        """生成得分分布柱状图"""
        try:
            percentages = [d['score']['score_percentage'] for d in all_data]

            fig, ax = plt.subplots(figsize=(10, 6))

            # 定义分数区间
            bins = [0, 60, 70, 80, 90, 100]
            labels = ['E级\n(<60%)', 'D级\n(60-70%)', 'C级\n(70-80%)', 'B级\n(80-90%)', 'A级\n(>=90%)']

            # 统计各区间人数
            hist, _ = np.histogram(percentages, bins=bins)

            # 绘制柱状图
            x = range(len(labels))
            bars = ax.bar(x, hist, color=['#d62728', '#ff7f0e', '#ffbb00', '#2ca02c', '#1f77b4'])

            ax.set_xlabel('评价等级', fontsize=12)
            ax.set_ylabel('企业数量', fontsize=12)
            ax.set_title('企业得分分布', fontsize=14, fontweight='bold')
            ax.set_xticks(x)
            ax.set_xticklabels(labels)

            # 在柱状图上显示数值
            for i, bar in enumerate(bars):
                height = bar.get_height()
                if height > 0:
                    ax.text(bar.get_x() + bar.get_width()/2., height,
                           f'{int(height)}家',
                           ha='center', va='bottom', fontsize=10)

            plt.tight_layout()

            chart_path = f'temp_distribution_{datetime.now().strftime("%Y%m%d%H%M%S")}.png'
            plt.savefig(chart_path, dpi=150, bbox_inches='tight')
            plt.close()

            return chart_path
        except Exception as e:
            print(f"[WARN] 生成分布图失败: {e}")
            return None

    def _create_dimension_comparison(self, doc, all_data):
        """创建各维度对比分析"""
        doc.add_heading('四、各维度对比分析', level=1)

        # 收集所有一级指标的数据
        all_level1_data = {}
        for data in all_data:
            for level1, level1_data in data['score']['score_by_level1'].items():
                if level1 not in all_level1_data:
                    all_level1_data[level1] = []
                all_level1_data[level1].append(level1_data['percentage'])

        # 计算各维度的平均得分
        doc.add_paragraph('各维度平均得分情况：\n')

        # 创建表格
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        header_cells[0].text = '维度'
        header_cells[1].text = '平均得分率'
        header_cells[2].text = '最高得分率'
        header_cells[3].text = '最低得分率'
        header_cells[4].text = '标准差'

        for level1, percentages in sorted(all_level1_data.items()):
            row_cells = table.add_row().cells
            row_cells[0].text = level1
            row_cells[1].text = f'{np.mean(percentages):.2f}%'
            row_cells[2].text = f'{np.max(percentages):.2f}%'
            row_cells[3].text = f'{np.min(percentages):.2f}%'
            row_cells[4].text = f'{np.std(percentages):.2f}%'

        # 生成对比柱状图
        chart_path = self._generate_dimension_comparison_chart(all_level1_data)
        if chart_path and os.path.exists(chart_path):
            doc.add_paragraph('\n各维度平均得分对比图：')
            doc.add_picture(chart_path, width=Inches(6.5))
            try:
                os.remove(chart_path)
            except:
                pass

    def _generate_dimension_comparison_chart(self, all_level1_data):
        """生成维度对比柱状图"""
        try:
            dimensions = list(all_level1_data.keys())
            avg_scores = [np.mean(all_level1_data[d]) for d in dimensions]

            fig, ax = plt.subplots(figsize=(12, 6))

            x = range(len(dimensions))
            bars = ax.bar(x, avg_scores, color='steelblue')

            ax.set_xlabel('维度', fontsize=12)
            ax.set_ylabel('平均得分率 (%)', fontsize=12)
            ax.set_title('各维度平均得分对比', fontsize=14, fontweight='bold')
            ax.set_xticks(x)
            ax.set_xticklabels(dimensions, rotation=15, ha='right')
            ax.set_ylim(0, 100)

            # 添加平均线
            overall_avg = np.mean(avg_scores)
            ax.axhline(y=overall_avg, color='r', linestyle='--', label=f'总体平均: {overall_avg:.2f}%')

            # 在柱状图上显示数值
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height,
                       f'{height:.1f}%',
                       ha='center', va='bottom', fontsize=9)

            ax.legend()
            plt.tight_layout()

            chart_path = f'temp_dimension_{datetime.now().strftime("%Y%m%d%H%M%S")}.png'
            plt.savefig(chart_path, dpi=150, bbox_inches='tight')
            plt.close()

            return chart_path
        except Exception as e:
            print(f"[WARN] 生成维度对比图失败: {e}")
            return None

    def _create_enterprise_ranking(self, doc, all_data):
        """创建企业排名"""
        doc.add_heading('五、企业排名', level=1)

        # 按得分率排序
        sorted_data = sorted(all_data, key=lambda x: x['score']['score_percentage'], reverse=True)

        doc.add_paragraph('按总得分率排名：\n')

        # 创建排名表
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        header_cells[0].text = '排名'
        header_cells[1].text = '企业名称'
        header_cells[2].text = '总得分'
        header_cells[3].text = '满分'
        header_cells[4].text = '得分率'
        header_cells[5].text = '评级'

        for idx, data in enumerate(sorted_data, 1):
            info = data['info']
            score = data['score']
            level, evaluation = self.calculator.get_evaluation_level(score['score_percentage'])

            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = info.get('企业名称', '')
            row_cells[2].text = f"{score['total_score']:.2f}"
            row_cells[3].text = f"{score['max_possible_score']:.2f}"
            row_cells[4].text = f"{score['score_percentage']:.2f}%"
            row_cells[5].text = f"{level}({evaluation})"

    def _create_industry_comparison(self, doc, all_data):
        """创建行业对比分析"""
        doc.add_heading('六、行业对比分析', level=1)

        # 按行业分组
        industry_data = {}
        for data in all_data:
            industry = data['info'].get('所属行业', '未知')
            if industry not in industry_data:
                industry_data[industry] = []
            industry_data[industry].append(data['score']['score_percentage'])

        if len(industry_data) > 1:
            doc.add_paragraph('各行业平均得分情况：\n')

            # 创建表格
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'

            header_cells = table.rows[0].cells
            header_cells[0].text = '行业'
            header_cells[1].text = '企业数'
            header_cells[2].text = '平均得分率'
            header_cells[3].text = '最高得分率'
            header_cells[4].text = '最低得分率'

            for industry, scores in sorted(industry_data.items()):
                row_cells = table.add_row().cells
                row_cells[0].text = industry
                row_cells[1].text = str(len(scores))
                row_cells[2].text = f'{np.mean(scores):.2f}%'
                row_cells[3].text = f'{np.max(scores):.2f}%'
                row_cells[4].text = f'{np.min(scores):.2f}%'
        else:
            doc.add_paragraph('参评企业所属行业较为单一，暂无行业对比数据。')

    def _create_common_issues_analysis(self, doc, all_data):
        """创建共性问题分析"""
        doc.add_heading('七、共性问题分析', level=1)

        # 统计各个问题的答题情况
        question_stats = {}

        for data in all_data:
            for question in data['score']['question_details']:
                q_id = question['seq_no']
                if q_id not in question_stats:
                    question_stats[q_id] = {
                        'question': question['question'],
                        'level1': question['level1'],
                        'zero_count': 0,
                        'total_count': 0
                    }
                question_stats[q_id]['total_count'] += 1
                if question['actual_score'] == 0 and question['base_score'] > 0:
                    question_stats[q_id]['zero_count'] += 1

        # 找出得分为0比例最高的问题（共性问题）
        common_issues = [
            (q_id, data) for q_id, data in question_stats.items()
            if data['zero_count'] / data['total_count'] >= 0.3  # 超过30%的企业得分为0
        ]

        common_issues.sort(key=lambda x: x[1]['zero_count'] / x[1]['total_count'], reverse=True)

        if common_issues:
            doc.add_paragraph(f'共识别出{len(common_issues)}个共性问题（超过30%的企业未达标）：\n')

            for idx, (q_id, data) in enumerate(common_issues[:20], 1):  # 显示前20个
                issue_text = f"{idx}. {data['question']}\n"
                issue_text += f"   所属维度：{data['level1']}\n"
                issue_text += f"   未达标企业比例：{data['zero_count']}/{data['total_count']} ({data['zero_count']/data['total_count']*100:.1f}%)"
                doc.add_paragraph(issue_text, style='List Number')
        else:
            doc.add_paragraph('未发现明显的共性问题，各企业表现较为均衡。')

    def _create_best_practices(self, doc, all_data):
        """创建最佳实践"""
        doc.add_heading('八、最佳实践', level=1)

        # 找出得分最高的企业
        top_enterprises = sorted(all_data, key=lambda x: x['score']['score_percentage'], reverse=True)[:3]

        doc.add_paragraph('以下是本次评价中表现优异的企业，可作为行业标杆学习借鉴：\n')

        for idx, data in enumerate(top_enterprises, 1):
            info = data['info']
            score = data['score']

            doc.add_heading(f'{idx}. {info.get("企业名称", "")}', level=3)

            practice_text = f"""
总得分率：{score['score_percentage']:.2f}%
企业类型：{info.get('企业类型', '')}
所属行业：{info.get('所属行业', '')}

优势维度：
            """
            doc.add_paragraph(practice_text.strip())

            # 列出优势维度（得分率>85%的）
            strong_dimensions = [
                (name, data_dict) for name, data_dict in score['score_by_level1'].items()
                if data_dict['percentage'] >= 85
            ]

            for dim_name, dim_data in sorted(strong_dimensions, key=lambda x: x[1]['percentage'], reverse=True):
                doc.add_paragraph(
                    f"- {dim_name}：{dim_data['percentage']:.2f}%",
                    style='List Bullet'
                )

            doc.add_paragraph()

    def _create_overall_recommendations(self, doc, all_data):
        """创建整体建议"""
        doc.add_heading('九、整体建议', level=1)

        # 收集所有维度数据，找出普遍薄弱的维度
        all_level1_data = {}
        for data in all_data:
            for level1, level1_data in data['score']['score_by_level1'].items():
                if level1 not in all_level1_data:
                    all_level1_data[level1] = []
                all_level1_data[level1].append(level1_data['percentage'])

        weak_dimensions = [
            (name, np.mean(scores)) for name, scores in all_level1_data.items()
            if np.mean(scores) < 70
        ]

        if weak_dimensions:
            doc.add_paragraph('根据整体评价结果，建议各参评企业重点关注以下方面的制度建设：\n')

            for idx, (dim_name, avg_score) in enumerate(sorted(weak_dimensions, key=lambda x: x[1]), 1):
                doc.add_paragraph(f'{idx}. {dim_name}（平均得分率：{avg_score:.2f}%）', style='List Number')

        doc.add_paragraph('\n')
        doc.add_heading('通用建议', level=3)

        general_recs = [
            '1. 加强顶层设计，建立健全现代企业制度体系',
            '2. 注重制度执行，避免"制度挂在墙上"的现象',
            '3. 定期开展制度培训，提升全员制度意识',
            '4. 建立制度评估和持续改进机制',
            '5. 学习借鉴行业标杆企业的优秀实践',
            '6. 结合企业实际，不断优化完善制度体系'
        ]

        for rec in general_recs:
            doc.add_paragraph(rec, style='List Bullet')

    def _create_score_details_appendix(self, doc, all_data):
        """创建得分明细附录"""
        doc.add_heading('附录：企业得分明细表', level=1)

        # 创建汇总表
        table = doc.add_table(rows=1, cols=12)
        table.style = 'Light Grid Accent 1'

        # 表头 - 基本信息 + 9个一级指标
        header_cells = table.rows[0].cells
        header_cells[0].text = '企业名称'
        header_cells[1].text = '总得分'
        header_cells[2].text = '得分率'

        # 获取所有一级指标名称
        if all_data:
            level1_names = list(all_data[0]['score']['score_by_level1'].keys())
            for i, name in enumerate(level1_names, 3):
                header_cells[i].text = name[:4]  # 简写

        # 填充数据
        for data in sorted(all_data, key=lambda x: x['score']['score_percentage'], reverse=True):
            info = data['info']
            score = data['score']

            row_cells = table.add_row().cells
            row_cells[0].text = info.get('企业名称', '')
            row_cells[1].text = f"{score['total_score']:.1f}"
            row_cells[2].text = f"{score['score_percentage']:.1f}%"

            for i, level1 in enumerate(level1_names, 3):
                if level1 in score['score_by_level1']:
                    row_cells[i].text = f"{score['score_by_level1'][level1]['percentage']:.1f}%"
                else:
                    row_cells[i].text = 'N/A'


if __name__ == '__main__':
    # 测试整体报告生成
    generator = OverallAnalysisReportGenerator()
    # 需要多个已填写的问卷文件
    # questionnaire_files = ['问卷1.xlsx', '问卷2.xlsx', '问卷3.xlsx']
    # report_path = generator.generate_report(questionnaire_files)
    # print(f"\n整体分析报告已生成: {report_path}")
    print("\n整体分析报告生成器已初始化")
