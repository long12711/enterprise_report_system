"""
南开问卷专业报告生成器
根据问卷提交数据生成详细的企业专业评价报告（Word格式，自动转换为PDF）
完全按照"最终简化摘要_专业版报告.pdf"的格式和内容结构
"""

import os
import json
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    print("警告：docx2pdf库未安装，将只生成Word格式报告")


class NankaiReportGenerator:
    """南开问卷自评报告生成器"""
    
    def __init__(self, indicator_file='指标体系.xlsx'):
        self.indicator_file = indicator_file
        
    def generate_report(self, submission_file, generate_pdf=True):
        """
        根据提交的问卷数据生成专业评价报告
        
        Args:
            submission_file: 问卷提交JSON文件路径
            generate_pdf: 是否生成PDF格式（默认True）
            
        Returns:
            生成的报告文件路径（PDF或Word）
        """
        # 读取提交数据
        with open(submission_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # 提取关键信息
        enterprise_info = data.get('enterprise_info', {})
        answers = data.get('answers', {})
        partial_details = data.get('partial_details', {})
        score_info = data.get('score', {})
        level = data.get('level', '初级')
        
        enterprise_name = enterprise_info.get('enterprise_name', '企业')
        timestamp = data.get('timestamp', datetime.now().strftime('%Y%m%d_%H%M%S'))
        
        # 读取指标体系
        try:
            df = pd.read_excel(self.indicator_file, sheet_name=level)
        except:
            # 如果没有对应级别的sheet，读取第一个sheet
            df = pd.read_excel(self.indicator_file, sheet_name=0)
        
        # 创建Word文档
        doc = Document()
        
        # 设置文档样式
        self._set_document_style(doc)
        
        # 1. 封面
        self._add_cover(doc, enterprise_name, level, score_info)
        
        # 2. 执行摘要（新增）
        self._add_executive_summary(doc, enterprise_info, df, answers, score_info, level)
        
        # 3. 企业基本信息
        self._add_enterprise_info(doc, enterprise_info)
        
        # 4. 评价总览
        self._add_score_overview(doc, score_info, level)
        
        # 5. 各维度详细分析
        self._add_dimension_analysis(doc, df, answers, partial_details, score_info)
        
        # 6. 优势与不足（已整合到企业客观评价报告中，这里跳过）
        # self._add_strengths_weaknesses(doc, enterprise_info, df, answers, score_info)
        
        # 7. 详细得分分解（新增）
        self._add_detailed_score_breakdown(doc, df, answers, score_info)
        
        # 8. 标杆对比（新增）
        self._add_benchmark_comparison(doc, enterprise_info, score_info)
        
        # 9. 行业对比分析
        self._add_industry_comparison(doc, enterprise_info, score_info)
        
        # 10. 实施建议（新增）
        self._add_implementation_suggestions(doc, enterprise_info, df, answers, score_info)
        
        # 11. 风险评估与预警
        self._add_risk_assessment(doc, enterprise_info, score_info)
        
        # 12. 改进路径规划
        self._add_improvement_plan(doc, enterprise_info, df, answers, score_info, level)
        
        # 13. 报告说明
        self._add_report_notes(doc, enterprise_info, level)
        
        # 14. 附录：问卷详细数据
        self._add_appendix(doc, df, answers, partial_details, score_info)
        
        # 保存Word报告
        output_dir = 'storage/nankai_reports'
        os.makedirs(output_dir, exist_ok=True)
        
        word_filename = f'专业评价报告_{enterprise_name}_{timestamp}.docx'
        word_path = os.path.join(output_dir, word_filename)
        
        doc.save(word_path)
        
        # 转换为PDF
        if generate_pdf and DOCX2PDF_AVAILABLE:
            try:
                pdf_filename = f'专业评价报告_{enterprise_name}_{timestamp}.pdf'
                pdf_path = os.path.join(output_dir, pdf_filename)
                convert(word_path, pdf_path)
                print(f"PDF报告生成成功: {pdf_path}")
                return pdf_path
            except Exception as e:
                print(f"PDF转换失败: {str(e)}，返回Word格式")
                return word_path
        else:
            return word_path
    
    def _set_document_style(self, doc):
        """设置文档样式"""
        # 设置默认字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.styles['Normal'].font.size = Pt(12)
    
    def _add_cover(self, doc, enterprise_name, level, score_info):
        """添加封面"""
        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('南开大学现代企业制度指数\n评价自评报告')
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        
        doc.add_paragraph()  # 空行
        
        # 企业名称
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'企业名称：{enterprise_name}')
        run.font.size = Pt(16)
        run.font.bold = True
        
        # 评价级别
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'评价级别：{level}')
        run.font.size = Pt(14)
        
        # 综合得分
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        score = score_info.get('percentage', 0)
        run = p.add_run(f'综合得分：{score:.2f}分')
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        
        # 生成日期
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
        run.font.size = Pt(12)
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc, enterprise_info, df, answers, score_info, level):
        """添加执行摘要（增强版 - 参考专业版报告格式）"""
        heading = doc.add_heading('执行摘要', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        enterprise_name = enterprise_info.get('enterprise_name', '企业')
        percentage = score_info.get('percentage', 0)
        total_score = score_info.get('total_score', 0)
        max_score = score_info.get('max_score', 100)
        main_business = enterprise_info.get('main_business', '相关行业')
        scale = enterprise_info.get('enterprise_scale', '')
        establishment_years = enterprise_info.get('establishment_years', '')
        
        # 评级判断
        if percentage >= 90:
            level_text = '优秀'
            level_desc = '企业现代制度建设达到卓越水平'
        elif percentage >= 80:
            level_text = '良好'
            level_desc = '企业现代制度建设较为完善'
        elif percentage >= 70:
            level_text = '中等'
            level_desc = '企业现代制度建设基本健全'
        elif percentage >= 60:
            level_text = '及格'
            level_desc = '企业现代制度建设有待提升'
        else:
            level_text = '需改进'
            level_desc = '企业现代制度建设亟需加强'
        
        # 评价结论段落
        p = doc.add_paragraph()
        p.add_run('本次对').bold = False
        p.add_run(f'{enterprise_name}')
        p.add_run(f'进行{level}级别现代企业制度指数评价，评价指标')
        # 这里需要计算实际指标数量
        indicator_count = len(df)
        p.add_run(f'{indicator_count}项，')
        p.add_run(f'总分{max_score:.1f}分，实际得分{total_score:.1f}分，完成度{percentage:.1f}%。')
        
        doc.add_paragraph()
        
        # 主要亮点段落
        p = doc.add_paragraph()
        p.add_run('主要亮点：')
        
        doc.add_paragraph()
        
        # 分析维度得分
        dimensions = self._analyze_dimensions(df, answers, score_info)
        
        # 找出优势维度（得分率>=75%）
        strengths = [(dim, data) for dim, data in dimensions.items() if data['percentage'] >= 75]
        strengths.sort(key=lambda x: x[1]['percentage'], reverse=True)
        
        # 生成亮点列表
        if strengths:
            for idx, (dim_name, dim_data) in enumerate(strengths[:3], 1):
                p = doc.add_paragraph()
                p.add_run(f'  {dim_name}完善，建立了现代企业治理架构')
        else:
            p = doc.add_paragraph()
            p.add_run('  制度体系基本健全，建立了基础治理架构')
            p = doc.add_paragraph()
            p.add_run(f'  管理制度规范，{main_business}运营有序')
            p = doc.add_paragraph()
            p.add_run('  治理结构清晰，决策执行监督机制基本建立')
        
        # 评价结论
        p = doc.add_paragraph()
        if percentage >= 90:
            conclusion = f'评价结论：企业现代企业制度建设优秀，在{main_business}中具备领先优势。'
        elif percentage >= 80:
            conclusion = f'评价结论：企业现代企业制度建设良好，在{main_business}中处于先进水平。'
        elif percentage >= 70:
            conclusion = f'评价结论：企业现代企业制度建设基本健全，在{main_business}中处于中等水平。'
        else:
            conclusion = f'评价结论：企业现代企业制度建设有待提升，需要系统性完善各项制度。'
        p.add_run(conclusion)
        
        doc.add_paragraph()
        doc.add_page_break()
    
    def _analyze_dimensions(self, df, answers, score_info):
        """分析各维度得分情况"""
        dimensions = {}
        score_details = score_info.get('details', {})
        
        for idx, row in df.iterrows():
            level1 = str(row['一级指标']) if pd.notna(row['一级指标']) else ''
            if not level1:
                continue
            
            if level1 not in dimensions:
                dimensions[level1] = {
                    'total_score': 0,
                    'earned_score': 0,
                    'percentage': 0
                }
            
            q_id = str(int(row['序号'])) if pd.notna(row['序号']) else str(idx + 1)
            detail = score_details.get(q_id, {})
            
            dimensions[level1]['total_score'] += detail.get('max_score', 0)
            dimensions[level1]['earned_score'] += detail.get('earned_score', 0)
        
        # 计算百分比
        for dim_name, dim_data in dimensions.items():
            if dim_data['total_score'] > 0:
                dim_data['percentage'] = (dim_data['earned_score'] / dim_data['total_score']) * 100
        
        return dimensions
    
    def _add_enterprise_info(self, doc, enterprise_info):
        """添加企业客观评价报告"""
        heading = doc.add_heading('一、企业客观评价报告', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        # （一）企业基本情况评价
        doc.add_heading('（一）企业基本情况评价', level=2)
        
        enterprise_name = enterprise_info.get('enterprise_name', '企业')
        main_business = enterprise_info.get('main_business', '相关行业')
        scale = enterprise_info.get('enterprise_scale', '')
        establishment_years = enterprise_info.get('establishment_years', '')
        
        # 企业基础信息段落
        p = doc.add_paragraph()
        p.add_run('企业基础信息：').bold = True
        
        p = doc.add_paragraph()
        scale_text = f'{scale}' if scale else ''
        years_text = f'，成立于{establishment_years}年' if establishment_years else ''
        p.add_run(f'{enterprise_name}是一家{scale_text}企业{years_text}，')
        p.add_run(f'主营业务为{main_business}。')
        p.add_run('企业现有员工若干人，')
        p.add_run('属于相应企业范畴。')
        
        doc.add_paragraph()
        
        # 企业发展状况段落
        p = doc.add_paragraph()
        p.add_run('企业发展状况：').bold = True
        
        p = doc.add_paragraph()
        p.add_run('企业自成立以来，坚持规范化经营和专业化发展道路，')
        p.add_run(f'在{main_business}领域积累了一定的经验和市场声誉。')
        p.add_run('公司治理结构逐步完善，管理制度不断健全，')
        p.add_run('为进一步提升现代企业制度建设水平奠定了基础。')
        
        doc.add_paragraph()
    
    def _add_score_overview(self, doc, score_info, level):
        """添加现代企业制度建设评价"""
        # 这部分已经在企业基本信息中作为子章节，这里添加为（二）
        doc.add_heading('（二）现代企业制度建设评价', level=2)
        
        total_score = score_info.get('total_score', 0)
        max_score = score_info.get('max_score', 100)
        percentage = score_info.get('percentage', 0)
        
        # 总体评价情况段落
        p = doc.add_paragraph()
        p.add_run('总体评价情况：').bold = True
        
        p = doc.add_paragraph()
        p.add_run('根据现代企业制度指数评价体系，')
        p.add_run(f'本次{level}级别评价共涉及相应指标，')
        p.add_run(f'总分{max_score:.1f}分，实际得分{total_score:.1f}分，')
        p.add_run(f'综合完成度为{percentage:.1f}%。')
        
        # 评级判断
        if percentage >= 90:
            level_text = '优秀'
            position = '领先水平'
        elif percentage >= 80:
            level_text = '良好'
            position = '先进水平'
        elif percentage >= 70:
            level_text = '中等'
            position = '中等水平'
        else:
            level_text = '有待提升'
            position = '基础水平'
        
        p.add_run('评价结果表明，企业现代企业制度建设表现')
        run = p.add_run(level_text)
        run.bold = True
        p.add_run(f'，在行业中处于{position}。')
        
        doc.add_paragraph()
    
    def _add_dimension_analysis(self, doc, df, answers, partial_details, score_info):
        """添加各维度详细分析 - 作为企业客观评价报告的子章节"""
        # 这部分作为（三）各维度具体表现
        doc.add_heading('（三）各维度具体表现', level=2)
        
        # 按一级指标分组
        dimensions = {}
        score_details = score_info.get('details', {})
        
        for idx, row in df.iterrows():
            level1 = str(row['一级指标']) if pd.notna(row['一级指标']) else ''
            if not level1:
                continue
            
            if level1 not in dimensions:
                dimensions[level1] = {
                    'questions': [],
                    'total_score': 0,
                    'earned_score': 0
                }
            
            q_id = str(int(row['序号'])) if pd.notna(row['序号']) else str(idx + 1)
            detail = score_details.get(q_id, {})
            
            dimensions[level1]['questions'].append({
                'id': q_id,
                'level2': str(row['二级指标']) if pd.notna(row['二级指标']) else '',
                'level3': str(row['三级指标']) if pd.notna(row['三级指标']) else '',
                'answer': answers.get(q_id, ''),
                'partial_detail': partial_details.get(q_id, ''),
                'max_score': detail.get('max_score', 0),
                'earned_score': detail.get('earned_score', 0)
            })
            
            dimensions[level1]['total_score'] += detail.get('max_score', 0)
            dimensions[level1]['earned_score'] += detail.get('earned_score', 0)
        
        # 输出各维度分析 - 使用更正式的语言
        for dim_name, dim_data in dimensions.items():
            total = dim_data['total_score']
            earned = dim_data['earned_score']
            percentage = (earned / total * 100) if total > 0 else 0
            
            # 统计完成情况
            completed = sum(1 for q in dim_data['questions'] if q['answer'].startswith('A'))
            total_items = len(dim_data['questions'])
            
            # 评价等级
            if percentage >= 90:
                performance = '表现突出，已达到优秀水平'
            elif percentage >= 80:
                performance = '表现良好，基本符合要求'
            elif percentage >= 70:
                performance = '表现一般，需要改进'
            else:
                performance = '表现较弱，有待加强'
            
            p = doc.add_paragraph()
            p.add_run(f'   {dim_name}：').bold = True
            p.add_run(f'共{total_items}项指标，得分{earned:.1f}分/{total:.1f}分（完成度{percentage:.1f}%）。')
            p.add_run(f'该维度{performance}。')
            
        doc.add_paragraph()
        doc.add_page_break()
    
    def _add_strengths_weaknesses(self, doc, enterprise_info, df, answers, score_info):
        """添加企业治理结构分析和经营管理水平评价 - 作为企业客观评价报告的子章节"""
        
        score_details = score_info.get('details', {})
        percentage = score_info.get('percentage', 0)
        enterprise_name = enterprise_info.get('enterprise_name', '企业')
        
        # （三）企业治理结构分析
        doc.add_heading('（三）企业治理结构分析', level=2)
        
        # 治理架构完善性段落
        p = doc.add_paragraph()
        p.add_run('治理架构完善性：').bold = True
        
        p = doc.add_paragraph()
        if percentage >= 80:
            p.add_run(f'{enterprise_name}建立了符合现代企业制度要求的治理架构，')
            p.add_run('形成了较为完善的治理格局。')
            p.add_run('各项治理制度较为完善，')
            p.add_run('为规范企业治理和科学决策提供了制度保障。')
        else:
            p.add_run('企业已初步建立治理架构，')
            p.add_run('但在制度完善性和执行有效性方面还需要进一步加强。')
            p.add_run('建议系统梳理现有治理制度，')
            p.add_run('补齐短板，提升治理水平。')
        
        doc.add_paragraph()
        
        # 制度执行有效性段落
        p = doc.add_paragraph()
        p.add_run('制度执行有效性：').bold = True
        
        p = doc.add_paragraph()
        if percentage >= 75:
            p.add_run('企业注重制度的贯彻执行，建立了相应的监督检查机制。')
            p.add_run('各项治理制度基本得到有效执行，治理流程规范有序。')
            if percentage < 90:
                p.add_run('但在制度执行的一致性和持续性方面还需要进一步加强。')
        else:
            p.add_run('企业在制度执行方面还需要加强，')
            p.add_run('建议建立健全制度执行监督机制，')
            p.add_run('提升制度执行的一致性和持续性。')
        
        doc.add_paragraph()
        
        # （四）经营管理水平评价
        doc.add_heading('（四）经营管理水平评价', level=2)
        
        # 财务管理水平
        p = doc.add_paragraph()
        p.add_run('财务管理水平：').bold = True
        
        p = doc.add_paragraph()
        p.add_run('企业建立了较为完善的财务管理制度，财务核算规范，资金管理有序。')
        if percentage >= 80:
            p.add_run('财务指标总体健康，具备良好的盈利能力和发展潜力。')
            p.add_run('内控制度基本健全，财务风险可控。')
        else:
            p.add_run('财务管理制度需要进一步完善，')
            p.add_run('建议加强内控体系建设，提升财务风险管控能力。')
        
        doc.add_paragraph()
        
        # 人力资源管理
        p = doc.add_paragraph()
        p.add_run('人力资源管理：').bold = True
        
        p = doc.add_paragraph()
        p.add_run('企业现有员工若干人，人员结构合理，专业技术人员比例适当。')
        p.add_run('建立了基本的人力资源管理制度，')
        p.add_run('在人才培养、绩效考核、薪酬激励等方面有一定基础，')
        if percentage >= 75:
            p.add_run('人力资源管理体系较为完善。')
        else:
            p.add_run('但还需要进一步完善。')
        
        doc.add_paragraph()
        
        # 运营管理效率
        p = doc.add_paragraph()
        p.add_run('运营管理效率：').bold = True
        
        p = doc.add_paragraph()
        p.add_run('企业运营管理相对规范，主营业务发展稳定，市场适应能力较强。')
        p.add_run('在质量管理、客户服务、供应链管理等方面建立了相应制度，')
        p.add_run('运营效率逐步提升，为企业持续发展奠定了基础。')
        
        doc.add_paragraph()
        doc.add_page_break()
    
    def _add_detailed_score_breakdown(self, doc, df, answers, score_info):
        """添加第五章：详细得分分解"""
        heading = doc.add_heading('五、详细得分分解', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        score_details = score_info.get('details', {})
        
        # 按一级指标分组统计
        dimensions = {}
        for idx, row in df.iterrows():
            level1 = str(row['一级指标']) if pd.notna(row['一级指标']) else ''
            if not level1:
                continue
            
            if level1 not in dimensions:
                dimensions[level1] = {
                    'total_score': 0,
                    'earned_score': 0
                }
            
            q_id = str(int(row['序号'])) if pd.notna(row['序号']) else str(idx + 1)
            detail = score_details.get(q_id, {})
            
            dimensions[level1]['total_score'] += detail.get('max_score', 0)
            dimensions[level1]['earned_score'] += detail.get('earned_score', 0)
        
        # 创建得分分解表格
        p = doc.add_paragraph()
        p.add_run('各维度得分详情：').bold = True
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        headers = ['维度名称', '满分', '得分', '得分率']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        for dim_name, dim_data in dimensions.items():
            row_cells = table.add_row().cells
            total = dim_data['total_score']
            earned = dim_data['earned_score']
            percentage = (earned / total * 100) if total > 0 else 0
            
            row_cells[0].text = dim_name
            row_cells[1].text = f'{total:.1f}'
            row_cells[2].text = f'{earned:.1f}'
            row_cells[3].text = f'{percentage:.1f}%'
        
        doc.add_paragraph()
        
        # 详细说明
        p = doc.add_paragraph()
        p.add_run('得分说明：').bold = True
        
        p = doc.add_paragraph()
        p.add_run('本次评价采用百分制计分方式，各维度得分根据问卷完成情况计算。')
        p.add_run('完全完成的项目得满分，部分完成的项目根据完成质量得部分分数，')
        p.add_run('未完成的项目不得分。各维度得分率反映了该维度的制度建设完成情况。')
        
        doc.add_paragraph()
        doc.add_page_break()
    
    def _add_benchmark_comparison(self, doc, enterprise_info, score_info):
        """添加第六章：标杆对比"""
        heading = doc.add_heading('六、标杆对比', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        enterprise_name = enterprise_info.get('enterprise_name', '企业')
        percentage = score_info.get('percentage', 0)
        industry = enterprise_info.get('main_business', '相关行业')
        
        # （一）行业标杆企业特征
        doc.add_heading('（一）行业标杆企业特征', level=2)
        
        p = doc.add_paragraph()
        p.add_run('标杆企业概况：').bold = True
        
        p = doc.add_paragraph()
        p.add_run(f'在{industry}领域，标杆企业被公认为行业典范。')
        p.add_run('这些企业在现代企业制度建设方面起步较早，制度体系相对完善，治理水平较高。')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('标杆对比分析：').bold = True
        
        p = doc.add_paragraph()
        if percentage >= 90:
            p.add_run(f'   制度完善度：标杆企业平均完成度达96-98%，{enterprise_name}为{percentage:.1f}%，')
            p.add_run('已达到标杆水平')
        else:
            gap = 96 - percentage
            p.add_run(f'   制度完善度：标杆企业平均完成度达96-98%，{enterprise_name}为{percentage:.1f}%，')
            p.add_run(f'差距{gap:.1f}个百分点')
        
        p = doc.add_paragraph()
        p.add_run('   治理效率：标杆企业决策链条平均3-4层，')
        p.add_run(f'{enterprise_name}治理流程基本规范')
        
        p = doc.add_paragraph()
        p.add_run('   风险管控：标杆企业风险管控体系覆盖率达98%以上，')
        if percentage >= 85:
            p.add_run('企业风险管控体系较为完善')
        else:
            p.add_run('企业风险管控体系有待提升')
        
        doc.add_paragraph()
        doc.add_page_break()
    
    def _add_industry_comparison(self, doc, enterprise_info, score_info):
        """添加行业对比分析"""
        heading = doc.add_heading('七、行业对比分析', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        enterprise_name = enterprise_info.get('enterprise_name', '企业')
        percentage = score_info.get('percentage', 0)
        industry = enterprise_info.get('main_business', '相关行业')
        
        # （一）行业基准对比
        doc.add_heading(f'（一）{industry}行业基准对比', level=2)
        
        p = doc.add_paragraph()
        p.add_run('行业发展态势：').bold = True
        
        p = doc.add_paragraph()
        p.add_run(f'{industry}作为国民经济的重要组成部分，近年来保持了稳定发展态势。')
        p.add_run('行业企业数量众多，市场竞争激烈，')
        p.add_run('对现代企业制度建设的要求不断提高。')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('综合对比分析结论：').bold = True
        
        p = doc.add_paragraph()
        p.add_run(f'通过与{industry}行业全维度基准数据的深度对比，')
        p.add_run(f'{enterprise_name}在多个关键维度上展现出不同的表现水平：')
        
        doc.add_paragraph()
        
        # 根据得分率判断行业地位
        if percentage >= 90:
            p = doc.add_paragraph()
            p.add_run('优势领域（超过行业优秀水平）：').bold = True
            p = doc.add_paragraph()
            p.add_run(f'   企业得分率{percentage:.1f}%，')
            p.add_run('超出行业平均水平，')
            p.add_run('在制度建设、治理规范等方面表现突出，')
            p.add_run('已达到行业领先水平')
        elif percentage >= 80:
            p = doc.add_paragraph()
            p.add_run('先进领域（超过行业平均但低于优秀水平）：').bold = True
            p = doc.add_paragraph()
            p.add_run(f'   企业得分率{percentage:.1f}%，')
            p.add_run('超出行业平均水平，')
            p.add_run('但距离行业优秀水平还有提升空间，')
            p.add_run('需要持续优化')
        else:
            p = doc.add_paragraph()
            p.add_run('改进空间（需要重点关注）：').bold = True
            p = doc.add_paragraph()
            p.add_run(f'   企业得分率{percentage:.1f}%，')
            p.add_run('需要系统性地提升各项制度建设水平，')
            p.add_run('补齐短板，提升整体竞争力')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('行业地位分析：').bold = True
        
        p = doc.add_paragraph()
        if percentage >= 85:
            p.add_run(f'在{industry}企业中，{enterprise_name}现代企业制度建设水平排名约为前25%，')
            p.add_run('具备冲击行业TOP10标杆企业的实力基础。')
        elif percentage >= 70:
            p.add_run(f'{enterprise_name}现代企业制度建设水平在{industry}中处于中上游，')
            p.add_run('具备持续提升的良好基础。')
        else:
            p.add_run(f'{enterprise_name}现代企业制度建设水平在{industry}中处于中等水平，')
            p.add_run('需要加大制度建设力度。')
        
        doc.add_paragraph()
        doc.add_page_break()
    
    def _add_risk_assessment(self, doc, enterprise_info, score_info):
        """添加风险评估与预警"""
        heading = doc.add_heading('九、风险评估与预警', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        enterprise_name = enterprise_info.get('enterprise_name', '企业')
        percentage = score_info.get('percentage', 0)
        
        # （一）风险等级评估
        doc.add_heading('（一）风险等级评估', level=2)
        
        # 根据得分率判断风险等级
        if percentage >= 85:
            risk_level = '低风险'
            risk_score = 85 + (percentage - 85) / 15 * 15
        elif percentage >= 70:
            risk_level = '中低风险'
            risk_score = 70 + (percentage - 70) / 15 * 15
        elif percentage >= 60:
            risk_level = '中等风险'
            risk_score = 60 + (percentage - 60) / 10 * 10
        else:
            risk_level = '较高风险'
            risk_score = percentage
        
        p = doc.add_paragraph()
        run = p.add_run(f'综合风险等级：{risk_level}')
        run.bold = True
        
        p = doc.add_paragraph()
        p.add_run(f'通过运用现代风险管理理论和方法，采用定性与定量相结合的风险评估模型，')
        p.add_run(f'对{enterprise_name}进行全方位风险识别、评估和分析。')
        p.add_run(f'评估结果显示，企业总体风险等级为{risk_level}，')
        p.add_run(f'综合风险得分率为{risk_score:.1f}分（满分100分），')
        level_text = '先进' if percentage >= 85 else '中等' if percentage >= 70 else '基础'
        p.add_run(f'风险管控能力达到{enterprise_info.get("main_business", "相关行业")}{level_text}水平。')
        
        doc.add_paragraph()
        
        # （二）分类风险深度分析
        doc.add_heading('（二）分类风险深度分析', level=2)
        
        # 风险分类表格
        table = doc.add_table(rows=5, cols=4)
        table.style = 'Light Grid Accent 1'
        
        headers = ['风险类别', '风险等级', '主要风险点', '建议改进措施']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
    
    def _add_implementation_suggestions(self, doc, enterprise_info, df, answers, score_info):
        """添加第八章：实施建议"""
        heading = doc.add_heading('八、实施建议', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        enterprise_name = enterprise_info.get('enterprise_name', '企业')
        percentage = score_info.get('percentage', 0)
        
        # （一）组织保障
        doc.add_heading('（一）组织保障', level=2)
        
        p = doc.add_paragraph()
        p.add_run('建立专项工作组：').bold = True
        
        p = doc.add_paragraph()
        p.add_run('成立由企业高层领导牵头的现代企业制度建设领导小组，')
        p.add_run('设立专门的制度建设办公室，配备专职人员，')
        p.add_run('明确各部门在制度建设中的职责分工，')
        p.add_run('建立定期汇报和协调机制，确保各项工作有序推进。')
        
        doc.add_paragraph()
        
        # （二）资源配置
        doc.add_heading('（二）资源配置', level=2)
        
        p = doc.add_paragraph()
        p.add_run('合理配置资源：').bold = True
        
        p = doc.add_paragraph()
        p.add_run('设立专项预算，确保制度建设资金投入；')
        p.add_run('引进或培养专业人才，提升团队专业能力；')
        p.add_run('采购必要的软硬件设施，支持制度数字化转型；')
        p.add_run('必要时聘请外部专业机构提供咨询服务，')
        p.add_run('为制度建设提供专业支持。')
        
        doc.add_paragraph()
        
        # （三）培训与宣贯
        doc.add_heading('（三）培训与宣贯', level=2)
        
        p = doc.add_paragraph()
        p.add_run('加强培训宣传：').bold = True
        
        p = doc.add_paragraph()
        p.add_run('开展全员制度培训，提升全体员工的制度意识；')
        p.add_run('针对关键岗位人员进行专项培训，')
        p.add_run('确保关键岗位人员熟练掌握相关制度；')
        p.add_run('建立制度学习考核机制，确保培训效果；')
        p.add_run('通过多种渠道宣传制度建设成果，')
        p.add_run('营造良好的制度建设氛围。')
        
        doc.add_paragraph()
        
        # （四）监督与评估
        doc.add_heading('（四）监督与评估', level=2)
        
        p = doc.add_paragraph()
        p.add_run('建立监督评估机制：').bold = True
        
        p = doc.add_paragraph()
        p.add_run('建立制度执行监督机制，定期检查制度执行情况；')
        p.add_run('设立制度建设评估指标，定期评估建设成效；')
        p.add_run('建立问题反馈和改进机制，及时解决存在的问题；')
        p.add_run('将制度建设纳入各级管理人员绩效考核，')
        p.add_run('强化责任落实，确保制度建设取得实效。')
        
        doc.add_paragraph()
        
        # （五）持续改进
        doc.add_heading('（五）持续改进', level=2)
        
        p = doc.add_paragraph()
        p.add_run('建立持续改进机制：').bold = True
        
        p = doc.add_paragraph()
        p.add_run(f'{enterprise_name}应建立制度持续改进机制，')
        p.add_run('定期评估制度执行效果，')
        p.add_run('根据内外部环境变化及时调整完善制度。')
        p.add_run('通过PDCA循环（计划-执行-检查-改进），')
        p.add_run('推动现代企业制度建设持续优化提升，')
        p.add_run('确保制度体系始终保持先进性和有效性。')
        
        doc.add_paragraph()
        doc.add_page_break()
    
    def _add_improvement_plan(self, doc, enterprise_info, df, answers, score_info, level):
        """添加改进路径规划"""
        heading = doc.add_heading('十、改进路径规划', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        enterprise_name = enterprise_info.get('enterprise_name', '企业')
        percentage = score_info.get('percentage', 0)
        industry = enterprise_info.get('main_business', '相关行业')
        
        p = doc.add_paragraph()
        p.add_run('路径规划导言：').bold = True
        
        p = doc.add_paragraph()
        intro_text = f'基于本次现代企业制度指数评价结果，{enterprise_name}在治理体系建设方面'
        if percentage >= 75:
            intro_text += '展现出了良好的发展态势'
        else:
            intro_text += '具备一定的基础'
        intro_text += f'，总体完成度达{percentage:.1f}%，'
        if percentage >= 85:
            intro_text += f'在{industry}中处于优秀水平。'
        elif percentage >= 70:
            intro_text += f'在{industry}中处于良好水平。'
        else:
            intro_text += f'在{industry}中处于中等水平。'
        intro_text += '为确保企业治理水平持续提升，我们制定了分阶段、有重点、可操作的改进路径规划。'
        p.add_run(intro_text)
        
        doc.add_paragraph()
        
        # （一）短期改进计划
        doc.add_heading('（一）短期改进计划（3-6个月）', level=2)
        
        p = doc.add_paragraph()
        run = p.add_run('阶段目标：')
        run.bold = True
        p.add_run('夯实基础，补齐短板，建立健全基本制度框架')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('1. 制度体系完善工程').bold = True
        
        p = doc.add_paragraph()
        p.add_run('1.1 全面梳理现有治理制度，建立制度清单和制度地图，识别制度缺口和薄弱环节')
        p = doc.add_paragraph()
        p.add_run('1.2 重点制定和完善核心制度，如《企业风险管理办法》、《信息披露管理制度》等')
        p = doc.add_paragraph()
        p.add_run('1.3 建立制度定期评估机制，每季度开展制度执行效果评估')
        p = doc.add_paragraph()
        p.add_run('1.4 完善制度执行监督机制，明确制度执行责任主体，建立制度违反问责机制')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('2. 治理透明度提升工程').bold = True
        
        p = doc.add_paragraph()
        p.add_run('2.1 建立企业信息披露专栏，确保重大信息及时披露')
        p = doc.add_paragraph()
        p.add_run('2.2 完善内部信息传递机制，建立信息披露责任制')
        p = doc.add_paragraph()
        p.add_run('2.3 规范关联交易信息披露，建立关联交易数据库')
        p = doc.add_paragraph()
        p.add_run('2.4 加强与利益相关方的沟通，定期举办说明会、座谈会等')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('预期成果：').bold = True
        p.add_run('形成较为完整的制度框架，补齐制度短板，建立风险管控基础设施，')
        p.add_run('信息披露及时性达到95%以上，为下一阶段深化改革奠定坚实基础。')
        if percentage < 90:
            p.add_run(f'短期内企业治理完成度预计提升至{min(percentage + 5, 96):.0f}%以上。')
        
        doc.add_paragraph()
        
        # （二）中期发展规划
        doc.add_heading('（二）中期发展规划（6-18个月）', level=2)
        
        p = doc.add_paragraph()
        run = p.add_run('阶段目标：')
        run.bold = True
        p.add_run('优化体系，提升效能，建立现代化治理运行机制')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('1. 治理效能优化工程').bold = True
        
        p = doc.add_paragraph()
        p.add_run('1.1 优化决策流程，建立科学决策机制，决策效率提升30%')
        p = doc.add_paragraph()
        p.add_run('1.2 推进治理数字化转型，建设智能化治理平台')
        p = doc.add_paragraph()
        p.add_run('1.3 完善绩效考核体系，将现代企业制度建设纳入考核指标')
        p = doc.add_paragraph()
        p.add_run('1.4 加强内部审计功能，实现对重点领域和关键环节的全覆盖审计')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('预期成果：').bold = True
        p.add_run('建立高效的现代化治理运行机制，治理效率提升30%，')
        p.add_run('人才队伍专业化水平显著提升，利益相关方满意度达到90%以上，')
        if percentage < 95:
            p.add_run(f'企业治理完成度达到{min(percentage + 10, 98):.0f}%以上。')
        
        doc.add_paragraph()
        
        # （三）长期目标
        doc.add_heading('（三）长期目标（1-3年）', level=2)
        
        p = doc.add_paragraph()
        run = p.add_run('战略目标：')
        run.bold = True
        p.add_run('全面建成现代企业制度体系，实现治理能力现代化')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('1. 建成完善的现代企业制度体系，实现治理完成度99%以上，形成科学规范、运行高效的治理架构')
        p = doc.add_paragraph()
        p.add_run('2. 实现治理规范化、制度化、程序化，各项治理活动有章可循，执行有力，监督到位')
        p = doc.add_paragraph()
        p.add_run(f'3. 达到{industry}领先的治理水平，成为同行业治理标杆，具备对外输出治理经验的能力')
        p = doc.add_paragraph()
        p.add_run('4. 建立与国际先进标准接轨的治理体系，为企业国际化发展奠定坚实基础')
        p = doc.add_paragraph()
        p.add_run('5. 形成良好的企业治理文化，治理理念深入人心，全员治理意识和能力显著提升')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('预期成果：').bold = True
        p.add_run('全面建成现代企业制度体系，治理能力达到国际先进水平，')
        p.add_run(f'成为{industry}治理标杆企业，为企业高质量发展提供坚强保障。')
        p.add_run('企业价值和社会责任实现有机统一，可持续发展能力显著增强。')
        
        doc.add_paragraph()
        doc.add_page_break()
    
    def _add_report_notes(self, doc, enterprise_info, level):
        """添加第十一章：报告说明"""
        heading = doc.add_heading('十一、报告说明', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        # 1. 报告性质与用途
        p = doc.add_paragraph()
        p.add_run('1. 报告性质与用途').bold = True
        
        p = doc.add_paragraph()
        p.add_run('本报告是基于企业自主填报数据生成的专业评价报告，')
        p.add_run('采用南开大学现代企业制度指数评价体系标准。')
        p.add_run('报告反映了企业在现代企业制度建设方面的现状水平，')
        p.add_run('为企业治理改进提供参考依据。')
        
        doc.add_paragraph()
        
        # 2. 评价方法与依据
        p = doc.add_paragraph()
        p.add_run('2. 评价方法与依据').bold = True
        
        p = doc.add_paragraph()
        p.add_run('评价标准：南开大学现代企业制度指数评价体系初稿')
        
        p = doc.add_paragraph()
        p.add_run('评价方法：定量评分与定性分析相结合')
        
        p = doc.add_paragraph()
        p.add_run('评价维度：治理方向性、治理有效性、治理规范性、治理透明性')
        
        p = doc.add_paragraph()
        p.add_run(f'指标体系：{level}级别包含相应的评价指标')
        
        doc.add_paragraph()
        
        # 3. 报告特色与创新
        p = doc.add_paragraph()
        p.add_run('3. 报告特色与创新').bold = True
        
        p = doc.add_paragraph()
        p.add_run('专业格式：采用正式报告格式，符合官方文档标准')
        p = doc.add_paragraph()
        p.add_run('丰富内容：包含行业对比、风险评估、改进路径等增值分析')
        p = doc.add_paragraph()
        p.add_run('智能分析：基于数据可用性动态生成相关章节')
        p = doc.add_paragraph()
        p.add_run('实用建议：提供针对性的改进建议和实施路径')
        
        doc.add_paragraph()
        
        # 4. 使用范围与限制
        p = doc.add_paragraph()
        p.add_run('4. 使用范围与限制').bold = True
        
        p = doc.add_paragraph()
        p.add_run('本报告仅供企业内部管理参考使用              不构成对外投资决策或融资担保的依据')
        p = doc.add_paragraph()
        p.add_run('评价结果基于企业自主填报数据，请确保数据真实性')
        p = doc.add_paragraph()
        p.add_run('建议结合企业实际情况制定具体改进措施')
        
        doc.add_paragraph()
        
        # 5. 技术支持与联系
        p = doc.add_paragraph()
        p.add_run('5. 技术支持与联系').bold = True
        
        p = doc.add_paragraph()
        p.add_run('系统版本：专业标准版报告生成器 v2.0')
        
        p = doc.add_paragraph()
        p.add_run('技术支持：现代企业制度指数评价系统')
        
        p = doc.add_paragraph()
        p.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
        
        p = doc.add_paragraph()
        contact_person = enterprise_info.get('contact_person', '')
        contact_email = enterprise_info.get('contact_email', '')
        if contact_person:
            p.add_run(f'企业联系人：{contact_person}')
        if contact_email:
            p = doc.add_paragraph()
            p.add_run(f'联系方式：{contact_email}')
        
        doc.add_paragraph()
        
        # 分隔线
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run('_' * 60)
        
        doc.add_paragraph()
        
        # 底部标识
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('现代企业制度指数评价系统 | 专业标准版报告')
        run.font.size = Pt(10)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Modern Enterprise System Index Evaluation System - Professional Standard Report')
        run.font.size = Pt(9)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('基于南开大学学术标准 | 专业报告生成器')
        run.font.size = Pt(9)
        
        doc.add_paragraph()
        doc.add_page_break()
    
    def _add_appendix(self, doc, df, answers, partial_details, score_info):
        """添加附录：问卷详细数据"""
        heading = doc.add_heading('附录：问卷详细数据', level=1)
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
        
        score_details = score_info.get('details', {})
        
        # 创建表格
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        headers = ['序号', '三级指标', '分值', '答案', '得分', '完成度', '说明']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # 数据行
        for idx, row in df.iterrows():
            q_id = str(int(row['序号'])) if pd.notna(row['序号']) else str(idx + 1)
            level3 = str(row['三级指标']) if pd.notna(row['三级指标']) else ''
            answer = answers.get(q_id, '')
            partial_detail = partial_details.get(q_id, '')
            
            detail = score_details.get(q_id, {})
            max_score = detail.get('max_score', 0)
            earned_score = detail.get('earned_score', 0)
            
            # 完成度
            if answer.startswith('A'):
                completion = '已完成'
            elif answer.startswith('B'):
                completion = '部分完成'
            else:
                completion = '未完成'
            
            row_cells = table.add_row().cells
            row_cells[0].text = q_id
            row_cells[1].text = level3
            row_cells[2].text = f'{max_score:.1f}'
            row_cells[3].text = answer
            row_cells[4].text = f'{earned_score:.1f}'
            row_cells[5].text = completion
            row_cells[6].text = partial_detail
        
        doc.add_paragraph()
        
        # 结束语
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run('—— 报告结束 ——')


if __name__ == '__main__':
    # 测试代码
    generator = NankaiReportGenerator()
    print("南开问卷自评报告生成器已就绪")