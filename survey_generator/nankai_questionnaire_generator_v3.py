"""
南开大学现代企业制度指数评价问卷生成器 V3
根据评分准则精确生成问卷选项

改进点：
1. 根据评分准则类型生成精确的选项
2. 支持多项累计、全部完成、分级评分等复杂类型
3. 选项与评分规则完全对应
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from datetime import datetime
from typing import Dict, List, Tuple
import re


class NankaiQuestionnaireGeneratorV3:
    """南开大学问卷生成器 V3 - 基于评分准则"""
    
    def __init__(self, excel_path: str):
        """
        初始化生成器
        
        Args:
            excel_path: 南开大学指标体系Excel文件路径
        """
        self.excel_path = excel_path
        self.data = {}
        self._load_data()
        
    def _load_data(self):
        """加载Excel数据"""
        try:
            # 读取三个级别的数据
            self.data['初级'] = pd.read_excel(self.excel_path, sheet_name='初级')
            self.data['中级'] = pd.read_excel(self.excel_path, sheet_name='中级')
            self.data['高级'] = pd.read_excel(self.excel_path, sheet_name='高级')
            print(f"[OK] 成功加载指标数据：初级 {len(self.data['初级'])} 条，中级 {len(self.data['中级'])} 条，高级 {len(self.data['高级'])} 条")
        except Exception as e:
            raise Exception(f"[ERROR] 加载Excel文件失败: {str(e)}")
    
    def _identify_rule_type(self, rule_text: str, score_value: str) -> str:
        """
        识别评分规则类型
        
        Returns:
            规则类型：binary、multi_item、all_required、negative、graded
        """
        if pd.isna(rule_text):
            return 'binary'
        
        rule_text = str(rule_text).strip()
        score_str = str(score_value).strip() if pd.notna(score_value) else "1"
        
        # 否决性指标（负分）
        if score_str.startswith('-'):
            return 'negative'
        
        # 全部完成型
        if '全部完成' in rule_text or '部分完成不得分' in rule_text:
            return 'all_required'
        
        # 多项累计评分
        if '项' in rule_text and ('得' in rule_text or '分' in rule_text):
            if re.search(r'\d+-\d+项', rule_text) or re.search(r'\d+项', rule_text):
                return 'multi_item'
        
        # 分级评分（分值范围如0-2）
        if '-' in score_str:
            parts = score_str.split('-')
            try:
                min_val = float(parts[0])
                max_val = float(parts[1])
                if max_val - min_val > 1:
                    return 'graded'
            except:
                pass
        
        # 二元评分（默认）
        return 'binary'
    
    def _parse_scoring_criteria(self, rule_text: str, score_value: str) -> Tuple[List[str], str, str]:
        """
        根据评分准则生成选项
        
        Args:
            rule_text: 评分准则文本
            score_value: 分值
            
        Returns:
            (选项列表, 填空题提示, 规则类型)
        """
        if pd.isna(rule_text):
            return (["A. 是", "B. 否"], "", "binary")
        
        rule_text = str(rule_text).strip()
        rule_type = self._identify_rule_type(rule_text, score_value)
        
        options = []
        fill_blank_hint = ""
        
        if rule_type == 'binary':
            # 二元评分：是/否
            if '建立' in rule_text or '制定' in rule_text or '设立' in rule_text:
                options = ["A. 是（已建立/制定/设立）", "B. 否（未建立/制定/设立）"]
            else:
                options = ["A. 是", "B. 否"]
        
        elif rule_type == 'all_required':
            # 全部完成型
            options = [
                "A. 全部完成",
                "B. 部分完成",
                "C. 均未完成"
            ]
            fill_blank_hint = "请说明完成情况：______"
        
        elif rule_type == 'multi_item':
            # 多项累计：从规则文本中提取项数要求
            options = self._extract_multi_item_options(rule_text)
            fill_blank_hint = "请说明具体完成了哪些项：______"
        
        elif rule_type == 'negative':
            # 否决性指标
            options = [
                "A. 是（存在违规情况）",
                "B. 否（不存在违规情况）"
            ]
            fill_blank_hint = "如存在违规，请说明：______"
        
        elif rule_type == 'graded':
            # 分级评分
            options = [
                "A. 完全达到要求",
                "B. 基本达到要求",
                "C. 部分达到要求",
                "D. 未达到要求"
            ]
            fill_blank_hint = "请说明达到程度：______"
        
        # 提取数据填空提示
        data_hints = []
        if "人数" in rule_text:
            data_hints.append("人数：______人")
        if "次数" in rule_text:
            data_hints.append("次数：______次")
        if "金额" in rule_text or "万元" in rule_text:
            data_hints.append("金额：______万元")
        if "比例" in rule_text or "%" in rule_text:
            data_hints.append("比例：______%")
        if "数量" in rule_text:
            data_hints.append("数量：______个")
        
        if data_hints:
            if fill_blank_hint:
                fill_blank_hint += "；" + "；".join(data_hints)
            else:
                fill_blank_hint = "；".join(data_hints)
        
        return (options, fill_blank_hint, rule_type)
    
    def _extract_multi_item_options(self, rule_text: str) -> List[str]:
        """
        从多项累计规则中提取选项
        
        例如："实现'1-2项'的得1分，实现'3项'的得2分"
        生成：["A. 完成3项", "B. 完成1-2项", "C. 均未完成"]
        """
        options = []
        
        # 查找项数描述
        # 模式1: "完成 X 项"
        pattern1 = r'完成\s*[【\[]?(\d+)[】\]]?\s*项'
        matches1 = re.findall(pattern1, rule_text)
        
        # 模式2: "X-Y项"
        pattern2 = r'(\d+)-(\d+)\s*项'
        matches2 = re.findall(pattern2, rule_text)
        
        # 模式3: "X项及以上"
        pattern3 = r'(\d+)\s*项及以上'
        matches3 = re.findall(pattern3, rule_text)
        
        if matches3:
            # 有"及以上"的情况
            num = matches3[0]
            options = [
                f"A. 完成{num}项及以上",
                f"B. 完成1-{int(num)-1}项",
                "C. 均未完成"
            ]
        elif matches1 and len(matches1) >= 2:
            # 有多个具体项数
            nums = sorted([int(n) for n in matches1], reverse=True)
            if len(nums) >= 2:
                options = [
                    f"A. 完成{nums[0]}项",
                    f"B. 完成{nums[1]}项",
                    "C. 完成更少或均未完成"
                ]
        elif matches2:
            # 有范围描述
            start, end = matches2[0]
            options = [
                f"A. 完成{end}项及以上",
                f"B. 完成{start}-{int(end)-1}项",
                "C. 均未完成"
            ]
        elif '全部完成' in rule_text:
            options = [
                "A. 全部完成",
                "B. 部分完成",
                "C. 均未完成"
            ]
        else:
            # 默认选项
            options = [
                "A. 完成",
                "B. 部分完成",
                "C. 均未完成"
            ]
        
        return options
    
    def _create_basic_info_section(self, doc: Document):
        """创建企业基本信息部分"""
        # 标题
        title = doc.add_paragraph()
        title_run = title.add_run("民营企业现代企业制度指数评价调查问卷")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # 说明文字
        intro = doc.add_paragraph()
        intro_run = intro.add_run("【填写说明】本问卷基于南开大学现代企业制度指数评价体系设计，请根据企业实际情况如实填写。")
        intro_run.font.size = Pt(10)
        intro_run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()
        
        # 制表时间
        p = doc.add_paragraph()
        p.add_run("制表时间：______年____月____日")
        
        doc.add_paragraph()
        
        # 企业基本信息
        p = doc.add_paragraph()
        p.add_run("一、企业基本信息").bold = True
        
        doc.add_paragraph()
        
        # 基本信息字段
        info_fields = [
            "1. 企业名称：______________________",
            "2. 企业主营业务：□ A. 第一产业  □ B. 第二产业  □ C. 第三产业  □ D. 多元化经营",
            "3. 企业规模：□ A. 大型  □ B. 中型  □ C. 小型  □ D. 微型",
            "4. 企业是否为民营企业500强：□ A. 是  □ B. 否",
            "5. 企业成立年限：______年",
            "6. 上一年度营业收入：______万元",
            "7. 上一年度研发投入金额：______万元，占营收比例：______%"
        ]
        
        for field in info_fields:
            doc.add_paragraph(field, style='List Number')
        
        doc.add_paragraph()
        doc.add_page_break()
    
    def _create_question_table(self, doc: Document, questions: List[Dict], level: str):
        """
        创建问题表格
        
        Args:
            doc: Word文档对象
            questions: 问题列表
            level: 级别（初级/中级/高级）
        """
        # 按一级指标分组
        grouped = {}
        for q in questions:
            primary = q['一级指标']
            if primary not in grouped:
                grouped[primary] = []
            grouped[primary].append(q)
        
        # 指标类型映射
        indicator_type_map = {
            '初级': '基础指标',
            '中级': '鼓励指标',
            '高级': '拔高指标'
        }
        
        # 添加级别标题
        level_title = doc.add_paragraph()
        level_run = level_title.add_run(f"二、{level}指标评价问卷（{indicator_type_map[level]}）")
        level_run.font.size = Pt(14)
        level_run.font.bold = True
        
        doc.add_paragraph()
        
        # 为每个一级指标创建表格
        for primary_indicator, qs in grouped.items():
            # 一级指标标题
            section_title = doc.add_paragraph()
            section_run = section_title.add_run(f"【{primary_indicator}】")
            section_run.font.size = Pt(12)
            section_run.font.bold = True
            section_run.font.color.rgb = RGBColor(0, 51, 102)
            
            doc.add_paragraph()
            
            # 创建表格
            table = doc.add_table(rows=1, cols=8)
            table.style = 'Table Grid'
            
            # 表头
            headers = ['序号', '一级指标', '二级指标', '三级指标', '评分准则', '分值', '选项', '补充数据/说明']
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # 设置表头格式
                for paragraph in header_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(255, 255, 255)
                # 设置表头背景色
                try:
                    from docx.oxml import parse_xml
                    shading_elm = parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format(qn('w')))
                    header_cells[i]._element.get_or_add_tcPr().append(shading_elm)
                except:
                    pass  # 如果设置背景色失败，跳过
            
            # 添加问题行
            for q in qs:
                row_cells = table.add_row().cells
                
                # 序号
                row_cells[0].text = str(int(q['序号'])) if pd.notna(q['序号']) else ""
                
                # 一级指标
                row_cells[1].text = str(q['一级指标']) if pd.notna(q['一级指标']) else ""
                
                # 二级指标
                row_cells[2].text = str(q['二级指标']) if pd.notna(q['二级指标']) else ""
                
                # 三级指标
                row_cells[3].text = str(q['三级指标']) if pd.notna(q['三级指标']) else ""
                
                # 评分准则
                scoring_criteria = ""
                for col_name in ['评分准则', '评分标准', '打分标准']:
                    if col_name in q and pd.notna(q[col_name]):
                        scoring_criteria = str(q[col_name])
                        break
                row_cells[4].text = scoring_criteria
                
                # 分值
                score_value = str(q.get('分值', '1')) if pd.notna(q.get('分值')) else "1"
                row_cells[5].text = score_value
                
                # 根据评分准则生成选项
                options, fill_hint, rule_type = self._parse_scoring_criteria(scoring_criteria, score_value)
                row_cells[6].text = '\n'.join(options)
                
                # 补充数据/说明
                supplement = ""
                
                # 添加填空题提示
                if fill_hint:
                    supplement += fill_hint
                
                # 添加佐证材料要求
                if pd.notna(q.get('佐证材料')) and q['佐证材料']:
                    evidence = str(q['佐证材料']).strip()
                    if evidence:
                        if supplement:
                            supplement += "\n\n"
                        supplement += f"【佐证材料】{evidence}"
                
                # 添加规则类型标注（便于理解）
                rule_type_label = {
                    'binary': '二元',
                    'multi_item': '多项',
                    'all_required': '全部',
                    'negative': '否决',
                    'graded': '分级'
                }.get(rule_type, '')
                
                if supplement:
                    supplement += f"\n\n[评分类型：{rule_type_label}]"
                else:
                    supplement = f"[评分类型：{rule_type_label}]"
                
                row_cells[7].text = supplement
                
                # 设置单元格格式
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                            run.font.name = '宋体'
            
            doc.add_paragraph()
    
    def generate_questionnaire(self, level: str, output_path: str = None) -> str:
        """
        生成指定级别的问卷
        
        Args:
            level: 级别（初级/中级/高级）
            output_path: 输出文件路径，如果为None则自动生成
            
        Returns:
            生成的文件路径
        """
        if level not in self.data:
            raise ValueError(f"无效的级别: {level}，必须是 初级/中级/高级 之一")
        
        print(f"\n开始生成{level}问卷...")
        
        # 创建Word文档
        doc = Document()
        
        # 设置默认字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 创建基本信息部分
        self._create_basic_info_section(doc)
        
        # 获取该级别的问题数据
        df = self.data[level]
        questions = df.to_dict('records')
        
        print(f"  - 共 {len(questions)} 个指标")
        
        # 创建问题表格
        self._create_question_table(doc, questions, level)
        
        # 生成输出路径
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_dir = "survey_generator/output/questionnaires"
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            output_path = os.path.join(output_dir, f"调查问卷_{level}_V3_{timestamp}.docx")
        
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 保存文档
        doc.save(output_path)
        print(f"[OK] 成功生成{level}问卷: {output_path}")
        
        return output_path
    
    def generate_all_questionnaires(self, output_dir: str = "survey_generator/output/questionnaires") -> Dict[str, str]:
        """
        生成所有三个级别的问卷
        
        Args:
            output_dir: 输出目录
            
        Returns:
            包含三个级别文件路径的字典
        """
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        results = {}
        
        print("\n" + "="*80)
        print("南开问卷生成器 V3 - 基于评分准则")
        print("="*80)
        
        for level in ['初级', '中级', '高级']:
            output_path = os.path.join(output_dir, f"调查问卷_{level}_V3_{timestamp}.docx")
            results[level] = self.generate_questionnaire(level, output_path)
        
        print("\n" + "="*80)
        print("问卷生成完成！")
        print("="*80)
        
        return results


def main():
    """主函数 - 示例用法"""
    # Excel文件路径
    excel_path = "survey_generator/南开大学-现代企业制度指数评价体系初稿2025.10.22（单独指标）(1).xlsx"
    
    try:
        # 创建生成器
        generator = NankaiQuestionnaireGeneratorV3(excel_path)
        
        # 生成所有级别的问卷
        results = generator.generate_all_questionnaires()
        
        print("\n生成的文件：")
        for level, path in results.items():
            print(f"  {level}: {path}")
        
    except Exception as e:
        print(f"\n[ERROR] 错误: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()