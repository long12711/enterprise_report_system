"""
南开大学现代企业制度指数评价问卷生成器

根据南开大学指标体系Excel文件生成三级（初级、中级、高级）调查问卷
每个问题包含：选择题 + 填空题 + 佐证材料提交
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from datetime import datetime
from typing import Dict, List, Tuple


class NankaiQuestionnaireGenerator:
    """南开大学问卷生成器"""
    
    def __init__(self, excel_path: str):
        """
        初始化生成器
        
        Args:
            excel_path: 南开大学指标体系Excel文件路径
        """
        self.excel_path = excel_path
        self.data = {}
        self._load_data()
        
    def _load_data(self):
        """加载Excel数据"""
        try:
            # 读取三个级别的数据
            self.data['初级'] = pd.read_excel(self.excel_path, sheet_name='初级')
            self.data['中级'] = pd.read_excel(self.excel_path, sheet_name='中级')
            self.data['高级'] = pd.read_excel(self.excel_path, sheet_name='高级')
            print(f"成功加载指标数据：初级 {len(self.data['初级'])} 条，中级 {len(self.data['中级'])} 条，高级 {len(self.data['高级'])} 条")
        except Exception as e:
            raise Exception(f"加载Excel文件失败: {str(e)}")
    
    def _parse_scoring_standard(self, standard: str) -> Tuple[List[str], str]:
        """
        解析打分标准，生成选择题选项和填空题提示
        
        Args:
            standard: 打分标准文本
            
        Returns:
            (选项列表, 填空题提示)
        """
        if pd.isna(standard) or not standard:
            return (["A. 是", "B. 否"], "")
        
        standard = str(standard).strip()
        options = []
        fill_blank_hint = ""
        
        # 解析不同类型的打分标准
        if "完成" in standard and ("项" in standard or "全部" in standard):
            # 类型1: "完成X项得分" 或 "全部完成得分"
            if "全部完成" in standard:
                options = ["A. 全部完成", "B. 部分完成", "C. 均未完成"]
            elif "完成 3 项及以上" in standard or "完成【3】项" in standard:
                options = ["A. 完成 3 项及以上", "B. 完成 1-2 项", "C. 均未完成"]
            elif "完成 4-5 项" in standard:
                options = ["A. 完成 4-5 项", "B. 完成 1-3 项", "C. 均未完成"]
            elif "完成两项" in standard or "完成 2 项" in standard:
                options = ["A. 完成两项", "B. 仅完成 1 项", "C. 均未完成"]
            elif "完成三项" in standard or "完成 3 项" in standard:
                options = ["A. 完成三项", "B. 完成 2 项", "C. 完成 1 项", "D. 均未完成"]
            elif "完成四项" in standard or "完成 4 项" in standard:
                options = ["A. 完成四项", "B. 完成 2-3 项", "C. 完成 1 项", "D. 均未完成"]
            elif "完成五项" in standard or "完成 5 项" in standard:
                options = ["A. 完成五项", "B. 完成 3-4 项", "C. 完成 1-2 项", "D. 均未完成"]
            else:
                options = ["A. 完成", "B. 部分完成", "C. 均未完成"]
                
        elif "达到" in standard and "行业" in standard:
            # 类型2: 达到行业平均值
            options = ["A. 达到行业平均值", "B. 接近行业平均值", "C. 低于行业平均值"]
            
        elif "建立" in standard or "制定" in standard or "设立" in standard:
            # 类型3: 是否建立/制定/设立
            options = ["A. 是", "B. 否"]
            
        elif "近" in standard and "年" in standard:
            # 类型4: 时间范围内的行为
            if "近 5 年" in standard:
                options = ["A. 近 5 年连续满足", "B. 近 5 年部分满足", "C. 未满足"]
            elif "近 3 年" in standard or "近三年" in standard:
                options = ["A. 近 3 年无负面记录", "B. 有负面记录"]
            else:
                options = ["A. 满足要求", "B. 部分满足", "C. 不满足"]
                
        else:
            # 默认选项
            options = ["A. 是", "B. 否"]
        
        # 提取填空题提示（从打分标准中提取关键数据项）
        if "人数" in standard:
            fill_blank_hint += "人数：______人；"
        if "次数" in standard:
            fill_blank_hint += "次数：______次；"
        if "金额" in standard or "万元" in standard:
            fill_blank_hint += "金额：______万元；"
        if "比例" in standard or "%" in standard:
            fill_blank_hint += "比例：______%；"
        if "数量" in standard:
            fill_blank_hint += "数量：______个；"
            
        return (options, fill_blank_hint)
    
    def _create_basic_info_section(self, doc: Document):
        """创建企业基本信息部分"""
        # 标题
        title = doc.add_paragraph()
        title_run = title.add_run("民营企业现代企业制度指数评价调查问卷")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # 制表时间
        p = doc.add_paragraph()
        p.add_run("制表时间：______年____月")
        
        doc.add_paragraph()
        
        # 企业基本信息
        p = doc.add_paragraph()
        p.add_run("企业基本信息：").bold = True
        
        doc.add_paragraph()
        
        # 基本信息字段
        info_fields = [
            "企业主营业务：A. 第一产业 B. 第二产业 C. 第三产业 D. 多元化经营",
            "企业规模：A. 大型 B. 中型 C. 小型 D. 微型",
            '企业是否为民营企业 500 强：A. 是 B. 否（企业规模选 "大型" 时必填）',
            "企业成立年限：______年 5. 上一年度营业收入：______万元 6. 上一年度研发投入金额：万元，占营收比例：%"
        ]
        
        for field in info_fields:
            doc.add_paragraph(field)
        
        doc.add_paragraph()
    
    def _create_question_table(self, doc: Document, questions: List[Dict], level: str):
        """
        创建问题表格
        
        Args:
            doc: Word文档对象
            questions: 问题列表
            level: 级别（初级/中级/高级）
        """
        # 按一级指标分组
        grouped = {}
        for q in questions:
            primary = q['一级指标']
            if primary not in grouped:
                grouped[primary] = []
            grouped[primary].append(q)
        
        # 指标类型映射
        indicator_type_map = {
            '初级': '基础指标',
            '中级': '鼓励指标',
            '高级': '拔高指标'
        }
        
        # 添加级别标题
        level_title = doc.add_paragraph()
        level_run = level_title.add_run(f"{level}指标评价问卷（{indicator_type_map[level]}）")
        level_run.font.size = Pt(14)
        level_run.font.bold = True
        level_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # 为每个一级指标创建表格
        for primary_indicator, qs in grouped.items():
            # 一级指标标题
            section_title = doc.add_paragraph()
            section_run = section_title.add_run(str(primary_indicator) if pd.notna(primary_indicator) else '')
            section_run.font.size = Pt(12)
            section_run.font.bold = True
            
            doc.add_paragraph()
            
            # 创建表格（添加评分准则列）
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            
            # 表头
            headers = ['序号', '一级指标', '二级指标', '三级指标', '评分准则', '选项', '补充数据 / 说明']
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # 设置表头格式
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
            
            # 添加问题行
            for q in qs:
                row_cells = table.add_row().cells
                
                # 序号
                row_cells[0].text = str(q['序号'])
                
                # 一级指标
                row_cells[1].text = str(q['一级指标'])
                
                # 二级指标
                row_cells[2].text = str(q['二级指标'])
                
                # 三级指标
                row_cells[3].text = str(q['三级指标'])
                
                # 评分准则（新增列）
                scoring_criteria = ""
                for col_name in ['评分准则', '评分标准', '打分标准']:
                    if col_name in q and not pd.isna(q[col_name]):
                        scoring_criteria = str(q[col_name])
                        break
                row_cells[4].text = scoring_criteria
                
                # 选项（从打分标准解析）
                options, fill_hint = self._parse_scoring_standard(q['打分标准'])
                row_cells[5].text = ' '.join(options)
                
                # 补充数据/说明
                supplement = ""
                
                # 添加填空题提示
                if fill_hint:
                    supplement += fill_hint
                
                # 添加佐证材料要求
                if not pd.isna(q['佐证材料']) and q['佐证材料']:
                    evidence = str(q['佐证材料']).strip()
                    if evidence:
                        if supplement:
                            supplement += "\n\n"
                        supplement += f"佐证材料：{evidence}"
                
                row_cells[6].text = supplement
                
                # 设置单元格格式
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
            
            doc.add_paragraph()
    
    def generate_questionnaire(self, level: str, output_path: str = None) -> str:
        """
        生成指定级别的问卷
        
        Args:
            level: 级别（初级/中级/高级）
            output_path: 输出文件路径，如果为None则自动生成
            
        Returns:
            生成的文件路径
        """
        if level not in self.data:
            raise ValueError(f"无效的级别: {level}，必须是 初级/中级/高级 之一")
        
        # 创建Word文档
        doc = Document()
        
        # 设置默认字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 创建基本信息部分
        self._create_basic_info_section(doc)
        
        # 获取该级别的问题数据
        df = self.data[level]
        questions = df.to_dict('records')
        
        # 创建问题表格
        self._create_question_table(doc, questions, level)
        
        # 生成输出路径
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"调查问卷_{level}_{timestamp}.docx"
        
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 保存文档
        doc.save(output_path)
        print(f"成功生成{level}问卷: {output_path}")
        
        return output_path
    
    def generate_all_questionnaires(self, output_dir: str = "output/questionnaires") -> Dict[str, str]:
        """
        生成所有三个级别的问卷
        
        Args:
            output_dir: 输出目录
            
        Returns:
            包含三个级别文件路径的字典
        """
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        results = {}
        
        for level in ['初级', '中级', '高级']:
            output_path = os.path.join(output_dir, f"调查问卷_{level}_{timestamp}.docx")
            results[level] = self.generate_questionnaire(level, output_path)
        
        return results


def main():
    """主函数 - 示例用法"""
    # Excel文件路径（使用相对于项目根目录的路径）
    excel_path = "../report_generator/南开大学-现代企业制度指数评价体系初稿2025.10.22（单独指标）(1).xlsx"
    
    # 创建生成器
    generator = NankaiQuestionnaireGenerator(excel_path)
    
    # 生成所有级别的问卷
    results = generator.generate_all_questionnaires()
    
    print("\n问卷生成完成！")
    for level, path in results.items():
        print(f"{level}: {path}")


if __name__ == "__main__":
    main()