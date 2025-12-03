"""
南开大学现代企业制度指数评价问卷生成器 V2
仿照参考文件格式生成问卷
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime
from typing import Dict, List, Tuple


class NankaiQuestionnaireGeneratorV2:
    """南开大学问卷生成器 V2 - 仿照参考格式"""
    
    def __init__(self, excel_path: str):
        """
        初始化生成器
        
        Args:
            excel_path: 南开大学指标体系Excel文件路径
        """
        self.excel_path = excel_path
        self.data = {}
        self._load_data()
        
    def _load_data(self):
        """加载Excel数据"""
        try:
            # 读取三个级别的数据
            self.data['初级'] = pd.read_excel(self.excel_path, sheet_name='初级')
            self.data['中级'] = pd.read_excel(self.excel_path, sheet_name='中级')
            self.data['高级'] = pd.read_excel(self.excel_path, sheet_name='高级')
            print(f"成功加载指标数据：初级 {len(self.data['初级'])} 条，中级 {len(self.data['中级'])} 条，高级 {len(self.data['高级'])} 条")
        except Exception as e:
            raise Exception(f"加载Excel文件失败: {str(e)}")
    
    def _set_cell_border(self, cell, **kwargs):
        """设置单元格边框"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # 创建边框元素
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'single')
            edge_element.set(qn('w:sz'), '4')
            edge_element.set(qn('w:space'), '0')
            edge_element.set(qn('w:color'), '000000')
            tcBorders.append(edge_element)
        
        tcPr.append(tcBorders)
    
    def _create_basic_info_section(self, doc: Document):
        """创建企业基本信息部分"""
        # 标题
        title = doc.add_paragraph()
        title_run = title.add_run("民营企业现代企业制度指数评价调查问卷")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.name = '宋体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # 制表时间
        p = doc.add_paragraph()
        run = p.add_run("制表时间：______年____月")
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        doc.add_paragraph()
        
        # 企业基本信息标题
        p = doc.add_paragraph()
        run = p.add_run("企业基本信息：")
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        doc.add_paragraph()
        
        # 基本信息字段
        info_fields = [
            "企业主营业务：A. 第一产业 B. 第二产业 C. 第三产业 D. 多元化经营",
            "企业规模：A. 大型 B. 中型 C. 小型 D. 微型",
            '企业是否为民营企业 500 强：A. 是 B. 否（企业规模选 "大型" 时必填）',
            "企业成立年限：______年 5. 上一年度营业收入：______万元 6. 上一年度研发投入金额：万元，占营收比例：%"
        ]
        
        for field in info_fields:
            p = doc.add_paragraph()
            run = p.add_run(field)
            run.font.name = '宋体'
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        doc.add_paragraph()
        doc.add_paragraph()
    
    def _create_question_table(self, doc: Document, questions: List[Dict], level: str):
        """
        创建问题表格 - 显示完整的题目和打分标准
        
        Args:
            doc: Word文档对象
            questions: 问题列表
            level: 级别（初级/中级/高级）
        """
        # 按一级指标分组 - 需要处理NaN值
        grouped = {}
        current_primary = None
        
        for q in questions:
            # 如果当前行有一级指标，更新当前一级指标
            if pd.notna(q['一级指标']):
                current_primary = q['一级指标']
            
            # 使用当前一级指标进行分组
            if current_primary:
                if current_primary not in grouped:
                    grouped[current_primary] = []
                grouped[current_primary].append(q)
        
        # 指标类型映射
        indicator_type_map = {
            '初级': '基础指标',
            '中级': '鼓励指标',
            '高级': '拔高指标'
        }
        
        # 添加级别标题
        level_title = doc.add_paragraph()
        level_run = level_title.add_run(f"{level}指标评价问卷（{indicator_type_map[level]}）")
        level_run.font.size = Pt(14)
        level_run.font.bold = True
        level_run.font.name = '宋体'
        level_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        level_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # 为每个一级指标创建表格
        for primary_indicator, qs in grouped.items():
            # 一级指标标题
            section_title = doc.add_paragraph()
            section_run = section_title.add_run(str(primary_indicator) if pd.notna(primary_indicator) else '')
            section_run.font.size = Pt(12)
            section_run.font.bold = True
            section_run.font.name = '宋体'
            section_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            doc.add_paragraph()
            
            # 创建表格 - 7列：序号、一级指标、二级指标、三级指标、打分标准、选项、补充数据/说明
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            
            # 表头
            headers = ['序号', '一级指标', '二级指标', '三级指标', '打分标准', '选项', '补充数据 / 说明']
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # 设置表头格式
                for paragraph in header_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                self._set_cell_border(header_cells[i])
            
            # 添加问题行
            for q in qs:
                row_cells = table.add_row().cells
                
                # 序号
                row_cells[0].text = str(int(q['序号'])) if pd.notna(q['序号']) else ''
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 一级指标
                row_cells[1].text = str(q['一级指标']) if pd.notna(q['一级指标']) else ''
                
                # 二级指标
                row_cells[2].text = str(q['二级指标']) if pd.notna(q['二级指标']) else ''
                
                # 三级指标
                row_cells[3].text = str(q['三级指标']) if pd.notna(q['三级指标']) else ''
                
                # 打分标准 - 显示完整的打分标准
                standard_text = str(q['打分标准']) if pd.notna(q['打分标准']) else ''
                row_cells[4].text = standard_text
                
                # 选项 - 从打分标准转换为简化选项
                options_text = self._parse_standard_to_options(q['打分标准'])
                row_cells[5].text = options_text
                
                # 补充数据/说明 - 从打分标准中提取填空项
                fill_text = self._extract_fill_blanks_from_standard(q['打分标准'])
                row_cells[6].text = fill_text if fill_text else ""
                
                # 设置单元格格式
                for i, cell in enumerate(row_cells):
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    self._set_cell_border(cell)
            
            doc.add_paragraph()
            doc.add_paragraph()
    
    def _parse_standard_to_options(self, standard: str) -> str:
        """
        将打分标准转换为简化的选项格式
        
        Args:
            standard: 打分标准文本
            
        Returns:
            简化的选项文本
        """
        if pd.isna(standard) or not standard:
            return "A. 符合 B. 部分符合 C. 不符合"
        
        standard = str(standard).strip()
        
        # 尝试从打分标准中提取选项描述
        import re
        
        # 检查是否包含明确的选项描述
        if 'A.' in standard and 'B.' in standard:
            # 提取A、B、C选项
            options = []
            for letter in ['A', 'B', 'C', 'D']:
                pattern = f'{letter}\\.([^A-D]+?)(?=[A-D]\\.|$)'
                match = re.search(pattern, standard, re.DOTALL)
                if match:
                    option_text = match.group(1).strip()
                    # 清理选项文本
                    option_text = re.sub(r'\s+', ' ', option_text)
                    option_text = option_text[:100]  # 限制长度
                    options.append(f"{letter}. {option_text}")
            
            if options:
                return ' '.join(options)
        
        # 默认返回统一格式
        return "A. 符合 B. 部分符合 C. 不符合"
    
    def _extract_fill_blanks_from_standard(self, standard: str) -> str:
        """
        从打分标准中提取需要填写的数据项
        
        Args:
            standard: 打分标准文本
            
        Returns:
            填空题文本
        """
        if pd.isna(standard) or not standard:
            return ""
        
        standard = str(standard).strip()
        fill_items = []
        
        import re
        
        # 识别各种数据类型
        if '党员' in standard and '人数' in standard:
            fill_items.append('党员人数：______人（仅 A/B 选项填写）')
        
        if '支部党员大会' in standard or '支委会' in standard or '党小组会' in standard or '党课' in standard:
            parts = []
            if '支部党员大会' in standard:
                parts.append('上一年度支部党员大会召开次数：______次')
            if '支委会' in standard:
                parts.append('支委会：______次')
            if '党小组会' in standard:
                parts.append('党小组会：______次')
            if '党课' in standard:
                parts.append('党课：______次')
            if parts:
                fill_items.append('；'.join(parts))
        
        if '经费' in standard and '场所' in standard:
            fill_items.append('上一年度党组织活动经费金额：元；活动场所面积：㎡（仅 A/B 选项填写）')
        elif '经费' in standard or '金额' in standard:
            fill_items.append('上一年度党组织活动经费金额：______元')
        elif '场所' in standard or '面积' in standard:
            fill_items.append('活动场所面积：______㎡')
        
        if '记录' in standard and '份数' in standard:
            fill_items.append('相关行为记录份数：______份；涉及事项数量：______项')
        
        if '认同度' in standard:
            fill_items.append('参与认同度调查员工人数：______人；认同人数：人；认同度：%')
        
        if '培训' in standard and '法治' in standard:
            fill_items.append('管理层法治培训年度次数：______次；培训时长：______小时')
        
        if '融入' in standard and ('战略' in standard or '生产' in standard or '培训' in standard or '考核' in standard):
            fill_items.append('融入环节：□战略管理 □生产经营 □员工培训 □考核评价；对应环节实施频次：______')
        
        # 通用识别
        if not fill_items:
            if '人数' in standard:
                fill_items.append('人数：______人')
            if '次数' in standard:
                fill_items.append('次数：______次')
            if '金额' in standard or '万元' in standard:
                fill_items.append('金额：______万元')
            if '比例' in standard or '%' in standard:
                fill_items.append('比例：______%')
        
        return '；'.join(fill_items) if fill_items else ''
    
    def generate_questionnaire(self, level: str, output_path: str = None) -> str:
        """
        生成指定级别的问卷
        
        Args:
            level: 级别（初级/中级/高级）
            output_path: 输出文件路径，如果为None则自动生成
            
        Returns:
            生成的文件路径
        """
        if level not in self.data:
            raise ValueError(f"无效的级别: {level}，必须是 初级/中级/高级 之一")
        
        # 创建Word文档
        doc = Document()
        
        # 设置默认字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 创建基本信息部分
        self._create_basic_info_section(doc)
        
        # 获取该级别的问题数据
        df = self.data[level]
        questions = df.to_dict('records')
        
        # 创建问题表格
        self._create_question_table(doc, questions, level)
        
        # 生成输出路径
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"调查问卷_{level}_{timestamp}.docx"
        
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 保存文档
        doc.save(output_path)
        print(f"成功生成{level}问卷: {output_path}")
        
        return output_path
    
    def generate_all_questionnaires(self, output_dir: str = "output/questionnaires") -> Dict[str, str]:
        """
        生成所有三个级别的问卷
        
        Args:
            output_dir: 输出目录
            
        Returns:
            包含三个级别文件路径的字典
        """
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        results = {}
        
        for level in ['初级', '中级', '高级']:
            output_path = os.path.join(output_dir, f"调查问卷_{level}_{timestamp}.docx")
            results[level] = self.generate_questionnaire(level, output_path)
        
        return results


def main():
    """主函数 - 示例用法"""
    # Excel文件路径
    excel_path = "南开大学-现代企业制度指数评价体系初稿2025.10.22（单独指标）(1).xlsx"
    
    # 创建生成器
    generator = NankaiQuestionnaireGeneratorV2(excel_path)
    
    # 生成所有级别的问卷
    results = generator.generate_all_questionnaires()
    
    print("\n问卷生成完成！")
    for level, path in results.items():
        print(f"{level}: {path}")


if __name__ == "__main__":
    main()